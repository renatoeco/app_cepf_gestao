import streamlit as st
import pandas as pd
import time
import datetime
import os
import uuid

import streamlit_shadcn_ui as ui
from st_rsuite import date_picker

# Geração de docx
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


from funcoes_auxiliares import (
    conectar_mongo_cepf_gestao,
    sidebar_projeto,

    # Google Drive
    obter_servico_drive,
    obter_ou_criar_pasta,
    obter_pasta_projeto,
    obter_pasta_recibos,
    enviar_arquivo_drive,
    gerar_link_drive,
    valor_por_extenso,
    numero_ordinal_pt,
    data_extenso_pt,
    enviar_email
)





st.set_page_config(page_title="Financeiro", page_icon=":material/payments:")





# ###################################################################################################
# SIDEBAR DA PÁGINA DO PROJETO
# ###################################################################################################

sidebar_projeto()



###########################################################################################################
# CONFIGURAÇÕES DO STREAMLIT
###########################################################################################################


# Traduzindo o texto do st.file_uploader
# Texto interno
st.markdown("""
<style>
/* Esconde o texto padrão */
[data-testid="stFileUploaderDropzone"] div div::before {
    content: "";
    color: rgba(49, 51, 63, 0.7);
    font-size: 0.9rem;
    font-weight: 400;
    position: absolute;
    top: 50px;              /* fixa no topo */
    left: 50%;
    transform: translate(-50%, 10%);
    pointer-events: none;
}
/* Esconde o texto original */
[data-testid="stFileUploaderDropzone"] div div span {
    visibility: hidden !important;
}
</style>
""", unsafe_allow_html=True)

# Traduzindo Botão do file_uploader
st.markdown("""
<style>
/* Alvo: apenas o botão dentro do componente de upload */
section[data-testid="stFileUploaderDropzone"] button[data-testid="stBaseButton-secondary"] {
    font-size: 0px !important;   /* esconde o texto original */
    padding-left: 14px !important;
    padding-right: 14px !important;
    min-width: 160px !important;
}
/* Insere o texto traduzido */
section[data-testid="stFileUploaderDropzone"] button[data-testid="stBaseButton-secondary"]::after {
    content: "Selecionar arquivo";
    font-size: 14px !important;
    color: inherit;
}
</style>
""", unsafe_allow_html=True)



###########################################################################################################
# CONEXÃO COM O BANCO DE DADOS MONGODB
###########################################################################################################

# Conecta-se ao banco de dados MongoDB (usa cache automático para melhorar performance)
db = conectar_mongo_cepf_gestao()

# Define as coleções específicas que serão utilizadas a partir do banco
col_projetos = db["projetos"]

col_organizacoes = db["organizacoes"]

col_editais = db["editais"]

col_ciclos = db["ciclos_investimento"]

col_investidores = db["investidores"]


###########################################################################################################
# TRATAMENTO DE DADOS
###########################################################################################################

# Logo hospedada no site do IEB para renderizar nos e-mails.
logo_cepf = "https://cepfcerrado.iieb.org.br/wp-content/uploads/2025/02/LogoConjuntaCEPFIEBGREEN-768x140.png"



codigo_projeto_atual = st.session_state.get("projeto_atual")

if not codigo_projeto_atual:
    st.error("Nenhum projeto selecionado.")
    st.stop()

# Capturando o projeto atual no bd
df_projeto = pd.DataFrame(
    list(
        col_projetos.find(
            {"codigo": codigo_projeto_atual}
        )
    )
)

if df_projeto.empty:
    st.error("Projeto não encontrado no banco de dados.")
    st.stop()


# Transformar o id em string
df_projeto = df_projeto.copy()
if "_id" in df_projeto.columns:
    df_projeto["_id"] = df_projeto["_id"].astype(str)


projeto = df_projeto.iloc[0]

# Capturando o financeiro no banco de dados
financeiro = projeto.get("financeiro", {})


# Organizações

df_organizacoes = pd.DataFrame(
    list(
        col_organizacoes.find()
    )
)

mapa_org_id_nome = {
    row["_id"]: row["nome_organizacao"]
    for _, row in df_organizacoes.iterrows()
}

###########################################################################################################
# FUNÇÕES
###########################################################################################################



# -----------------------------------
# Formata float para moeda R$
# -----------------------------------
def format_brl(valor):
    return (
        f"R$ {valor:,.2f}"
        .replace(",", "X")
        .replace(".", ",")
        .replace("X", ".")
    ) if valor else ""





# Função para notificar os usuários internos quando os valores das parcelas ficam desatualizados
# em decorrência do cadastro de aditivo ou devolução, que alteram o valor do projeto.
def notifica_parcelas_desencontradas():

    # --------------------------------------------------
    # NOTIFICAÇÃO: valor total ajustado diferente das parcelas
    # --------------------------------------------------
    if usuario_interno:


        # -----------------------------------
        # Parcelas cadastradas
        # -----------------------------------
        parcelas = financeiro.get("parcelas", [])

        # -----------------------------------
        # Só executa validação se houver parcelas cadastradas
        # -----------------------------------
        if not parcelas:
            return


        # -----------------------------------
        # Recuperar valores financeiros
        # -----------------------------------
        valor_total_base = financeiro.get("valor_total") or 0.0
        valor_aditivo = financeiro.get("valor_aditivo") or 0.0
        valor_devolucao = financeiro.get("valor_devolucao") or 0.0

        # -----------------------------------
        # Calcular valor total ajustado
        # -----------------------------------
        valor_total_ajustado = valor_total_base + valor_aditivo - valor_devolucao



        soma_parcelas = sum(
            p.get("valor", 0)
            for p in parcelas
            if p.get("valor") is not None
        )


        # -----------------------------------
        # Verificar inconsistência
        # -----------------------------------
        if round(soma_parcelas, 2) != 0 and round(soma_parcelas, 2) != round(valor_total_ajustado, 2):

            st.warning(
                "O valor total das parcelas está diferente do valor total atualizado do projeto após ajustes financeiros. **Atualize o cronograma de parcelas.**",
                icon=":material/warning:"
            )





# ==================================================
# Aprova remanejamento manualmente
# • muda status para aceito
# • grava data
# • efetiva orçamento
# ==================================================
def aprovar_remanejamento(
    codigo_projeto,
    idx,
    item
):
    """
    Aprova o remanejamento selecionado.
    """

    projeto = col_projetos.find_one({"codigo": codigo_projeto})
    financeiro = projeto.get("financeiro", {})
    lista = financeiro.get("remanejamentos_financeiros", [])

    if idx >= len(lista):
        return

    # --------------------------------------------------
    # Atualiza status + data
    # --------------------------------------------------
    lista[idx]["status_remanejamento"] = "aceito"
    lista[idx]["data_aprov_remanej"] = datetime.datetime.now(datetime.UTC)

    col_projetos.update_one(
        {"codigo": codigo_projeto},
        {
            "$set": {
                "financeiro.remanejamentos_financeiros": lista
            }
        }
    )

    # --------------------------------------------------
    # Efetiva impacto no orçamento
    # --------------------------------------------------
    efetivar_remanejamento(
        col_projetos,
        codigo_projeto,
        financeiro,
        lista[idx].get("reduzidas", []),
        lista[idx].get("aumentadas", [])
    )


    projeto_atualizado = col_projetos.find_one({"codigo": codigo_projeto})
    enviar_email_remanejamento_aprovado(
        projeto_atualizado,
        lista[idx]
    )



    st.rerun()





# ==================================================
# Atualiza aceite técnico/financeiro imediatamente, na análise do remanejamento
# ==================================================
# ==================================================
# Atualiza aceite técnico/financeiro do remanejamento
# (mesmo padrão usado nos relatórios)
# ==================================================
def atualizar_aceite_remanejamento(
    projeto_codigo,
    idx,
    campo,          # "aceite_tecnico" ou "aceite_financeiro"
    checkbox_key
):
    """
    Disparado automaticamente quando checkbox muda.

    • Lê estado real do checkbox via session_state
    • Atualiza item do array em memória
    • Regrava lista completa no Mongo (forma segura)
    """

    marcado = st.session_state.get(checkbox_key, False)

    nome = st.session_state.get("nome", "Usuário")
    data = datetime.datetime.now().strftime("%d/%m/%Y")

    # --------------------------------------------------
    # Carrega documento atualizado do banco
    # --------------------------------------------------
    projeto = col_projetos.find_one({"codigo": projeto_codigo})
    financeiro = projeto.get("financeiro", {})
    lista = financeiro.get("remanejamentos_financeiros", [])

    if idx >= len(lista):
        return

    # --------------------------------------------------
    # Atualiza item em memória
    # --------------------------------------------------
    if marcado:
        lista[idx][campo] = f"Aceito por {nome} em {data}"
    else:
        lista[idx].pop(campo, None)

    # --------------------------------------------------
    # Salva lista inteira (Mongo não atualiza por índice)
    # --------------------------------------------------
    col_projetos.update_one(
        {"codigo": projeto_codigo},
        {
            "$set": {
                "financeiro.remanejamentos_financeiros": lista
            }
        }
    )







# ==================================================
# Envia e-mail para o beneficiárioquando remanejamento é recusado
# ==================================================
def enviar_email_remanejamento_recusado(
    projeto,
    item_remanejamento
):
    """
    Envia e-mail para todos os contatos cadastrados
    na chave 'contatos' do projeto quando o
    remanejamento for recusado.
    """

    contatos = projeto.get("contatos", [])

    # --------------------------------------------------
    # Coletar e-mails
    # --------------------------------------------------
    destinatarios = [
        c.get("email")
        for c in contatos
        if c.get("email")
    ]

    if not destinatarios:
        return

    codigo = projeto.get("codigo")
    nome_projeto = projeto.get("nome_do_projeto")

    # -----------------------------------
    # Recupera nome da organização via mapa global
    # -----------------------------------
    organizacao = mapa_org_id_nome.get(
        projeto.get("id_organizacao"),
        ""
    )


    reduzidas = item_remanejamento.get("reduzidas", [])
    aumentadas = item_remanejamento.get("aumentadas", [])
    justificativa = item_remanejamento.get("justificativa", "")
    motivo_recusa = item_remanejamento.get("motivo_recusa", "")

    total_reduzido = sum(r.get("valor_reduzido", 0) for r in reduzidas)

    # --------------------------------------------------
    # Montar listas HTML
    # --------------------------------------------------
    def montar_lista_html(itens, campo_valor):
        if not itens:
            return "<p>Nenhuma</p>"

        linhas = ""
        for i in itens:
            linhas += f"<li>{i['nome_despesa']}: {format_brl(i[campo_valor])}</li>"

        return f"<ul>{linhas}</ul>"

    lista_reduzidas = montar_lista_html(reduzidas, "valor_reduzido")
    lista_aumentadas = montar_lista_html(aumentadas, "valor_aumentado")

    assunto = "Remanejamento recusado"

    logo = logo_cepf

    corpo_html = f"""
    <!DOCTYPE html>
    <html>
    <head>
    <meta charset="utf-8">
    <style>

    body {{
        font-family: Arial, Helvetica, sans-serif;
        background-color: #f5f5f5;
    }}

    .container {{
        max-width: 760px;
        margin: 0 auto;
        background: white;
        border-top: 6px solid #C82333; /* vermelho */
        padding: 30px;
    }}

    .logo {{
        text-align: center;
        margin-bottom: 20px;
    }}

    .highlight {{
        color: #C82333;
        font-weight: bold;
    }}

    ul {{
        list-style: none;
        padding-left: 0;
        margin: 0;
    }}

    li {{
        margin: 0;
        padding: 0;
    }}

    </style>
    </head>

    <body>

    <div class="container">

        <div class="logo">
            <img src="{logo}" height="60">
        </div>

        <p>
        Foi <span class="highlight"><strong>recusado</strong></span>
        o remanejamento de
        <span class="highlight">{format_brl(total_reduzido)}</span>
        no orçamento do projeto
        <span class="highlight">{codigo} - {nome_projeto}</span>
        da organização
        <span class="highlight">{organizacao}</span>,
        conforme detalhado a seguir:
        </p>

        <br>

        <table width="100%">
        <tr>
            <td valign="top" width="50%">
                <strong>Despesas reduzidas</strong>
                {lista_reduzidas}
            </td>

            <td valign="top" width="50%">
                <strong>Despesas aumentadas</strong>
                {lista_aumentadas}
            </td>
        </tr>
        </table>

        <br>

        <p><strong>Justificativa original:</strong></p>
        <p>{justificativa}</p>

        <br>

        <p><strong>Motivo da recusa:</strong></p>
        <p>{motivo_recusa}</p>

        <br>
        <p>Sistema Veredas</p>

    </div>

    </body>
    </html>
    """

    enviar_email(corpo_html, destinatarios, assunto)







# ==================================================
# Envia e-mail para o beneficiário quando remanejamento é aprovado
# ==================================================
def enviar_email_remanejamento_aprovado(
    projeto,
    item_remanejamento
):
    """
    Envia e-mail para todos os contatos cadastrados
    na chave 'contatos' do projeto.
    """

    contatos = projeto.get("contatos", [])

    # --------------------------------------------------
    # Coletar e-mails dos contatos
    # --------------------------------------------------
    destinatarios = [
        c.get("email")
        for c in contatos
        if c.get("email")
    ]

    if not destinatarios:
        return

    codigo = projeto.get("codigo")
    nome_projeto = projeto.get("nome_do_projeto")
    
    organizacao = mapa_org_id_nome.get(projeto.get("id_organizacao"), "")
    
    reduzidas = item_remanejamento.get("reduzidas", [])
    aumentadas = item_remanejamento.get("aumentadas", [])
    justificativa = item_remanejamento.get("justificativa", "")

    # --------------------------------------------------
    # Soma total reduzido
    # --------------------------------------------------
    total_reduzido = sum(r.get("valor_reduzido", 0) for r in reduzidas)

    # --------------------------------------------------
    # Montar listas HTML (sem bolinha)
    # --------------------------------------------------
    def montar_lista_html(itens, campo_valor):
        if not itens:
            return "<p>Nenhuma</p>"

        linhas = ""
        for i in itens:
            linhas += f"<li>{i['nome_despesa']}: {format_brl(i[campo_valor])}</li>"

        return f"<ul>{linhas}</ul>"

    lista_reduzidas = montar_lista_html(reduzidas, "valor_reduzido")
    lista_aumentadas = montar_lista_html(aumentadas, "valor_aumentado")

    # --------------------------------------------------
    # Assunto
    # --------------------------------------------------
    assunto = "Remanejamento aprovado"

    logo = logo_cepf

    corpo_html = f"""
    <!DOCTYPE html>
    <html>
    <head>
    <meta charset="utf-8">
    <style>

    body {{
        font-family: Arial, Helvetica, sans-serif;
        background-color: #f5f5f5;
    }}

    .container {{
        max-width: 760px;
        margin: 0 auto;
        background: white;
        border-top: 6px solid #A0C256;
        padding: 30px;
    }}

    .logo {{
        text-align: center;
        margin-bottom: 20px;
    }}

    .highlight {{
        color: #A0C256;
        font-weight: bold;
    }}

    ul {{
        list-style: none;
        padding-left: 0;
        margin: 0;
    }}

    li {{
        margin: 0;
        padding: 0;
    }}

    </style>
    </head>

    <body>

    <div class="container">

        <div class="logo">
            <img src="{logo}" height="60">
        </div>

        <p>
        Foi <span class="highlight"><strong>aprovado</strong></span>
        o remanejamento de
        <span class="highlight">{format_brl(total_reduzido)}</span>
        no orçamento do projeto
        <span class="highlight">{codigo} - {nome_projeto}</span>
        da organização
        <span class="highlight">{organizacao}</span>,
        conforme detalhado a seguir:
        </p>

        <br>

        <table width="100%">
        <tr>
            <td valign="top" width="50%">
                <strong>Despesas reduzidas</strong>
                {lista_reduzidas}
            </td>

            <td valign="top" width="50%">
                <strong>Despesas aumentadas</strong>
                {lista_aumentadas}
            </td>
        </tr>
        </table>

        <br>

        <p><strong>Justificativa:</strong></p>
        <p>{justificativa}</p>

        <br>

        <p>
        O orçamento do projeto já está atualizado no sistema.
        </p>

        <br>
        <p>Sistema Veredas</p>

    </div>

    </body>
    </html>
    """

    enviar_email(corpo_html, destinatarios, assunto)






# ==================================================
# Envia notificação de remanejamento para equipe/admin
# ==================================================
def enviar_email_remanejamento(
    db,
    codigo_projeto,
    sigla,
    nome_projeto,
    organizacao,
    reduzidas,
    aumentadas,
    status_remanejamento
):
    """
    Envia e-mail para:
    • tipo_usuario in ["admin", "equipe"]
    • pessoa vinculada ao projeto
    """

    col_pessoas = db["pessoas"]

    # --------------------------------------------------
    # Buscar destinatários
    # --------------------------------------------------
    pessoas = list(
        col_pessoas.find({
            "status": "ativo",
            "tipo_usuario": {"$in": ["admin", "equipe"]},
            "projetos": codigo_projeto
        })
    )

    if not pessoas:
        return

    emails = [p["e_mail"] for p in pessoas if p.get("e_mail")]

    # --------------------------------------------------
    # Montar tabelas HTML
    # --------------------------------------------------
    def montar_lista_html(itens, campo_valor):
        if not itens:
            return "<p>Nenhuma</p>"

        linhas = ""
        for i in itens:
            linhas += f"<li>{i['nome_despesa']}: {format_brl(i[campo_valor])}</li>"

        return f"<ul>{linhas}</ul>"



    lista_reduzidas = montar_lista_html(reduzidas, "valor_reduzido")
    lista_aumentadas = montar_lista_html(aumentadas, "valor_aumentado")

    # --------------------------------------------------
    # Mensagem condicional
    # --------------------------------------------------
    if status_remanejamento == "aceito":
        mensagem_status = (
            "O remanejamento <strong>foi aceito automaticamente</strong> e o orçamento já está atualizado."
        )
    else:
        mensagem_status = (
            "<b>AÇÃO NECESSÁRIA: Esse remanejamento depende de análise e aprovação</b><br><br>"
            "Visite a página de remanejamentos no Sistema de Gestão de Projetos para dar continuidade."
        )

    # --------------------------------------------------
    # Assunto
    # --------------------------------------------------
    assunto = f"Solicitação de remanejamento - {codigo_projeto} - {sigla}"

    # --------------------------------------------------
    # Enviar para cada pessoa (personalizado)
    # --------------------------------------------------
    logo = logo_cepf

    with st.spinner("Enviando..."):


        for pessoa in pessoas:

            nome = pessoa.get("nome_completo", "").split()[0]

            corpo_html = f"""
            <!DOCTYPE html>
            <html>
            <head>
            <meta charset="utf-8">
            <style>

            body {{
                font-family: Arial, Helvetica, sans-serif;
                background-color: #f5f5f5;
            }}

            .container {{
                max-width: 760px;   /* mais largo */
                margin: 0 auto;
                background: white;
                border-top: 6px solid #A0C256;
                padding: 30px;
            }}

            .logo {{
                text-align: center;   /* centraliza imagem */
                margin-bottom: 20px;
            }}

            .highlight {{
                color: #A0C256;
                font-weight: bold;
            }}

            /* remove bolinhas e espaçamentos */
            ul {{
                list-style: none;
                padding-left: 0;
                margin: 0;
            }}

            li {{
                margin: 0;
                padding: 0;
            }}

            </style>
            </head>

            <body>

            <div class="container">

                <br>

                <div class="logo">
                    <img src="{logo}" height="60">
                </div>

                <br>

                <p>Olá <strong>{nome}</strong>,</p>

                <p>
                O projeto <span class="highlight">{codigo_projeto} - {sigla} - {nome_projeto}</span>,
                da organização <span class="highlight">{organizacao}</span>,
                enviou uma nova solicitação de remanejamento financeiro.
                </p>

                <br>

                <p><strong>Resumo do remanejamento:</strong></p>

                <table width="100%">
                <tr>
                    <td valign="top" width="50%">
                    <strong>Reduções</strong>
                    {lista_reduzidas}
                    </td>

                    <td valign="top" width="50%">
                    <strong>Aumentos</strong>
                    {lista_aumentadas}
                    </td>
                </tr>
                </table>

                <br>

                <p>{mensagem_status}</p>

                <br>
                <p>Sistema Veredas</p>

            </div>

            </body>
            </html>
            """


            enviar_email(corpo_html, [pessoa["e_mail"]], assunto)










# ==================================================
# Função utilitária
# Efetiva o remanejamento no orçamento
# (redistribui valores entre as despesas)
# ==================================================
def efetivar_remanejamento(
    col_projetos,
    codigo_projeto_atual,
    financeiro,
    reduzidas,
    aumentadas
):
    """
    Aplica o remanejamento diretamente nas linhas do orçamento.

    Regras:
    • diminui valor_total das despesas reduzidas
    • aumenta valor_total das despesas aumentadas
    • NÃO altera financeiro.valor_total (total do projeto)
    """

    orcamento_atual = financeiro.get("orcamento", []) or []

    # percorre cada linha do orçamento
    for item in orcamento_atual:

        nome = item.get("nome_despesa")
        valor_atual = item.get("valor_total", 0) or 0

        # aplicar reduções
        for r in reduzidas:
            if r["nome_despesa"] == nome:
                valor_atual -= float(r["valor_reduzido"])

        # aplicar aumentos
        for a in aumentadas:
            if a["nome_despesa"] == nome:
                valor_atual += float(a["valor_aumentado"])

        item["valor_total"] = valor_atual

    # salva apenas o orçamento atualizado
    col_projetos.update_one(
        {"codigo": codigo_projeto_atual},
        {
            "$set": {
                "financeiro.orcamento": orcamento_atual
            }
        }
    )







def obter_nome_investidor(db, projeto):
    """
    Resolve o texto do investidor para uso no recibo, no formato:
    Nome do Investidor - SIGLA / Nome do Investidor - SIGLA

    Não usa st.stop().
    Retorna:
        - nome_investidor (str ou None)
        - erros (list[str])
    """

    # Coleções
    col_editais = db["editais"]
    col_ciclos = db["ciclos_investimento"]
    col_investidores = db["investidores"]

    nome_investidor = None
    erros = []

    # --------------------------------------------------
    # 1. Buscar edital do projeto
    # --------------------------------------------------
    codigo_edital = projeto.get("edital")

    if not codigo_edital:
        erros.append("Projeto sem código de edital.")
        return nome_investidor, erros

    edital = col_editais.find_one(
        {"codigo_edital": codigo_edital},
        {"ciclo_investimento": 1}
    )

    if not edital:
        erros.append(f"Edital '{codigo_edital}' não encontrado.")
        return nome_investidor, erros

    codigo_ciclo = edital.get("ciclo_investimento")

    if not codigo_ciclo:
        erros.append("Edital sem ciclo de investimento definido.")
        return nome_investidor, erros

    # --------------------------------------------------
    # 2. Buscar ciclo de investimento
    # --------------------------------------------------
    ciclo = col_ciclos.find_one(
        {"codigo_ciclo": codigo_ciclo},
        {"investidores": 1}
    )

    if not ciclo:
        erros.append(f"Ciclo de investimento '{codigo_ciclo}' não encontrado.")
        return nome_investidor, erros

    siglas = ciclo.get("investidores", [])

    if not siglas:
        erros.append("Nenhum investidor definido no ciclo de investimento.")
        return nome_investidor, erros

    # --------------------------------------------------
    # 3. Resolver investidores (nome + sigla)
    # --------------------------------------------------
    nomes_formatados = []

    for sigla in siglas:
        investidor = col_investidores.find_one(
            {"sigla_investidor": sigla},
            {"nome_investidor": 1, "sigla_investidor": 1}
        )

        if not investidor:
            erros.append(f"Investidor com sigla '{sigla}' não encontrado.")
            continue

        nome = investidor.get("nome_investidor")
        sigla_inv = investidor.get("sigla_investidor")

        if not nome or not sigla_inv:
            erros.append(f"Investidor '{sigla}' com dados incompletos.")
            continue

        # Formato final exigido no recibo
        nomes_formatados.append(f"{nome} - {sigla_inv}")

    if not nomes_formatados:
        erros.append("Não foi possível resolver o nome de nenhum investidor.")
        return nome_investidor, erros

    # --------------------------------------------------
    # 4. Texto final
    # --------------------------------------------------
    nome_investidor = " / ".join(nomes_formatados)

    return nome_investidor, erros






def calcular_gasto(item):
    lancamentos = item.get("lancamentos", [])
    return sum(
        l.get("valor_despesa", 0)
        for l in lancamentos
        if l.get("valor_despesa") is not None
    )



def gerar_recibo_docx(
    caminho_arquivo,
    valor_parcela,
    numero_parcela,
    nome_projeto,
    data_assinatura_contrato,
    contatos,
    nome_organizacao,
    cnpj_organizacao,
    contrato_nome,
    nome_investidor  # NOVO PARÂMETRO
):

    """
    Gera um arquivo DOCX de recibo com texto padrão do projeto.
    """


    # ============================
    # VALIDAÇÃO DE CAMPOS OBRIGATÓRIOS
    # ============================

    campos_obrigatorios = {
        "Nome/número do contrato": contrato_nome,
        "Data de assinatura do contrato": data_assinatura_contrato,
        "Nome do projeto": nome_projeto,
        "Valor da parcela": valor_parcela,
        "Número da parcela": numero_parcela,
        "Organização": nome_organizacao,
        "CNPJ da organização": cnpj_organizacao,
        "Contatos para assinatura": contatos,
        "Investidor": nome_investidor,

    }

    campos_faltando = []

    for nome, valor in campos_obrigatorios.items():
        if valor is None:
            campos_faltando.append(nome)
        elif isinstance(valor, str) and not valor.strip():
            campos_faltando.append(nome)
        elif isinstance(valor, list) and len(valor) == 0:
            campos_faltando.append(nome)

    if campos_faltando:
        st.error(
            "Não foi possível gerar o recibo. "
            "Os seguintes campos estão faltando:\n\n"
            + "\n".join([f"- {c}" for c in campos_faltando]),
            icon=":material/error:"
        )
        return False


    doc = Document()

    # ============================
    # TÍTULO
    # ============================
    titulo = doc.add_paragraph("Recibo")
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titulo.runs[0].bold = True
    titulo.runs[0].font.size = Pt(14)

    doc.add_paragraph("")
    doc.add_paragraph("")
    doc.add_paragraph("")

    # ============================
    # TEXTO PRINCIPAL
    # ============================
    valor_fmt = f"R$ {valor_parcela:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
    valor_extenso = valor_por_extenso(valor_parcela)
    ordinal = numero_ordinal_pt(numero_parcela)

    data_assinatura = data_extenso_pt(data_assinatura_contrato)

    data_hoje = data_extenso_pt(datetime.datetime.today())

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 2.3
    p.paragraph_format.space_after = Pt(12)

    p.add_run(
        "Recebi do Instituto Internacional de Educação do Brasil (IEB), "
        "a quantia de "
    )

    r = p.add_run(f"{valor_fmt} ({valor_extenso})")
    r.bold = True

    p.add_run(
        ", referente à "
        f"{ordinal} parcela de Recursos destinados a apoiar o projeto titulado "
    )

    r = p.add_run(nome_projeto)
    r.bold = True
    r.italic = True

    p.add_run(
        ", sob o Mecanismo de Pequenos Apoios, conforme o contrato de subvenção nº "
    )

    r = p.add_run(contrato_nome)
    r.bold = True
    r.italic = True


    p.add_run(
        f", assinado em {data_assinatura}, no âmbito do "
    )

    r = p.add_run(nome_investidor)
    r.bold = True

    p.add_run(
        ".\n\n"
        f"Brasília-DF, {data_hoje}"
    )

    # Ajusta o tamanho da fonte do parágrafo
    for run in p.runs:
        run.font.size = Pt(12.5)



    # ============================
    # ASSINATURAS
    # ============================

    def add_assinatura_centralizada(doc, texto, bold=False):
        p = doc.add_paragraph(texto)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.3
        if bold:
            p.runs[0].bold = True
        return p  # retorna para ajuste de fonte depois


    for contato in contatos:

        # Espaço ANTES do bloco
        p = doc.add_paragraph("")
        p.paragraph_format.space_after = Pt(18)

        # Linha de assinatura
        p = add_assinatura_centralizada(doc, "_" * 65)

        # Organização
        p = add_assinatura_centralizada(doc, nome_organizacao, bold=True)

        # CNPJ
        p = add_assinatura_centralizada(doc, f"CNPJ {cnpj_organizacao}", bold=True)

        # Nome
        p = add_assinatura_centralizada(doc, contato.get("nome", ""))

        # Função
        p = add_assinatura_centralizada(doc, contato.get("funcao", ""))

        # Parágrafo vazio entre assinaturas
        p = doc.add_paragraph("")

        # Ajuste de fonte do último parágrafo criado
        for run in p.runs:
            run.font.size = Pt(12.5)



    # ============================
    # SALVAR
    # ============================

    os.makedirs(os.path.dirname(caminho_arquivo), exist_ok=True)

    doc.save(caminho_arquivo)
    return True


    # # ============================
    # # SALVAR
    # # ============================
    # doc.save(caminho_arquivo)
    # return True





# Atualiza datas das parcelas, que devem ser 15 dias após o relatório anterior
def atualizar_datas_parcelas(col_projetos, codigo_projeto):
    # Busca o projeto no MongoDB
    projeto = col_projetos.find_one({"codigo": codigo_projeto})

    parcelas = projeto.get("financeiro", {}).get("parcelas", [])
    relatorios = projeto.get("relatorios", [])

    # Se não houver dados suficientes, sai
    if not parcelas or not relatorios:
        return

    # Mapeia relatórios por número (1, 2, 3...)
    mapa_relatorios = {
        r["numero"]: r
        for r in relatorios
        if r.get("numero") is not None
    }

    novas_parcelas = []

    for p in parcelas:
        numero = p.get("numero")

        # Parcela 1 não muda
        if numero == 1:
            novas_parcelas.append(p)
            continue

        # Relatório anterior (parcela 2 → relatório 1, etc)
        relatorio_ref = mapa_relatorios.get(numero - 1)

        if relatorio_ref and relatorio_ref.get("data_prevista"):
            data_relatorio = pd.to_datetime(
                relatorio_ref["data_prevista"],
                format="%d/%m/%Y",
                errors="coerce"
            )

            if pd.notnull(data_relatorio):
                nova_data = (
                    data_relatorio + datetime.timedelta(days=15)
                ).strftime("%d/%m/%Y")
            else:
                nova_data = None
        else:
            nova_data = None

        # Copia a parcela original
        parcela_atualizada = p.copy()

        # Atualiza apenas a data_prevista
        parcela_atualizada["data_prevista"] = nova_data

        novas_parcelas.append(parcela_atualizada)

    # Atualiza no MongoDB
    col_projetos.update_one(
        {"codigo": codigo_projeto},
        {"$set": {"financeiro.parcelas": novas_parcelas}}
    )






def criar_parcelas_a_partir_relatorios(col_projetos, codigo_projeto):
    """
    Cria parcelas automaticamente com base nos relatórios.

    Regra:
    - Nº parcelas = nº relatórios + 1
    - Parcela 1: permanece sem data automática
    - Parcela N (N >= 2): relatório N-1 + 15 dias
    """

    projeto = col_projetos.find_one({"codigo": codigo_projeto})

    if not projeto:
        return

    relatorios = projeto.get("relatorios", [])
    parcelas_existentes = projeto.get("financeiro", {}).get("parcelas", [])

    if not relatorios:
        return

    # -----------------------------------
    # Ordena relatórios
    # -----------------------------------
    relatorios = sorted(
        [r for r in relatorios if r.get("numero") is not None],
        key=lambda x: x["numero"]
    )

    # -----------------------------------
    # Mapa de parcelas existentes
    # -----------------------------------
    mapa_existente = {
        p.get("numero"): p
        for p in parcelas_existentes
        if p.get("numero") is not None
    }

    total_parcelas = len(relatorios) + 1
    novas_parcelas = []

    for i in range(1, total_parcelas + 1):

        parcela_antiga = mapa_existente.get(i, {})
        parcela_atualizada = parcela_antiga.copy()

        parcela_atualizada["numero"] = i




        # -----------------------------------
        # Parcela 1 mantém data existente (se houver)
        # -----------------------------------
        if i == 1:

            data_existente = parcela_antiga.get("data_prevista")

            if data_existente:
                data_dt = pd.to_datetime(
                    data_existente,
                    format="%d/%m/%Y",
                    errors="coerce"
                )
            else:
                data_dt = None




        else:
            relatorio_ref = relatorios[i - 2]  # índice ajustado

            data_relatorio = relatorio_ref.get("data_prevista")

            if data_relatorio:
                data_dt = pd.to_datetime(
                    data_relatorio,
                    format="%d/%m/%Y",
                    errors="coerce"
                )

                if pd.notnull(data_dt):
                    data_dt = data_dt + datetime.timedelta(days=15)
                else:
                    data_dt = None
            else:
                data_dt = None

        # -----------------------------------
        # Formatar data
        # -----------------------------------
        parcela_atualizada["data_prevista"] = (
            data_dt.strftime("%d/%m/%Y")
            if data_dt is not None and pd.notnull(data_dt)
            else None
        )

        novas_parcelas.append(parcela_atualizada)

    # -----------------------------------
    # Salvar no MongoDB
    # -----------------------------------
    col_projetos.update_one(
        {"codigo": codigo_projeto},
        {
            "$set": {
                "financeiro.parcelas": novas_parcelas
            }
        }
    )








# ==========================================================================================
# DIÁLOGO: VER RELATOS FINANCEIROS
# ==========================================================================================
@st.dialog("Lançamentos de despesa", width="large")
def dialog_relatos_fin():

    despesa = st.session_state.get("despesa_selecionada")

    if not isinstance(despesa, dict):
        st.warning("Nenhuma despesa selecionada.")
        return

    nome_despesa = (
        despesa.get("despesa")
        or despesa.get("Despesa")
        or "Despesa sem nome"
    )

    st.markdown(f"### {nome_despesa}")
    st.write("")

    # ==========================================================
    # BUSCA DOS LANÇAMENTOS DA DESPESA
    # ==========================================================
    lancamentos = []

    for d in projeto.get("financeiro", {}).get("orcamento", []):
        if d.get("nome_despesa") == nome_despesa:
            lancamentos = d.get("lancamentos", [])
            break

    if not lancamentos:
        st.caption("Esta despesa ainda não possui lançamentos.")
        return

    # ==========================================================
    # RENDERIZAÇÃO DOS LANÇAMENTOS
    # ==========================================================
    for lanc in lancamentos:

        with st.container(border=True):

            id_despesa = lanc.get("id_lanc_despesa", "").upper()
            num_relatorio = lanc.get("relatorio_numero")



            # ==================================================
            # BADGE DE STATUS
            # ==================================================

            status_despesa_db = lanc.get("status_despesa", "em_analise")
            tem_devolutiva = bool(lanc.get("devolutiva"))

            if status_despesa_db == "aberto" and tem_devolutiva:
                badge = {
                    "label": "Pendente",
                    "bg": "#F8D7DA",
                    "color": "#721C24"
                }
            elif status_despesa_db == "aberto":
                badge = {
                    "label": "Aberto",
                    "bg": "#FFF3CD",
                    "color": "#856404"
                }
            elif status_despesa_db == "aceito":
                badge = {
                    "label": "Aceito",
                    "bg": "#D4EDDA",
                    "color": "#155724"
                }
            else:
                badge = {
                    "label": "Em análise",
                    "bg": "#D1ECF1",
                    "color": "#0C5460"
                }


            # ==================================================
            # CABEÇALHO COM BADGE
            # ==================================================

            col_header1, col_header2 = st.columns([9, 1])

            with col_header1:
                st.markdown(f"**{id_despesa}** (R{num_relatorio})")

            with col_header2:
                st.markdown(
                    f"""
                    <div style="margin-top:6px;">
                        <span style="
                            background:{badge['bg']};
                            color:{badge['color']};
                            padding:4px 10px;
                            border-radius:20px;
                            font-size:12px;
                            font-weight:600;
                        ">
                            {badge['label']}
                        </span>
                    </div>
                    """,
                    unsafe_allow_html=True
                )


            # ==================================================
            # DESCRIÇÃO
            # ==================================================
            st.write(lanc.get("descricao_despesa", ""))




            col1, col2 = st.columns([1, 2])

            c1, c2 = col1.columns([1, 3])

            c1.write("**Data:**")
            c2.write(lanc.get("data_despesa", "-"))

            c1.write("**Fornecedor:**")
            c2.write(lanc.get("fornecedor", "-"))

            c1.write("**CPF/CNPJ:**")
            c2.write(lanc.get("cpf_cnpj", "-"))

            valor = lanc.get("valor_despesa", 0)
            valor_br = f"{valor:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")

            c1.write("**Valor (R$):**")
            c2.write(valor_br)

            anexos = lanc.get("anexos", [])
            if anexos:
                col2.markdown("**Anexos:**")
                for a in anexos:
                    link = gerar_link_drive(a["id_arquivo"])
                    col2.markdown(f"[{a['nome_arquivo']}]({link})")






###########################################################################################################
# INTERFACE PRINCIPAL DA PÁGINA
###########################################################################################################



# Logo do sidebar
st.logo("images/ieb_logo.svg", size='large')

# Título da página e identificação
col_titulo, col_identificacao = st.columns([3, 2])

with col_titulo:
    st.header("Financeiro")

with col_identificacao:
    st.markdown(
        f"<div style='text-align: right; margin-top: 30px;'>{df_projeto['codigo'].values[0]} - {df_projeto['sigla'].values[0]}</div>",
        unsafe_allow_html=True
    )




# Filtragem de usuários internos
usuario_interno = st.session_state.tipo_usuario in ["admin", "equipe"]



if usuario_interno:
    cron_desemb, orcamento, remanejamentos, recibos = st.tabs(
        ["Cronograma", "Orçamento", "Remanejamentos", "Recibos"]
    )
else:
    cron_desemb, orcamento, remanejamentos = st.tabs(
        ["Cronograma", "Orçamento", "Remanejamentos"]
    )




# ##############################################################
# ABA CRONOGRAMA
# ##############################################################

with cron_desemb:

    # Capturando os dados no bd
    valor_atual = financeiro.get("valor_total")




    # --------------------------------------------------
    # PERMISSÃO E MODO DE EDIÇÃO
    # --------------------------------------------------
    usuario_interno = st.session_state.tipo_usuario in ["admin", "equipe"]

    with st.container(horizontal=True, horizontal_alignment="right"):
        if usuario_interno:
            modo_edicao = st.toggle("Modo de edição", key="editar_cronograma")
        else:
            modo_edicao = False





    

    # --------------------------------------------------
    # MODO VISUALIZAÇÃO CRONOGRAMA
    # --------------------------------------------------
    if not modo_edicao:

        st.markdown('### Cronograma de Parcelas e Relatórios')

        st.write("")


        # -----------------------------------
        # Validação de dados mínimos para montagem do cronograma
        # -----------------------------------

        parcelas = financeiro.get("parcelas", [])
        relatorios = projeto.get("relatorios", [])

        if not parcelas or not relatorios:
            st.caption(
                "Cadastre o **Valor do projeto**, os **Relatórios** e as **Parcelas** para montar o cronograma."
            )
            st.stop()


        # Notifica se as parcelas estão com valores errados, após cadastro de Aditivo / Devolução de valor
        notifica_parcelas_desencontradas()





        # OPÇÃO 1 DE LAYOUT DA TABELA


        st.write("")

        # -----------------------------
        # Construir cronograma
        # -----------------------------
        linhas_cronograma = []

        # ===== Parcelas =====
        parcelas = financeiro.get("parcelas", [])

        for p in parcelas:
            numero = p.get("numero")
            percentual = p.get("percentual")
            valor = p.get("valor")
            data_prevista = p.get("data_prevista")
            data_realizada = p.get("data_realizada")

            linhas_cronograma.append(
                {
                    "evento": f"Parcela {numero}",
                    "Entregas": "",
                    "Valor R$": (
                        f"{valor:,.2f}".replace(",", "X")
                        .replace(".", ",")
                        .replace("X", ".")
                        if valor is not None else ""
                    ),
                    "Data prevista": pd.to_datetime(
                        data_prevista,
                        format="%d/%m/%Y",  # formato brasileiro
                        errors="coerce"
                    ),
                    "Data realizada": (
                        pd.to_datetime(data_realizada).strftime("%d/%m/%Y")
                        if data_realizada else ""
                    ),
                }
            )

        # ===== Relatórios =====
        relatorios = projeto.get("relatorios", [])

        for r in relatorios:
            numero = r.get("numero")
            entregas = r.get("entregas", [])
            data_prevista = r.get("data_prevista")
            data_realizada = r.get("data_envio")

            linhas_cronograma.append(
                {
                    "evento": f"Relatório {numero}",
                    "Entregas": " / ".join(entregas) if entregas else "",

                    # "Entregas": "\n".join(entregas) if entregas else "",
                    # "Entregas": "<br>".join(entregas) if entregas else "",
                    "Valor R$": "",
                    # "Percentual": "",
                    "Data prevista": pd.to_datetime(
                        data_prevista,
                        format="%d/%m/%Y",  # garante interpretação no padrão brasileiro
                        errors="coerce"
                    ),
                    "Data realizada": (
                        pd.to_datetime(data_realizada).strftime("%d/%m/%Y")
                        if data_realizada else ""
                    ),
                }
            )

        # -----------------------------
        # DataFrame final
        # -----------------------------
        df_cronograma = pd.DataFrame(linhas_cronograma)


        if df_cronograma.empty:
            st.caption("Não há dados financeiros para exibição.")
        else:
            
            # Ordenar por data prevista
            df_cronograma = df_cronograma.sort_values(
                by="Data prevista",
                ascending=True
            )

            # -----------------------------------
            # Formatar data prevista para exibição
            # (tratando valores nulos com texto)
            # -----------------------------------
            df_cronograma["Data prevista"] = df_cronograma["Data prevista"].apply(
                lambda x: x.strftime("%d/%m/%Y") if pd.notnull(x) else "sem data definida"
            )

            # df_cronograma["Data prevista"] = df_cronograma["Data prevista"].dt.strftime(
            #     "%d/%m/%Y"
            # )

            # Renomear a coluna evento
            df_cronograma = df_cronograma.rename(
                columns={"evento": "Evento"}
            )


            # -----------------------------
            # Tabela
            # -----------------------------
            
            ui.table(df_cronograma)
        






    # --------------------------------------------------
    # MODO EDIÇÃO CRONOGRAMA
    # --------------------------------------------------
    else:



        notifica_parcelas_desencontradas()



        # Define a ordem de edição do cronograma
        opcao_editar_cron = st.radio(
            "O que deseja editar?",
            [
                "Valor do projeto",
                "Relatórios",
                "Parcelas"
            ],
            horizontal=True
        )

        # st.divider()
        st.write('')
        st.write('')








        # -------------------------------------------------------
        # Editar o valor total do projeto
        # -------------------------------------------------------
        if opcao_editar_cron == "Valor do projeto":

            col1, col2 = st.columns(2)

            # -----------------------------------
            # Recuperar valores atuais do banco
            # -----------------------------------
            valor_aditivo_atual = financeiro.get("valor_aditivo", 0.0)
            valor_devolucao_atual = financeiro.get("valor_devolucao", 0.0)

            # -----------------------------------
            # Coluna 1: Inputs
            # -----------------------------------
            with col1:

                with st.form("form_valor_total", border=False):

                    st.markdown("#### Valor inicial do projeto")

                    # -----------------------------------
                    # Input valor total
                    # -----------------------------------
                    valor_total = st.number_input(
                        "Valor total do projeto (R$)",
                        min_value=0.0,
                        step=1000.0,
                        format="%.2f",
                        value=float(valor_atual) if valor_atual is not None else 0.0,
                        width=300
                    )

                    st.write('')
                    st.write('')

                    st.markdown("#### Ajustes financeiros")

                    # -----------------------------------
                    # Layout em duas colunas para aditivo e devolução
                    # -----------------------------------
                    col_aditivo, col_devolucao = st.columns(2)

                    with col_aditivo:
                        valor_aditivo = st.number_input(
                            "Aditivo (R$)",
                            step=100.0,
                            format="%.2f",
                            value=float(valor_aditivo_atual) if valor_aditivo_atual is not None else 0.0,
                            width=300
                        )

                    with col_devolucao:
                        valor_devolucao = st.number_input(
                            "Devolução (R$)",
                            step=100.0,
                            format="%.2f",
                            value=float(valor_devolucao_atual) if valor_devolucao_atual is not None else 0.0,
                            width=300
                        )

                    st.write('')

                    salvar = st.form_submit_button(
                        "Salvar",
                        icon=":material/save:",
                        width=150,
                        type="primary"
                    )

                    # -----------------------------------
                    # Persistência no banco
                    # -----------------------------------
                    if salvar:
                        col_projetos.update_one(
                            {"codigo": codigo_projeto_atual},
                            {
                                "$set": {
                                    "financeiro.valor_total": float(valor_total),
                                    "financeiro.valor_aditivo": float(valor_aditivo),
                                    "financeiro.valor_devolucao": float(valor_devolucao)
                                }
                            }
                        )

                        st.success("Valores financeiros salvos com sucesso!", icon=":material/check:")
                        time.sleep(3)
                        st.rerun()

            # -----------------------------------
            # Coluna 2: Exibição do valor final
            # -----------------------------------
            with col2:

                # -----------------------------------
                # Garantir valores base
                # -----------------------------------
                valor_total_base = float(valor_atual) if valor_atual is not None else 0.0

                valor_aditivo = valor_aditivo_atual or 0.0
                valor_devolucao = valor_devolucao_atual or 0.0

                # -----------------------------------
                # Calcular valor final atualizado
                # -----------------------------------
                valor_atualizado = valor_total_base + valor_aditivo - valor_devolucao

                # -----------------------------------
                # Exibir apenas se houver ajuste
                # -----------------------------------
                if valor_aditivo != 0 or valor_devolucao != 0:

                    st.metric(
                        "Valor final com Aditivo / Devolução",
                        value=format_brl(valor_atualizado),
                    )





        # -------------------------------------------------------
        # Editar Relatórios
        # -------------------------------------------------------

        if opcao_editar_cron == "Relatórios":

            st.markdown("#### Relatórios")

            # --------------------------------------------------
            # Coletar entregas do projeto
            # --------------------------------------------------

            entregas_projeto = []

            plano = projeto.get("plano_trabalho", {})
            for componente in plano.get("componentes", []):
                for entrega in componente.get("entregas", []):
                    if entrega.get("entrega"):
                        entregas_projeto.append(entrega["entrega"])

            entregas_projeto = sorted(set(entregas_projeto))

            if not entregas_projeto:
                st.warning(
                    "Nenhuma entrega cadastrada para este projeto. "
                    "Cadastre as entregas no Plano de Trabalho antes de continuar.",
                    icon=":material/warning:"
                )

            # --------------------------------------------------
            # Dados existentes
            # --------------------------------------------------

            relatorios_existentes = projeto.get("relatorios", [])


            # -----------------------------------
            # Caso NÃO haja relatórios no banco
            # -----------------------------------
            if not relatorios_existentes:

                df_relatorios_base = pd.DataFrame(
                    columns=["numero", "entregas", "data_prevista"]
                )

            else:
                colunas = ["numero", "entregas", "data_prevista"]

                df_relatorios_base = pd.DataFrame(relatorios_existentes)
                df_relatorios_base = df_relatorios_base.reindex(columns=colunas)

                # Garantir colunas
                for col in ["numero", "entregas", "data_prevista"]:
                    if col not in df_relatorios_base.columns:
                        df_relatorios_base[col] = None

                # Converter data
                df_relatorios_base["data_prevista"] = pd.to_datetime(
                    df_relatorios_base["data_prevista"],
                    format="%d/%m/%Y",
                    errors="coerce"
                )

                # Garantir lista
                df_relatorios_base["entregas"] = df_relatorios_base["entregas"].apply(
                    lambda x: x if isinstance(x, list) else []
                )

                # Ordenar por número
                df_relatorios_base = df_relatorios_base.sort_values(
                    by="numero"
                ).reset_index(drop=True)

            # --------------------------------------------------
            # Editor
            # --------------------------------------------------

            df_editado = st.data_editor(
                df_relatorios_base,
                num_rows="dynamic",
                column_config={
                    "numero": st.column_config.NumberColumn(
                        "Número",
                        min_value=1,
                        step=1,
                        width=80
                    ),
                    "entregas": st.column_config.MultiselectColumn(
                        "Entregas",
                        options=[""] + entregas_projeto,
                        help="Selecione as entregas relacionadas a este relatório",
                        width=800
                    ),
                    "data_prevista": st.column_config.DateColumn(
                        "Data prevista",
                        format="DD/MM/YYYY",
                        width=150
                    ),
                },
                key="editor_relatorios",
                hide_index=True
            )

            st.write("")

            # --------------------------------------------------
            # Salvar
            # --------------------------------------------------

            if st.button("Salvar relatórios", icon=":material/save:"):

                df_salvar = df_editado.copy()

                # Remover linhas vazias
                df_salvar = df_salvar.dropna(
                    subset=["numero"],
                    how="any"
                )


                # -----------------------------------
                # Validação de datas previstas
                # -----------------------------------
                erros_datas = []

                for idx, row in df_salvar.iterrows():

                    data = row.get("data_prevista")

                    data_convertida = pd.to_datetime(
                        data,
                        format="%d/%m/%Y",
                        errors="coerce"
                    )

                    if pd.isna(data_convertida):
                        erros_datas.append(idx + 1)

                # -----------------------------------
                # Se houver erro, não salva
                # -----------------------------------
                if erros_datas:
                    linhas = ", ".join(map(str, erros_datas))

                    st.warning(
                        f"Data inválida na(s) linha(s): {linhas}. Utilize o formato DD/MM/AAAA.",
                        icon=":material/warning:"
                    )
                else:



                    # Ordenar por número
                    df_salvar = df_salvar.sort_values(
                        by="numero"
                    ).reset_index(drop=True)

                    relatorios_salvar = []

                    for _, row in df_salvar.iterrows():

                        entregas = [e for e in row["entregas"] if e]

                        relatorios_salvar.append(
                            {
                                "numero": int(row["numero"]),
                                "entregas": entregas,
                                "data_prevista": (
                                    None
                                    if pd.isna(row["data_prevista"])
                                    else pd.to_datetime(row["data_prevista"]).strftime("%d/%m/%Y")
                                ),
                            }
                        )

                    col_projetos.update_one(
                        {"codigo": codigo_projeto_atual},
                        {
                            "$set": {
                                "relatorios": relatorios_salvar
                            }
                        }
                    )

                    # Atualiza parcelas com base nos relatórios
                    criar_parcelas_a_partir_relatorios(col_projetos, codigo_projeto_atual)
                    

                    st.success("Relatórios salvos com sucesso!", icon=":material/check:")
                    time.sleep(3)
                    st.rerun()









        # -------------------------------------------------------
        # Editar Parcelas
        # -------------------------------------------------------

        if opcao_editar_cron == "Parcelas":

            st.markdown("#### Parcelas")

            st.caption(
                "A data da parcela 1 é definida **manualmente**. "
                "As demais são calculadas automaticamente a partir dos relatórios."
            )

            # -----------------------------------
            # Valor total ajustado
            # -----------------------------------
            valor_total_base = valor_atual if valor_atual is not None else 0.0

            valor_aditivo = financeiro.get("valor_aditivo") or 0.0
            valor_devolucao = financeiro.get("valor_devolucao") or 0.0

            valor_total = valor_total_base + valor_aditivo - valor_devolucao

            # -----------------------------------
            # Dados atuais
            # -----------------------------------
            parcelas = financeiro.get("parcelas", [])

            if parcelas:
                df_parcelas = pd.DataFrame(parcelas)

                # Conversão segura da data (sem recalcular nada)
                df_parcelas["data_prevista"] = pd.to_datetime(
                    df_parcelas["data_prevista"],
                    format="%d/%m/%Y",
                    errors="coerce"
                )

                # Garantir colunas essenciais
                for col in ["numero", "valor"]:
                    if col not in df_parcelas.columns:
                        df_parcelas[col] = None

            else:
                df_parcelas = pd.DataFrame(
                    columns=[
                        "numero",
                        "valor",
                        "data_prevista",
                    ]
                )

            # -----------------------------------
            # Ordenação lógica (por número)
            # -----------------------------------
            if not df_parcelas.empty:
                df_parcelas = df_parcelas.sort_values(
                    by="numero",
                    ascending=True
                ).reset_index(drop=True)

            # -----------------------------------
            # Garantir coluna valor
            # -----------------------------------
            df_parcelas["valor"] = df_parcelas["valor"].fillna(0.0)

            # -----------------------------------
            # Calcular percentual
            # -----------------------------------
            if valor_total > 0:
                df_parcelas["percentual"] = (
                    df_parcelas["valor"] / valor_total * 100
                )
            else:
                df_parcelas["percentual"] = 0.0

            # -----------------------------------
            # Formatar valor
            # -----------------------------------
            df_parcelas["valor_fmt"] = df_parcelas["valor"].apply(
                lambda x: f"R$ {x:,.2f}"
                .replace(",", "X")
                .replace(".", ",")
                .replace("X", ".")
                if pd.notna(x) else ""
            )

            # -----------------------------------
            # Editor
            # -----------------------------------
            df_editado_parcelas = st.data_editor(
                df_parcelas[
                    [
                        "numero",
                        "valor_fmt",
                        "percentual",
                        "data_prevista",
                    ]
                ],
                num_rows="dynamic",
                width=800,
                column_config={
                    "numero": st.column_config.NumberColumn(
                        "Número",
                        min_value=1,
                        step=1,
                        width=60
                    ),
                    "valor_fmt": st.column_config.TextColumn(
                        "Valor (R$)",
                        width=150
                    ),
                    "percentual": st.column_config.NumberColumn(
                        "Percentual (auto)",
                        format="%.2f%%",
                        disabled=True,
                        width=100
                    ),
                    "data_prevista": st.column_config.DateColumn(
                        "Data prevista",
                        format="DD/MM/YYYY",
                        width=150
                    ),
                },
                key="editor_parcelas",
            )

            st.write("")

            # -----------------------------------
            # Converter valor BRL para float
            # -----------------------------------
            def parse_brl(valor):
                """
                Converte string monetária brasileira para float.
                """
                if not valor:
                    return 0.0

                return float(
                    str(valor)
                    .replace("R$", "")
                    .replace(".", "")
                    .replace(",", ".")
                    .strip()
                )

            # -----------------------------------
            # Reconstruir DataFrame após edição
            # -----------------------------------
            df_parcelas = df_editado_parcelas.copy()

            # Conversão segura da data
            df_parcelas["data_prevista"] = pd.to_datetime(
                df_parcelas["data_prevista"],
                errors="coerce"
            )

            # Converter valores
            df_parcelas["valor"] = df_parcelas["valor_fmt"].apply(parse_brl)

            # -----------------------------------
            # Recalcular percentual
            # -----------------------------------
            if valor_total > 0:
                df_parcelas["percentual"] = (
                    df_parcelas["valor"] / valor_total * 100
                )
            else:
                df_parcelas["percentual"] = 0.0

            # -----------------------------------
            # Totais
            # -----------------------------------
            soma_valores = df_parcelas["valor"].sum()

            soma_fmt = (
                f"R$ {soma_valores:,.2f}"
                .replace(",", "X")
                .replace(".", ",")
                .replace("X", ".")
            )

            total_fmt = (
                f"R$ {valor_total:,.2f}"
                .replace(",", "X")
                .replace(".", ",")
                .replace("X", ".")
            )

            st.write(f"**Total das parcelas:** {soma_fmt.replace('$', '\\$')}")
            st.write(f"**Valor total ajustado do projeto:** {total_fmt.replace('$', '\\$')}")

            # -----------------------------------
            # Salvar
            # -----------------------------------
            if st.button("Salvar parcelas", icon=":material/save:"):

                df_salvar = df_parcelas.dropna(
                    subset=["valor", "data_prevista"],
                    how="any"
                ).copy()

                # -----------------------------------
                # Validar soma
                # -----------------------------------
                soma_valores = df_salvar["valor"].sum()

                if round(soma_valores, 2) != round(valor_total, 2):

                    soma_fmt = (
                        f"R$ {soma_valores:,.2f}"
                        .replace(",", "X")
                        .replace(".", ",")
                        .replace("X", ".")
                    )

                    total_fmt = (
                        f"R$ {valor_total:,.2f}"
                        .replace(",", "X")
                        .replace(".", ",")
                        .replace("X", ".")
                    )

                    st.error(
                        f"Erro: A soma das parcelas ({soma_fmt.replace('$', '\\$')}) deve ser igual ao valor total do projeto ({total_fmt.replace('$', '\\$')}).",
                        icon=":material/error:"
                    )
                    st.stop()

                # Ordenar por número (não por data)
                df_salvar = df_salvar.sort_values(
                    by="numero",
                    ascending=True
                ).reset_index(drop=True)

                # -----------------------------------
                # Preservar dados antigos
                # -----------------------------------
                mapa_existente = {
                    p.get("numero"): p
                    for p in parcelas
                    if p.get("numero") is not None
                }

                parcelas_final = []

                for _, row in df_salvar.iterrows():

                    numero = int(row["numero"]) if not pd.isna(row["numero"]) else None

                    parcela_antiga = mapa_existente.get(numero, {})

                    parcela_atualizada = parcela_antiga.copy()

                    parcela_atualizada.update({
                        "numero": numero,
                        "percentual": float(row["percentual"]),
                        "valor": float(row["valor"]),
                        "data_prevista": (
                            pd.to_datetime(row["data_prevista"])
                            .strftime("%d/%m/%Y")
                        ),
                    })

                    parcelas_final.append(parcela_atualizada)

                col_projetos.update_one(
                    {"codigo": codigo_projeto_atual},
                    {
                        "$set": {
                            "financeiro.parcelas": parcelas_final
                        }
                    }
                )

                st.success("Parcelas salvas com sucesso!", icon=":material/check:")
                time.sleep(3)
                st.rerun()















# --------------------------------------------------
# ABA ORÇAMENTO
# --------------------------------------------------


with orcamento:


    # ==================================================
    # PERMISSÃO E MODO DE EDIÇÃO
    # ==================================================

    with st.container(horizontal=True, horizontal_alignment="right"):
        if usuario_interno:
            modo_edicao = st.toggle("Modo de edição", key="editar_orcamento")
        else:
            modo_edicao = False



    st.markdown("### Orçamento")
    st.write("")








    # ==================================================
    # MODO VISUALIZAÇÃO — ORÇAMENTO AGRUPADO POR CATEGORIA
    # ==================================================
    if not modo_edicao:



        # ==================================================
        # Notificação de total do orçamento diferente do valor total do projeto, somente para usuário interno
        # ==================================================

        if usuario_interno:


            # -----------------------------------
            # Valor total ajustado do projeto
            # -----------------------------------
            valor_total_base = financeiro.get("valor_total") or 0.0
            valor_aditivo = financeiro.get("valor_aditivo") or 0.0
            valor_devolucao = financeiro.get("valor_devolucao") or 0.0

            valor_total_projeto = valor_total_base + valor_aditivo - valor_devolucao



            
            orcamento_salvo = financeiro.get("orcamento", [])

            if valor_total_projeto and orcamento_salvo:

                soma_orcamento = sum(
                    item.get("valor_total", 0)
                    for item in orcamento_salvo
                    if item.get("valor_total") is not None
                )

                if round(soma_orcamento, 2) != round(valor_total_projeto, 2):

                    soma_fmt = (
                        f"R\\$ {soma_orcamento:,.2f}"
                        .replace(",", "X")
                        .replace(".", ",")
                        .replace("X", ".")
                    )

                    total_fmt = (
                        f"R\\$ {valor_total_projeto:,.2f}"
                        .replace(",", "X")
                        .replace(".", ",")
                        .replace("X", ".")
                    )


                    st.warning(
                        f"O total do orçamento ({soma_fmt}) é diferente do valor total do projeto ({total_fmt}).",
                        icon=":material/warning:"
                    )

                    st.write('')





        # -----------------------------------
        # Valor total ajustado (com aditivo e devolução)
        # -----------------------------------
        valor_total_base = financeiro.get("valor_total") or 0.0
        valor_aditivo = financeiro.get("valor_aditivo") or 0.0
        valor_devolucao = financeiro.get("valor_devolucao") or 0.0

        valor_total = valor_total_base + valor_aditivo - valor_devolucao



        orcamento = financeiro.get("orcamento", [])
        parcelas = financeiro.get("parcelas", [])

        gasto_total = 0
        valor_recebido = 0

        # -----------------------------
        # Calcular gasto total (se houver orçamento)
        # -----------------------------
        if orcamento:
            for item in orcamento:
                gasto_total += sum(
                    l.get("valor_despesa", 0)
                    for l in item.get("lancamentos", [])
                    if l.get("valor_despesa") is not None
                )

        # -----------------------------
        # Calcular valor recebido (se houver parcelas)
        # -----------------------------
        if parcelas:
            valor_recebido += sum(
                p.get("valor", 0)
                for p in parcelas
                if p.get("data_realizada") not in [None, ""]
            )

        saldo_total = valor_total - gasto_total




        # --------------------------------------------------
        # Exibição das métricas em 3 colunas
        # --------------------------------------------------

        col1, col2, col3 = st.columns(3)

        col1.metric(
            label="Valor total do projeto",
            value=f"R$ {valor_total:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
        )

        col2.metric(
            label="Gasto",
            value=f"R$ {gasto_total:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
        )

        col3.metric(
            label="Saldo",
            value=f"R$ {saldo_total:,.2f}".replace(",", "X").replace(".", ",").replace("X", "."),
            delta=None if saldo_total >= 0 else "Negativo",
            delta_color="inverse" if saldo_total < 0 else "normal"
        )





        # --------------------------------------------------
        # BARRAS DE PROGRESSO: recebido x gasto
        # --------------------------------------------------

        # -----------------------------
        # Cálculo do valor recebido
        # -----------------------------
        parcelas = financeiro.get("parcelas", [])

        valor_recebido = sum(
            p.get("valor", 0)
            for p in parcelas
            if p.get("data_realizada") not in [None, ""]
        )

        # -----------------------------
        # Percentuais (0 a 1)
        # -----------------------------

        if valor_total > 0:
            pct_recebido = min(valor_recebido / valor_total, 1)
            pct_gasto = min(gasto_total / valor_total, 1)
        else:
            pct_recebido = 0
            pct_gasto = 0





        # -----------------------------
        # Barra de valor recebido
        # -----------------------------
        st.progress(
            pct_recebido,
            text=(
                f"Valor recebido: R$ {valor_recebido:,.2f}"
                .replace(",", "X").replace(".", ",").replace("X", ".")
                # + f"  ({pct_recebido*100:.1f}%)"
            )
        )

        # -----------------------------
        # Barra de valor gasto
        # -----------------------------
        st.progress(
            pct_gasto,
            text=(
                f"Valor gasto: R$ {gasto_total:,.2f}"
                .replace(",", "X").replace(".", ",").replace("X", ".")
                # + f"  ({pct_gasto*100:.1f}%)"
            )
        )



        st.write("")

        # --------------------------------------------------
        # ESTADOS DO DIÁLOGO (mantidos como no seu código)
        # --------------------------------------------------
        if "despesa_selecionada" not in st.session_state:
            st.session_state["despesa_selecionada"] = None

        if "despesa_selecionada_tabela_key" not in st.session_state:
            st.session_state["despesa_selecionada_tabela_key"] = None

        if "abrir_dialogo_despesa" not in st.session_state:
            st.session_state["abrir_dialogo_despesa"] = False

        # --------------------------------------------------
        # Dados do orçamento
        # --------------------------------------------------
        orcamento = financeiro.get("orcamento", [])

        if not orcamento:
            st.caption("Nenhuma despesa cadastrada no orçamento.")
            st.stop()

        # --------------------------------------------------
        # Cálculo de gasto e saldo (inalterado)
        # --------------------------------------------------
        for item in orcamento:
            gasto = calcular_gasto(item)
            valor_item = item.get("valor_total", 0) or 0

            item["gasto"] = gasto
            item["saldo"] = valor_item - gasto

        # --------------------------------------------------
        # Criação do DataFrame
        # --------------------------------------------------
        df_orcamento = pd.DataFrame(orcamento)

        # coluna auxiliar numérica (NÃO exibida)
        df_orcamento["saldo_num"] = df_orcamento["saldo"]

        # --------------------------------------------------
        # Garantir colunas mínimas
        # --------------------------------------------------
        for col in [
            "categoria",
            "nome_despesa",
            "descricao_despesa",
            "unidade",
            "quantidade",
            "valor_unitario",
            "valor_total",
        ]:
            if col not in df_orcamento.columns:
                df_orcamento[col] = None

        # --------------------------------------------------
        # Formatação monetária (apenas exibição)
        # --------------------------------------------------
        def fmt_moeda(x):
            if x in [None, ""]:
                return ""
            return f"R$ {x:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")

        df_orcamento["Valor unitário"] = df_orcamento["valor_unitario"].apply(fmt_moeda)
        df_orcamento["Valor total"] = df_orcamento["valor_total"].apply(fmt_moeda)
        df_orcamento["Gasto"] = df_orcamento["gasto"].apply(fmt_moeda)
        df_orcamento["Saldo"] = df_orcamento["saldo"].apply(fmt_moeda)


        # --------------------------------------------------
        # Agrupamento por categoria (ordem alfabética)
        # --------------------------------------------------
        categorias = sorted(
            df_orcamento["categoria"]
            .dropna()
            .unique()
            .tolist(),
            key=str.lower  # ordenação case-insensitive
        )



        # --------------------------------------------------
        # CALLBACK - Seleção de despesa

        # --------------------------------------------------
        def criar_callback_selecao_orcamento(dataframe_orc, chave_tabela):

            def handle_selecao():

                estado_tabela = st.session_state.get(chave_tabela, {})
                selecao = estado_tabela.get("selection", {})
                linhas = selecao.get("rows", [])

                if not linhas:
                    return

                idx = linhas[0]
                linha = dataframe_orc.iloc[idx]

                despesa_escolhida = {
                    "categoria": linha.get("categoria", ""),
                    "nome_despesa": linha.get("nome_despesa", ""),
                    "descricao_despesa": linha.get("descricao_despesa", ""),
                    "unidade": linha.get("unidade", ""),
                    "quantidade": linha.get("quantidade", 0),
                    "valor_unitario": linha.get("valor_unitario", 0),
                    "valor_total": linha.get("valor_total", 0),
                    "indice": idx,
                }

                # Compatibilidade com o diálogo
                despesa_escolhida["despesa"] = despesa_escolhida["nome_despesa"]
                despesa_escolhida["Despesa"] = despesa_escolhida["nome_despesa"]

                st.session_state["despesa_selecionada"] = despesa_escolhida
                st.session_state["despesa_selecionada_tabela_key"] = chave_tabela
                st.session_state["abrir_dialogo_despesa"] = True

            return handle_selecao

        # --------------------------------------------------
        # RENDERIZAÇÃO POR CATEGORIA
        # --------------------------------------------------
        for categoria in categorias:

            st.write("")
            st.markdown(f"##### {categoria}")

            # Filtra categoria
            df_cat = df_orcamento[df_orcamento["categoria"] == categoria].copy()

            # Renomeia colunas para exibição
            df_vis = df_cat.rename(columns={
                "nome_despesa": "Despesa",
                "descricao_despesa": "Descrição",
                "unidade": "Unidade",
                "quantidade": "Quantidade",
            })

            colunas_vis = [
                "Despesa",
                "Descrição",
                # "Unidade",
                # "Quantidade",
                # "Valor unitário",
                "Valor total",
                "Gasto",
                "Saldo",
            ]

            key_df = f"df_vis_orcamento_{categoria}"

            callback_selecao = criar_callback_selecao_orcamento(
                df_cat,
                key_df
            )

            # --------------------------------------------------
            # Estilo: saldo negativo em vermelho
            # --------------------------------------------------
            def estilo_saldo(col):
                estilos = []
                for idx in col.index:
                    if df_cat.loc[idx, "saldo_num"] < 0:
                        estilos.append("color: red;")
                    else:
                        estilos.append("")
                return estilos

            df_estilizado = (
                df_vis[colunas_vis]
                .style
                .apply(estilo_saldo, subset=["Saldo"])
            )

            # --------------------------------------------------
            # DataFrame final
            # --------------------------------------------------
            st.dataframe(
                df_estilizado,
                hide_index=True,
                selection_mode="single-row",
                key=key_df,
                on_select=callback_selecao,
                column_config={
                    "Despesa": st.column_config.TextColumn(width=220),
                    "Descrição": st.column_config.TextColumn(width=420),
                    "Unidade": st.column_config.TextColumn(width=120),
                    "Quantidade": st.column_config.NumberColumn(width=80),
                    "Valor unitário": st.column_config.TextColumn(width=120),
                    "Valor total": st.column_config.TextColumn(width=120),
                    "Gasto": st.column_config.TextColumn(width=120),
                    "Saldo": st.column_config.TextColumn(width=120),
                }
            )

        # --------------------------------------------------
        # ABRIR DIÁLOGO
        # --------------------------------------------------
        if st.session_state.get("abrir_dialogo_despesa"):
            dialog_relatos_fin()
            st.session_state["abrir_dialogo_despesa"] = False






    # ==================================================
    # MODO EDIÇÃO — CRUD DO ORÇAMENTO
    # ==================================================
    if modo_edicao:

        # -----------------------------------
        # Buscar categorias de despesa
        # -----------------------------------
        col_categorias_despesa = db["categorias_despesa"]

        categorias = list(
            col_categorias_despesa
            .find({}, {"categoria": 1})
            .sort("categoria", 1)
        )

        opcoes_categorias = [c["categoria"] for c in categorias]

        if not opcoes_categorias:
            st.warning(
                "Não há categorias de despesa cadastradas. "
                "Cadastre primeiro nas configurações auxiliares."
            )
            st.stop()

        # -----------------------------------
        # Dados atuais do orçamento
        # -----------------------------------
        orcamento_atual = financeiro.get("orcamento", [])

        if orcamento_atual:
            df_orcamento = pd.DataFrame(orcamento_atual)
        else:
            df_orcamento = pd.DataFrame(
                columns=[
                    "categoria",
                    "nome_despesa",
                    "descricao_despesa",
                    "unidade",
                    "quantidade",
                    "valor_unitario",
                    "id_despesa"
                ]
            )

        # -----------------------------------
        # Garantir colunas
        # -----------------------------------
        for col in [
            "categoria",
            "nome_despesa",
            "descricao_despesa",
            "unidade",
            "quantidade",
            "valor_unitario",
            "id_despesa"
        ]:
            if col not in df_orcamento.columns:
                df_orcamento[col] = None

        # -----------------------------------
        # Preencher valores nulos
        # -----------------------------------
        df_orcamento["quantidade"] = df_orcamento["quantidade"].fillna(0)
        df_orcamento["valor_unitario"] = df_orcamento["valor_unitario"].fillna(0)

        # -----------------------------------
        # Calcular valor total
        # -----------------------------------
        df_orcamento["valor_total"] = (
            df_orcamento["quantidade"] * df_orcamento["valor_unitario"]
        )

        # -----------------------------------
        # Formatação
        # -----------------------------------
        def format_decimal(valor):
            if pd.isna(valor):
                return ""
            valor = float(valor)
            return str(int(valor)) if valor.is_integer() else str(valor).replace(".", ",")

        df_orcamento["quantidade_fmt"] = df_orcamento["quantidade"].apply(format_decimal)
        df_orcamento["valor_unitario_fmt"] = df_orcamento["valor_unitario"].apply(format_brl)
        df_orcamento["valor_total_fmt"] = df_orcamento["valor_total"].apply(format_brl)

        # -----------------------------------
        # Ordenar
        # -----------------------------------
        df_orcamento = df_orcamento.sort_values("categoria", ignore_index=True)

        # -----------------------------------
        # Inicializar estado do editor
        # -----------------------------------
        if "df_orcamento_editor" not in st.session_state:
            st.session_state["df_orcamento_editor"] = df_orcamento.copy()

        # -----------------------------------
        # Editor
        # -----------------------------------
        df_editado_orc = st.data_editor(
            st.session_state["df_orcamento_editor"][
                [
                    "id_despesa",
                    "categoria",
                    "nome_despesa",
                    "descricao_despesa",
                    "unidade",
                    "quantidade_fmt",
                    "valor_unitario_fmt",
                    "valor_total_fmt",
                ]
            ],
            num_rows="dynamic",
            height="content",
            column_config={
                "id_despesa": None,
                "categoria": st.column_config.SelectboxColumn(
                    "Categoria de despesa",
                    options=opcoes_categorias,
                    required=True
                ),
                "nome_despesa": st.column_config.TextColumn(
                    "Despesa",
                    required=True
                ),
                "descricao_despesa": st.column_config.TextColumn("Descrição"),
                "unidade": st.column_config.TextColumn("Unidade"),
                "quantidade_fmt": st.column_config.TextColumn("Quantidade"),
                "valor_unitario_fmt": st.column_config.TextColumn("Valor unitário (R$)"),
                "valor_total_fmt": st.column_config.TextColumn(
                    "Valor total (auto)",
                    disabled=True
                ),
            },
            key="editor_orcamento",
        )


        # -----------------------------------
        # Conversões
        # -----------------------------------
        def parse_brl(valor):
            if not valor:
                return 0.0
            return float(str(valor).replace("R$", "").replace(".", "").replace(",", ".").strip())

        def parse_decimal(valor):
            if not valor:
                return 0.0
            return float(str(valor).replace(".", "").replace(",", ".").strip())




        # -----------------------------------
        # Resumo financeiro do orçamento
        # -----------------------------------

        # -----------------------------------
        # Valor total do projeto (ajustado)
        # -----------------------------------
        valor_total_base = financeiro.get("valor_total") or 0.0
        valor_aditivo = financeiro.get("valor_aditivo") or 0.0
        valor_devolucao = financeiro.get("valor_devolucao") or 0.0

        valor_total = valor_total_base + valor_aditivo - valor_devolucao

        # -----------------------------------
        # Soma das despesas (a partir do editor)
        # -----------------------------------
        df_temp = df_editado_orc.copy()

        df_temp["quantidade"] = df_temp["quantidade_fmt"].apply(parse_decimal)
        df_temp["valor_unitario"] = df_temp["valor_unitario_fmt"].apply(parse_brl)

        df_temp["valor_total"] = (
            df_temp["quantidade"] * df_temp["valor_unitario"]
        )

        soma_despesas = df_temp["valor_total"].sum()

        # -----------------------------------
        # Formatação
        # -----------------------------------
        valor_total_fmt = format_brl(valor_total)
        soma_despesas_fmt = format_brl(soma_despesas)







        # -----------------------------------
        # BOTÃO ATUALIZAR
        # -----------------------------------
        with st.container(horizontal=True, horizontal_alignment="right"):

            if st.button("Atualizar tabela", icon=":material/sync:", width=200):

                df_temp = df_editado_orc.copy()

                df_temp["quantidade"] = df_temp["quantidade_fmt"].apply(parse_decimal)
                df_temp["valor_unitario"] = df_temp["valor_unitario_fmt"].apply(parse_brl)

                df_temp["valor_total"] = (
                    df_temp["quantidade"] * df_temp["valor_unitario"]
                )

                df_temp["valor_total_fmt"] = df_temp["valor_total"].apply(format_brl)

                # Atualiza estado corretamente (sem erro de widget)
                st.session_state["df_orcamento_editor"] = df_temp

                st.rerun()

            





        # -----------------------------------
        # Exibição do resumo financeiro em baixo
        # -----------------------------------

        col1, col2, col3 = st.columns(3)

        # -----------------------------------
        # Valor do projeto
        # -----------------------------------
        col1.markdown(
            f"""
            <div>
                <div style="font-size:14px;">Valor do projeto</div>
                <div style="font-size:22px; font-weight:600;">
                    {valor_total_fmt}
                </div>
            </div>
            """,
            unsafe_allow_html=True
        )

        # -----------------------------------
        # Soma das despesas
        # -----------------------------------
        col2.markdown(
            f"""
            <div>
                <div style="font-size:14px;">Soma das despesas</div>
                <div style="font-size:22px; font-weight:600;">
                    {soma_despesas_fmt}
                </div>
            </div>
            """,
            unsafe_allow_html=True
        )

        # -----------------------------------
        # Diferença (com cor dinâmica)
        # -----------------------------------
        diferenca = valor_total - soma_despesas

        if diferenca != 0:

            diferenca_fmt = format_brl(diferenca)

            # Definir cor
            cor = "red" if diferenca < 0 else "green"

            col3.markdown(
                f"""
                <div>
                    <div style="font-size:14px;">Diferença</div>
                    <div style="font-size:22px; font-weight:600; color:{cor};">
                        {diferenca_fmt}
                    </div>
                </div>
                """,
                unsafe_allow_html=True
            )

        st.write('')
        st.write('')


        # -----------------------------------
        # SALVAR
        # -----------------------------------

        with st.container(horizontal=True, horizontal_alignment="right"):
          
            botao_salvar_orcamento = st.button("Salvar orçamento", icon=":material/save:", type="primary", width=200)
            
            
            
        
        if botao_salvar_orcamento:


            # -----------------------------------
            # Atualizar tabela antes de salvar (mesma lógica do botão Atualizar)
            # -----------------------------------
            df_temp = df_editado_orc.copy()

            df_temp["quantidade"] = df_temp["quantidade_fmt"].apply(parse_decimal)
            df_temp["valor_unitario"] = df_temp["valor_unitario_fmt"].apply(parse_brl)

            df_temp["valor_total"] = (
                df_temp["quantidade"] * df_temp["valor_unitario"]
            )

            df_temp["valor_total_fmt"] = df_temp["valor_total"].apply(format_brl)

            # Atualizar estado visual
            st.session_state["df_orcamento_editor"] = df_temp


            df_salvar = df_editado_orc.dropna(
                subset=["categoria", "nome_despesa"],
                how="any"
            ).copy()

            # -----------------------------------
            # Validação
            # -----------------------------------
            campos_obrigatorios = [
                "categoria",
                "nome_despesa",
                "descricao_despesa",
                "unidade",
                "quantidade_fmt",
                "valor_unitario_fmt",
            ]

            nomes_legiveis = {
                "categoria": "Categoria",
                "nome_despesa": "Despesa",
                "descricao_despesa": "Descrição",
                "unidade": "Unidade",
                "quantidade_fmt": "Quantidade",
                "valor_unitario_fmt": "Valor unitário",
            }

            erros = []

            for idx, row in df_salvar.iterrows():
                for campo in campos_obrigatorios:
                    valor = row.get(campo)
                    if pd.isna(valor) or str(valor).strip() == "":
                        erros.append(f"Linha {idx+1}: campo '{nomes_legiveis.get(campo)}' não preenchido.")

            if erros:
                st.error("Existem campos obrigatórios não preenchidos.", icon=":material/error:")
                for erro in erros:
                    st.write(f"- {erro}")
                st.stop()

            # -----------------------------------
            # Conversões finais
            # -----------------------------------
            df_salvar["quantidade"] = df_salvar["quantidade_fmt"].apply(parse_decimal)
            df_salvar["valor_unitario"] = df_salvar["valor_unitario_fmt"].apply(parse_brl)
            df_salvar["valor_total"] = df_salvar["quantidade"] * df_salvar["valor_unitario"]


            # -----------------------------------
            # Validação: soma das despesas vs valor do projeto
            # -----------------------------------

            # Valor total do projeto (ajustado)
            valor_total_base = financeiro.get("valor_total") or 0.0
            valor_aditivo = financeiro.get("valor_aditivo") or 0.0
            valor_devolucao = financeiro.get("valor_devolucao") or 0.0

            valor_total_projeto = valor_total_base + valor_aditivo - valor_devolucao

            # Soma das despesas
            soma_despesas = df_salvar["valor_total"].sum()

            # -----------------------------------
            # Comparação com tolerância de centavos
            # -----------------------------------
            if round(soma_despesas, 2) != round(valor_total_projeto, 2):

                soma_fmt = format_brl(soma_despesas).replace('$', '\\$')
                total_fmt = format_brl(valor_total_projeto).replace('$', '\\$')

                st.error(
                    f"A soma das despesas ({soma_fmt}) deve ser igual ao valor do projeto ({total_fmt}). **O orçamento não foi salvo**.",
                    icon=":material/error:"
                )

                st.stop()



            # -----------------------------------
            # Merge com banco (preservando dados)
            # -----------------------------------
            mapa_antigo = {
                item.get("id_despesa"): item
                for item in orcamento_atual
                if item.get("id_despesa")
            }

            novo_orcamento = []

            for _, row in df_salvar.iterrows():

                id_despesa = row.get("id_despesa")
                if pd.isna(id_despesa) or not id_despesa:
                    id_despesa = str(uuid.uuid4())
                else:
                    id_despesa = str(id_despesa)

                item_existente = mapa_antigo.get(id_despesa, {})

                item_atualizado = item_existente.copy()

                item_atualizado.update({
                    "id_despesa": id_despesa,
                    "categoria": row["categoria"],
                    "nome_despesa": row["nome_despesa"],
                    "descricao_despesa": row.get("descricao_despesa"),
                    "unidade": row.get("unidade"),
                    "quantidade": float(row["quantidade"]),
                    "valor_unitario": float(row["valor_unitario"]),
                    "valor_total": float(row["valor_total"]),
                })

                # Garantir vínculo com lançamentos
                for lanc in item_atualizado.get("lancamentos", []):
                    lanc["id_despesa"] = id_despesa

                novo_orcamento.append(item_atualizado)

            # -----------------------------------
            # Persistência
            # -----------------------------------
            col_projetos.update_one(
                {"codigo": codigo_projeto_atual},
                {"$set": {"financeiro.orcamento": novo_orcamento}}
            )

            st.success("Orçamento salvo com sucesso!", icon=":material/check:")
            time.sleep(3)
            st.rerun()







# --------------------------------------------------
# ABA RECIBOS
# --------------------------------------------------

if usuario_interno:

    # ==================================================
    # CONTROLE DE ESTADO
    # Permite abrir apenas um uploader de recibo por vez
    # ==================================================
    if "recibo_aberto_parcela" not in st.session_state:
        st.session_state["recibo_aberto_parcela"] = None

    # Controla a substituição dos botões de gerar recibo para baixar recibo.
    if "recibos_gerados" not in st.session_state:
        st.session_state["recibos_gerados"] = {}


    with recibos:

        st.markdown("### Recibos")
        st.caption("Ao **guardar o recibo**, é necessário informar a **data de pagamento** (para a evolução do cronograma do projeto).")
        st.write("")

        # --------------------------------------------------
        # Parcelas do projeto
        # --------------------------------------------------
        parcelas = financeiro.get("parcelas", [])

        if not parcelas:
            st.caption("Não há parcelas cadastradas.")
            st.stop()



        # --------------------------------------------------
        # Nome do investidor para o recibo
        # --------------------------------------------------
        nome_investidor, erros_investidor = obter_nome_investidor(db, projeto)



        # ==================================================
        # Layout único das colunas
        # (alterar aqui muda tudo)
        # ==================================================
        LAYOUT_COLUNAS_RECIBOS = [1.5, 2, 2, 2, 5]

        # ==================================================
        # Conexão com Google Drive
        # ==================================================
        servico_drive = obter_servico_drive()

        pasta_projeto_id = obter_pasta_projeto(
            servico_drive,
            projeto["codigo"],
            projeto["sigla"]
        )

        pasta_recibos_id = obter_pasta_recibos(
            servico_drive,
            pasta_projeto_id
        )

        # ==================================================
        # LINHAS (uma por parcela)
        # ==================================================
        for parcela in parcelas:

            numero = parcela.get("numero")

            col1, col2, col3, col4, col5 = st.columns(LAYOUT_COLUNAS_RECIBOS)

            # --------------------------------------------------
            # COLUNA 1 — Identificação da parcela
            # --------------------------------------------------
            col1.write(f"**Parcela {numero}**")

            # --------------------------------------------------
            # COLUNA 2 — Gerar recibo
            # --------------------------------------------------
 

            with col2:

                codigo = projeto["codigo"]
                chave = f"{numero}_{codigo}"
                contatos = projeto.get("contatos", [])


                caminho = f"/tmp/recibo_parcela_{numero}_{codigo}.docx"

                # Inicializa estado
                if "recibos_gerados" not in st.session_state:
                    st.session_state["recibos_gerados"] = {}

                # Filtra contatos aptos para assinatura
                contatos_assinam = [
                    c for c in contatos
                    if c.get("assina_docs", False) is True
                ]



                # ==========================
                # GERAR RECIBO
                # ==========================
                if chave not in st.session_state["recibos_gerados"]:

                    if st.button(
                        "Gerar recibo",
                        key=f"gerar_{chave}",
                        width="stretch",
                        icon=":material/receipt_long:",
                        type="secondary"
                    ):

                        # --------------------------------------------------
                        # 1. Validação: contatos que assinam
                        # --------------------------------------------------
                        if not contatos_assinam:
                            st.error(
                                "Não há contatos cadastrados com a responsabilidade de assinar documentos. "
                                "Cadastre ao menos um contato com essa opção marcada antes de gerar o recibo."
                            )

                        else:
                            # --------------------------------------------------
                            # 2. Resolver nome do investidor
                            # (função utilitária, sem st.stop)
                            # --------------------------------------------------
                            nome_investidor, erros_investidor = obter_nome_investidor(db, projeto)

                            if not nome_investidor:
                                # Mostra erro, mas NÃO interrompe a aplicação
                                st.error(
                                    "Não foi possível identificar o investidor do projeto:\n\n"
                                    + "\n".join([f"- {e}" for e in erros_investidor]),
                                    icon=":material/error:"
                                )

                            else:

                                # --------------------------------------------------
                                # 3. Buscar organização (CNPJ)
                                # --------------------------------------------------

                                # Busca a organização diretamente pelo _id armazenado no projeto
                                organizacao_doc = col_organizacoes.find_one(
                                    {"_id": projeto["id_organizacao"]},
                                    {"cnpj": 1}
                                )



                                if not organizacao_doc:
                                    st.error(
                                        f"Não foi encontrado CNPJ para a organização '{projeto['organizacao']}'."
                                    )

                                else:
                                    cnpj_organizacao = organizacao_doc.get("cnpj")


                                    # --------------------------------------------------
                                    # 4. Gerar recibo DOCX
                                    # --------------------------------------------------

                                    # Recupera nome da organização via mapa
                                    nome_organizacao = mapa_org_id_nome.get(projeto["id_organizacao"], "")

                                    sucesso = gerar_recibo_docx(
                                        caminho_arquivo=caminho,
                                        valor_parcela=parcela.get("valor", 0),
                                        numero_parcela=numero,
                                        nome_projeto=projeto["nome_do_projeto"],
                                        data_assinatura_contrato=projeto.get("contrato_data_assinatura"),
                                        contatos=contatos_assinam,
                                        nome_organizacao=nome_organizacao,
                                        cnpj_organizacao=cnpj_organizacao,
                                        contrato_nome=projeto.get("contrato_nome"),
                                        nome_investidor=nome_investidor
                                    )


                                    # --------------------------------------------------
                                    # 5. Pós-geração
                                    # --------------------------------------------------
                                    if sucesso:
                                        st.session_state["recibos_gerados"][chave] = caminho
                                        st.rerun()



       
       
                # ==========================
                # BAIXAR RECIBO
                # ==========================
                else:
                    with open(caminho, "rb") as f:
                        st.download_button(
                            label="Baixar recibo",
                            data=f,
                            file_name=f"recibo_parcela_{numero}_{codigo}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key=f"baixar_{chave}",
                            icon=":material/download:",
                            type="primary",
                            use_container_width=True
                        )

                    st.caption("Clique para baixar")



            # --------------------------------------------------
            # COLUNA 3 — Guardar recibo (abre uploader)
            # --------------------------------------------------
            if col3.button(
                "Guardar recibo",
                key=f"abrir_uploader_{numero}",
                width="stretch",
                icon=":material/save:"
            ):
                st.session_state["recibo_aberto_parcela"] = numero


            # --------------------------------------------------
            # COLUNA 4 + 5 — Link do recibo (se existir)
            # Agora o recibo está DENTRO da parcela
            # --------------------------------------------------
            recibo = parcela.get("recibo")

            if recibo:
                id_recibo = recibo.get("id_recibo")
                nome_arquivo = recibo.get("nome_arquivo", "Recibo")

                if id_recibo:
                    col4.write(":material/check: Recibo salvo")
                    link = gerar_link_drive(id_recibo)
                    col5.markdown(f"[{nome_arquivo}]({link})")








            # ==================================================
            # BLOCO DE UPLOAD (abre somente para a parcela ativa)
            # ==================================================
            if st.session_state["recibo_aberto_parcela"] == numero:
                st.write("")

                with st.container(border=True):

                    st.markdown(f"**Enviar recibo da Parcela {numero}**")

                    # --------------------------------------------------
                    # Data do pagamento (obrigatória)
                    # --------------------------------------------------

                    col1, col2 = st.columns([1, 4])

                    with col1:
                        data_pagamento = date_picker(
                            label="Data do pagamento",
                            format="dd/MM/yyyy",
                            locale="pt_BR",
                            one_tap=True,
                            key=f"data_pagamento_{numero}",
                        )



                    # --------------------------------------------------
                    # Upload do arquivo
                    # --------------------------------------------------
                    arquivo = st.file_uploader(
                        "Selecione o arquivo do recibo:",
                        type=["pdf", "png", "jpg", "jpeg", "docx", "doc"],
                        key=f"uploader_recibo_{numero}"
                    )

                    st.write("")

                    with st.container(horizontal=True):

                        # ----------------------------
                        # SALVAR
                        # ----------------------------
                        if st.button(
                            "Salvar recibo",
                            key=f"salvar_recibo_{numero}",
                            type="primary",
                            icon=":material/save:",
                            width=180
                        ):

                            # -----------------------------------
                            # Validação do arquivo
                            # -----------------------------------
                            if not arquivo:
                                st.error("Selecione um arquivo antes de salvar.", icon=":material/warning:")
                                st.stop()

                            # -----------------------------------
                            # Validação da data
                            # -----------------------------------
                            if not data_pagamento:
                                st.error("Selecione a data do pagamento.")
                                st.stop()

                            # -----------------------------------
                            # Upload no Google Drive
                            # -----------------------------------
                            id_arquivo = enviar_arquivo_drive(
                                servico_drive,
                                pasta_recibos_id,
                                arquivo
                            )

                            if not id_arquivo:
                                st.stop()



                            # -----------------------------------
                            # SALVAR NO MONGODB
                            # -----------------------------------
                            col_projetos.update_one(
                                {
                                    "codigo": codigo_projeto_atual,
                                    "financeiro.parcelas.numero": numero
                                },
                                {
                                    "$set": {
                                        "financeiro.parcelas.$.recibo": {
                                            "id_recibo": id_arquivo,
                                            "nome_arquivo": arquivo.name
                                        },
                                        "financeiro.parcelas.$.data_realizada": data_pagamento.strftime("%d/%m/%Y")
                                    }
                                }
                            )



                            # -----------------------------------
                            # Feedback + reset de estado
                            # -----------------------------------
                            st.success("Recibo salvo com sucesso!", icon=":material/check:")
                            st.session_state["recibo_aberto_parcela"] = None
                            time.sleep(3)
                            st.rerun()

                        # ----------------------------
                        # CANCELAR
                        # ----------------------------
                        if st.button(
                            "Cancelar",
                            key=f"cancelar_{numero}",
                            icon=":material/close:",
                            width=180
                        ):
                            st.session_state["recibo_aberto_parcela"] = None
                            st.rerun()


            st.divider()








# --------------------------------------------------
# ABA REMANEJAMENTOS
# --------------------------------------------------






with remanejamentos:

    st.markdown("### Remanejamentos")





    @st.fragment
    def fragmento_remanejamento(financeiro):

        with st.container(border=True):

            st.markdown("##### Nova solicitação de remanejamento")

            st.write('')

            # ==================================================
            # Segurança
            # ==================================================
            orcamento = financeiro.get("orcamento", []) or []

            if not orcamento:
                st.info("Não há linhas de despesa cadastradas no orçamento.")
                return

            # ==================================================
            # Montar opções + saldos
            # ==================================================
            opcoes_base = [" "]
            mapa_saldos = {}

            for item in orcamento:

                nome = item.get("nome_despesa", "Despesa")
                valor_total = item.get("valor_total", 0) or 0

                gasto = sum(
                    l.get("valor_despesa", 0)
                    for l in item.get("lancamentos", [])
                    if l.get("valor_despesa") is not None
                )

                saldo = valor_total - gasto

                saldo_fmt = (
                    f"R$ {saldo:,.2f}"
                    .replace(",", "X")
                    .replace(".", ",")
                    .replace("X", ".")
                )

                label = f"{nome} (saldo: {saldo_fmt})"

                opcoes_base.append(label)
                mapa_saldos[label] = saldo

            # ==================================================
            # Estados
            # ==================================================
            if "reducoes" not in st.session_state:
                st.session_state["reducoes"] = [{"despesa": " ", "valor": 0.0}]

            if "aumentos" not in st.session_state:
                st.session_state["aumentos"] = [{"despesa": " ", "valor": 0.0}]

            reducoes = st.session_state["reducoes"]
            aumentos = st.session_state["aumentos"]

            # ==================================================
            # Layout principal
            # ==================================================
            col_form, col_resumo = st.columns([3, 1], gap="large")

            # ==================================================
            # Função render linhas
            # ==================================================
            def render_linhas(lista, prefixo, tipo):

                total = 0

                for i in range(len(lista)):

                    usadas = {
                        x["despesa"]
                        for x in (reducoes + aumentos)
                        if x["despesa"] != " "
                    }

                    atual = lista[i]["despesa"]

                    opcoes_local = [
                        o for o in opcoes_base
                        if o == " " or o == atual or o not in usadas
                    ]

                    c1, c2, c3 = st.columns([3, 1, 1])

                    with c1:
                        escolha = st.selectbox(
                            "Despesa",
                            options=opcoes_local,
                            key=f"{prefixo}_sel_{i}"
                        )

                    lista[i]["despesa"] = escolha

                    if escolha != " ":

                        saldo = mapa_saldos.get(escolha, 0)


                        with c2:
                            label_valor = "Valor reduzido" if tipo == "reduzir" else "Valor aumentado"

                            valor = st.number_input(
                                label_valor,
                                min_value=0.0,
                                max_value=float(saldo) if tipo == "reduzir" else None,
                                step=100.0,
                                key=f"{prefixo}_val_{i}"
                            )



                        lista[i]["valor"] = valor

                        with c3:
                            saldo_final = saldo - valor if tipo == "reduzir" else saldo + valor

                            st.markdown(
                                f"Saldo final:  \nR$ {saldo_final:,.2f}"
                                .replace(",", "X").replace(".", ",").replace("X", ".")
                            )

                        total += valor

                    else:
                        lista[i]["valor"] = 0

                # ==================================================
                # Botão adicionar linha CENTRALIZADO
                # ==================================================
                with st.container(horizontal=True, horizontal_alignment="center"):
                    if st.button(
                        "+ Adicionar linha",
                        key=f"{prefixo}_add",
                        icon=":material/add:",
                        type="tertiary"
                    ):
                        lista.append({"despesa": " ", "valor": 0.0})
                        st.rerun()

                return total

            # ==================================================
            # COLUNA ESQUERDA — FORMULÁRIO
            # ==================================================
            with col_form:

                st.write("**:material/arrow_downward: Reduzir despesas**")
                total_reduzido = render_linhas(reducoes, "red", "reduzir")

                st.divider()

                st.write("**:material/arrow_upward: Aumentar despesas**")
                total_aumentado = render_linhas(aumentos, "aum", "aumentar")

                st.divider()

                with st.container(horizontal=True, vertical_alignment="bottom"):

                    justificativa = st.text_area(
                        "**Justificativa do remanejamento:**"
                    )

                    st.button(
                        "Salvar justificativa",
                    )

            # ==================================================
            # COLUNA DIREITA — RESUMO
            # ==================================================
            with col_resumo:

                st.write('')
                st.write('**Resumo do remanejamento**')
                st.write('')

                fmt = lambda v: f"R$ {v:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")

                st.metric(":material/arrow_downward: Total reduzido", fmt(total_reduzido))
                st.metric(":material/arrow_upward: Total aumentado", fmt(total_aumentado))
                
                saldo_remanejamento = total_reduzido - total_aumentado

                st.write("")
                st.write("")



                # ==================================================
                # MENSAGENS DINÂMICAS
                # ==================================================


                # --------------------------------------------------
                # CASO 1 — Nenhuma despesa foi preenchida ainda
                # --------------------------------------------------
                # Se ambos totais são zero, o usuário ainda não
                # selecionou linhas para reduzir ou aumentar.
                # Mostramos apenas uma instrução simples.
                # --------------------------------------------------
                if total_reduzido == 0 and total_aumentado == 0:

                    st.markdown("<span style='color:green; font-size:120%'>*Selecione despesas para reduzir e aumentar os valores*</span>", unsafe_allow_html=True)


                # --------------------------------------------------
                # CASO 2 — Valores preenchidos, mas NÃO batem
                # --------------------------------------------------
                # Aqui existe diferença entre o total reduzido
                # e o total aumentado.
                #
                # saldo_remanejamento = reduzido - aumentado
                #
                # > 0  → ainda pode aumentar
                # < 0  → aumentou demais, precisa reduzir
                # --------------------------------------------------
                elif saldo_remanejamento != 0:

                    if saldo_remanejamento > 0:
                        # Ainda sobra valor para aumentar

                        st.markdown(f"""
                                    <span style=
                                        'color:green; 
                                        font-size:120%; 
                                        margin-bottom:0
                                    '>
                                        :material/arrow_upward: Pode aumentar:
                                    </span>
                                    <br>
                                    <span style=
                                        'color:green; 
                                        font-size:220%
                                    '>
                                        R$ {saldo_remanejamento:,.2f} 
                                    </span>
                                    """
                                    , unsafe_allow_html=True)



                    else:
                        # Passou do limite

                        st.markdown(f"""
                                    <span style=
                                        'color:red; 
                                        font-size:120%; 
                                        margin-bottom:0
                                    '>
                                        :material/arrow_downward: Passou um pouco. <br>Reduza mais despesas ou diminua o aumento.
                                    </span>
                                    <br>
                                    <span style=
                                        'color:red; 
                                        font-size:220%
                                    '>
                                        R$ {saldo_remanejamento:,.2f} 
                                    </span>
                                    """
                                    , unsafe_allow_html=True)




                # --------------------------------------------------
                # CASO 3 — Valores batem (saldo == 0)
                # --------------------------------------------------
                # Agora o remanejamento está matematicamente correto.
                # Próximo passo: validar a justificativa.
                # --------------------------------------------------
                else:

                    # ----------------------------------------------
                    # Subcaso 3A — sem justificativa
                    # ----------------------------------------------
                    # Mesmo com as contas corretas, NÃO liberamos
                    # o botão de envio sem justificativa.
                    # ----------------------------------------------
                    if not justificativa.strip():


                        st.markdown(f"""
                                    <span style=
                                        'color:green; 
                                        font-size:120%; 
                                        margin-bottom:0
                                    '>
                                        :material/check: As contas estão corretas
                                    </span>
                                    """
                                    , unsafe_allow_html=True)



                        st.markdown(f"""
                                    <span style=
                                        'color:red; 
                                        font-size:120%
                                    '>
                                        :material/warning: Preencha a justificativa

                                    </span>
                                    """
                                    , unsafe_allow_html=True)




                    # ----------------------------------------------
                    # Subcaso 3B — tudo válido
                    # ----------------------------------------------
                    # • totais batem
                    # • justificativa preenchida
                    # → liberar botão de envio
                    # ----------------------------------------------
                    else:

                        st.markdown(f"""
                                    <span style=
                                        'color:green; 
                                        font-size:120%; 
                                        margin-bottom:0
                                    '>
                                        :material/check: Tudo certo. Pode enviar.
                                    </span>
                                    """
                                    , unsafe_allow_html=True)

                        st.write('')

                        # ------------------------------------------
                        # Botão de envio da solicitação
                        # ------------------------------------------
                        if st.button(
                            "Enviar solicitação",
                            icon=":material/mail:",
                            type="primary",
                            use_container_width=True
                        ):

   
                            # --------------------------------------
                            # Montar lista de REDUÇÕES
                            # --------------------------------------
                            # Guarda somente:
                            # • nome_despesa
                            # • valor_reduzido
                            # Ignora linhas vazias ou valor 0
                            # --------------------------------------
                            reduzidas = [
                                {
                                    "nome_despesa": r["despesa"].split(" (")[0],
                                    "valor_reduzido": r["valor"]
                                }
                                for r in reducoes if r["valor"] > 0
                            ]


                            # --------------------------------------
                            # Montar lista de AUMENTOS
                            # --------------------------------------
                            aumentadas = [
                                {
                                    "nome_despesa": a["despesa"].split(" (")[0],
                                    "valor_aumentado": a["valor"]
                                }
                                for a in aumentos if a["valor"] > 0
                            ]



                            # ------------------------------------------------------
                            # REGRA AUTOMÁTICA DE STATUS
                            # ------------------------------------------------------
                            # • soma o total já remanejado no projeto
                            # • adiciona o valor desta solicitação
                            # • compara com 15% do valor total do projeto
                            # ------------------------------------------------------

                            valor_remanejado_atual = financeiro.get("valor_remanejado", 0) or 0
                            valor_total_projeto = financeiro.get("valor_total", 0) or 0

                            novo_total_remanejado = valor_remanejado_atual + float(total_reduzido)

                            limite_remanejamento = valor_total_projeto * 0.15

                            if novo_total_remanejado <= limite_remanejamento:
                                status_remanejamento = "aceito"
                            else:
                                status_remanejamento = "em_analise"


                            # ------------------------------------------------------
                            # Estrutura do registro no MongoDB
                            # ------------------------------------------------------
                            registro = {
                                "data_solicit_remanej": datetime.datetime.now(datetime.UTC),
                                "status_remanejamento": status_remanejamento,
                                "justificativa": justificativa.strip(),
                                "reduzidas": reduzidas,
                                "aumentadas": aumentadas
                            }




                            # --------------------------------------
                            # Salvar no banco
                            # --------------------------------------
                            col_projetos.update_one(
                                {"codigo": codigo_projeto_atual},
                                {
                                    "$push": {
                                        "financeiro.remanejamentos_financeiros": registro
                                    },
                                    "$inc": {
                                        "financeiro.valor_remanejado": float(total_reduzido)
                                    }
                                }
                            )


                            # --------------------------------------------------
                            # Se aprovado automaticamente, efetiva no orçamento
                            # --------------------------------------------------
                            if status_remanejamento == "aceito":
                                efetivar_remanejamento(
                                    col_projetos,
                                    codigo_projeto_atual,
                                    financeiro,
                                    reduzidas,
                                    aumentadas
                                )




                            # -----------------------------------
                            # Recupera nome da organização via mapa
                            # -----------------------------------
                            organizacao_nome = mapa_org_id_nome.get(
                                projeto.get("id_organizacao"),
                                ""
                            )

                            # -----------------------------------
                            # Enviar e-mail
                            # -----------------------------------
                            enviar_email_remanejamento(
                                db,
                                projeto["codigo"],
                                projeto["sigla"],
                                projeto["nome_do_projeto"],
                                organizacao_nome,
                                reduzidas,
                                aumentadas,
                                status_remanejamento
                            )


                            # --------------------------------------
                            # Feedback visual + reset de estado
                            # --------------------------------------
                            st.success("Solicitação enviada.", icon=":material/check:")


                            # limpa listas para novo formulário
                            st.session_state["reducoes"] = [{"despesa": " ", "valor": 0.0}]
                            st.session_state["aumentos"] = [{"despesa": " ", "valor": 0.0}]

                            # Aguarda
                            time.sleep(3)

                            # Fecha o fragment de novo remanejamento
                            st.session_state["mostrar_remanejamento"] = False


                            # recarrega apenas o fragmento
                            st.rerun()













    # --------------------------------------------------
    # Estado de visibilidade do fragment
    # --------------------------------------------------
    if "mostrar_remanejamento" not in st.session_state:
        st.session_state["mostrar_remanejamento"] = False



    # --------------------------------------------------
    # Botão só para beneficiário
    # --------------------------------------------------
    if st.session_state.get("tipo_usuario") == "beneficiario":

        st.write("")


        # Verifica se já existe remanejamento em análise
        remanejamentos = financeiro.get("remanejamentos_financeiros", []) or []

        tem_pendente = any(
            r.get("status_remanejamento") == "em_analise"
            for r in remanejamentos
        )


        if st.button(
            "Solicitar remanejamento",
            icon=":material/compare_arrows:",
            disabled=tem_pendente
        ):

            st.session_state["mostrar_remanejamento"] = not st.session_state["mostrar_remanejamento"]


    # --------------------------------------------------
    # Fragmento (formulário)
    # --------------------------------------------------
    if st.session_state["mostrar_remanejamento"]:
        fragmento_remanejamento(financeiro)






    # --------------------------------------------------
    # Histórico de remanejamentos
    # --------------------------------------------------




    st.write("")
    st.write("")
    st.write("##### Histórico de remanejamentos")


    # --------------------------------------------------
    # Aviso quando remanejamentos excedem 15% do orçamento
    # --------------------------------------------------
    valor_remanejado = financeiro.get("valor_remanejado", 0) or 0
    valor_total = financeiro.get("valor_total", 0) or 0

    if valor_total and valor_remanejado > valor_total * 0.15:

        st.markdown(
            "<span style='color: #cb410b'>Os remanejamentos excederam 15% do orçamento do projeto. Portanto, as próximas solicitações serão analisadas pela equipe do Fundo.</span>",
            unsafe_allow_html=True
        )





    lista_remanej = financeiro.get("remanejamentos_financeiros", [])


    if not lista_remanej:
        st.caption("Nenhum remanejamento até o momento.")
    else:

        for idx in range(len(lista_remanej) - 1, -1, -1):
            item = lista_remanej[idx]

            data_solic = item.get("data_solicit_remanej")
            data_aprov = item.get("data_aprov_remanej")  # pode não existir
            status = item.get("status_remanejamento", "-")




            with st.container(border=True):

                # ==================================================
                # Função formatação de datas
                # ==================================================
                def fmt_data(dt):
                    if not dt:
                        return "-"
                    try:
                        return dt.astimezone().strftime("%d/%m/%Y")
                    except:
                        return str(dt)


                dist_colunas = [2,2,1]

                # ==================================================
                # LINHA 1 — DATAS + STATUS
                # ==================================================
                col1, col2, col3 = st.columns(dist_colunas)

                # ------------------------------------------
                # Coluna 1 — Data solicitação
                # ------------------------------------------
                with col1:
                    st.write(f"**Data da solicitação:** {fmt_data(data_solic)}")

                # ------------------------------------------
                # Coluna 2 — Data aceite
                # ------------------------------------------
                with col2:
                    if data_aprov:
                        st.write(f"**Data de aceite:** {fmt_data(data_aprov)}")


                # ------------------------------------------
                # Coluna 3 — Badge status
                # ------------------------------------------
                with col3:


                    if status == "aceito":
                        badge = {
                            "label": "Aceito",
                            "bg": "#D4EDDA",
                            "color": "#155724"
                        }

                    elif status == "em_analise":
                        badge = {
                            "label": "Em análise",
                            "bg": "#FFF3CD",
                            "color": "#856404"
                        }

                    elif status == "recusado":
                        badge = {
                            "label": "Recusado",
                            "bg": "#F8D7DA",     # vermelho claro
                            "color": "#721C24"  # vermelho escuro
                        }



                    st.markdown(
                        f"""
                        <div style="margin-bottom:6px;">
                            <span style="
                                background:{badge['bg']};
                                color:{badge['color']};
                                padding:4px 10px;
                                border-radius:20px;
                                font-size:12px;
                                font-weight:600;
                            ">
                                {badge['label']}
                            </span>
                        </div>
                        """,
                        unsafe_allow_html=True
                    )


                st.write('')





                # ==================================================
                # LINHA 2 — REDUÇÕES | AUMENTOS | AÇÕES
                # ==================================================
                col1, col2, col3 = st.columns(dist_colunas)


                # ------------------------------------------
                # Coluna 1 — Reduzidas
                # ------------------------------------------
                with col1:

                    st.write("**Despesas reduzidas:**")

                    reduzidas = item.get("reduzidas", [])

                    if not reduzidas:
                        st.caption("Nenhuma")
                    else:
                        for r in reduzidas:
                            st.write(
                                f"- {r['nome_despesa']}: "
                                + f"R$ {r['valor_reduzido']:,.2f}"
                                .replace(",", "X").replace(".", ",").replace("X", ".")
                            )


                # ------------------------------------------
                # Coluna 2 — Aumentadas
                # ------------------------------------------
                with col2:

                    st.write("**Despesas aumentadas:**")

                    aumentadas = item.get("aumentadas", [])

                    if not aumentadas:
                        st.caption("Nenhuma")
                    else:
                        for a in aumentadas:
                            st.write(
                                f"- {a['nome_despesa']}: "
                                + f"R$ {a['valor_aumentado']:,.2f}"
                                .replace(",", "X").replace(".", ",").replace("X", ".")
                            )


                # ------------------------------------------
                # Coluna 3 — Aceites + botão
                # ------------------------------------------


                with col3:

                    # --------------------------------------------------
                    # Ações somente para admin/equipe
                    # --------------------------------------------------


                    # --------------------------------------------------
                    # Mostrar ações somente para admin/equipe
                    # e somente se ainda NÃO aceito
                    # --------------------------------------------------
                    if st.session_state.get("tipo_usuario") in ["admin", "equipe"] and status != "aceito":

                        # ==================================================
                        # PRE-CARREGAR estados a partir do banco
                        # (ESSENCIAL para checkbox iniciar marcado)
                        # ==================================================

                        tec_key = f"tec_{idx}"
                        fin_key = f"fin_{idx}"

                        if tec_key not in st.session_state:
                            st.session_state[tec_key] = "aceite_tecnico" in item

                        if fin_key not in st.session_state:
                            st.session_state[fin_key] = "aceite_financeiro" in item


                        # ==================================================
                        # CHECKBOX TÉCNICO
                        # ==================================================
                        st.checkbox(
                            "Aceite técnico",
                            key=tec_key,
                            on_change=atualizar_aceite_remanejamento,
                            args=(codigo_projeto_atual, idx, "aceite_tecnico", tec_key)
                        )

                        if item.get("aceite_tecnico"):
                            st.caption(item["aceite_tecnico"])


                        # ==================================================
                        # CHECKBOX FINANCEIRO
                        # ==================================================
                        st.checkbox(
                            "Aceite financeiro",
                            key=fin_key,
                            on_change=atualizar_aceite_remanejamento,
                            args=(codigo_projeto_atual, idx, "aceite_financeiro", fin_key)
                        )

                        if item.get("aceite_financeiro"):
                            st.caption(item["aceite_financeiro"])


                        # ==================================================
                        # BOTÃO APROVAR
                        # ==================================================

                        habilitar = st.session_state[tec_key] and st.session_state[fin_key]


                        if st.button(
                            "Aprovar remanejamento",
                            disabled=not habilitar,
                            key=f"aprovar_{idx}",
                            type="primary",
                            icon=":material/check:",
                            width="stretch"
                        ):
                            aprovar_remanejamento(
                                codigo_projeto_atual,
                                idx,
                                item
                            )



                        # ==================================================
                        # BOTÃO RECUSAR (abre campo de justificativa)
                        # ==================================================
                        recusar_key = f"abrir_recusa_{idx}"

                        if recusar_key not in st.session_state:
                            st.session_state[recusar_key] = False


                        if st.button(
                            "Recusar remanejamento",
                            key=f"recusar_btn_{idx}",
                            type="secondary",
                            icon=":material/close:",
                            width="stretch"
                        ):
                            st.session_state[recusar_key] = True


                        # ==================================================
                        # FORMULÁRIO DE RECUSA
                        # ==================================================
                        if st.session_state[recusar_key]:

                            motivo = st.text_area(
                                "Justificativa da recusa:",
                                key=f"motivo_recusa_{idx}"
                            )

                            if st.button(
                                "Confirmar recusa",
                                key=f"confirmar_recusa_{idx}",
                                type="primary",
                                width="stretch"
                            ):

                                if not motivo.strip():
                                    st.warning("Informe a justificativa da recusa.")
                                else:

                                    nome = st.session_state.get("nome", "Usuário")
                                    data = datetime.datetime.now().strftime("%d/%m/%Y")

                                    log_recusa = f"Recusado por {nome} em {data}"

                                    col_projetos.update_one(
                                        {"codigo": codigo_projeto_atual},
                                        {
                                            "$set": {
                                                f"financeiro.remanejamentos_financeiros.{idx}.status_remanejamento": "recusado",
                                                f"financeiro.remanejamentos_financeiros.{idx}.motivo_recusa": motivo.strip(),
                                                f"financeiro.remanejamentos_financeiros.{idx}.log_recusa": log_recusa
                                            }
                                        }
                                    )

                                    projeto_atualizado = col_projetos.find_one({"codigo": codigo_projeto_atual})

                                    enviar_email_remanejamento_recusado(
                                        projeto_atualizado,
                                        projeto_atualizado["financeiro"]["remanejamentos_financeiros"][idx]
                                    )
              
                                    st.session_state[recusar_key] = False
                                    st.rerun()


                    # --------------------------------------------------
                    # Mostrar motivo da recusa para beneficiário
                    # --------------------------------------------------

                    if status == "recusado":

                        st.write("")

                        if item.get("log_recusa"):
                            st.caption(item["log_recusa"])

                        if item.get("motivo_recusa"):
                            # st.write("**Motivo da recusa:**")
                            st.write(item.get("motivo_recusa"))




              


                st.write('')


                # ==================================================
                # JUSTIFICATIVA (fora das colunas)
                # ==================================================
                st.write(f"**Justificativa:** {item.get('justificativa', '')}")







