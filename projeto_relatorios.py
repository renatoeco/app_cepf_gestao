import streamlit as st
import pandas as pd
import streamlit_antd_components as sac
import time
import datetime
from collections import defaultdict
import uuid
from io import BytesIO
from zoneinfo import ZoneInfo 
from docx import Document
import tempfile
import os
from st_rsuite import date_picker

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE



from funcoes_auxiliares import (
    conectar_mongo_cepf_gestao,
    sidebar_projeto,
    obter_servico_drive,
    obter_ou_criar_pasta,
    obter_pasta_pesquisas,
    obter_pasta_projeto,
    obter_pasta_relatos_financeiros,
    obter_pasta_relatorios,
    enviar_arquivo_drive,
    gerar_link_drive,
    enviar_email
)



st.set_page_config(page_title="Relatórios", page_icon=":material/edit_note:")





###########################################################################################################
# CONFIGURAÇÕES DO STREAMLIT
###########################################################################################################


# Traduzindo o texto do st.file_uploader
# Texto interno
st.markdown("""
<style>
/* Esconde o texto padrão */
[data-testid="stFileUploaderDropzone"] div div::before {
    content: "";
    color: rgba(49, 51, 63, 0.7);
    font-size: 0.9rem;
    font-weight: 400;
    position: absolute;
    top: 50px;              /* fixa no topo */
    left: 50%;
    transform: translate(-50%, 10%);
    pointer-events: none;
}
/* Esconde o texto original */
[data-testid="stFileUploaderDropzone"] div div span {
    visibility: hidden !important;
}
</style>
""", unsafe_allow_html=True)

# Traduzindo Botão do file_uploader
st.markdown("""
<style>
/* Alvo: apenas o botão dentro do componente de upload */
section[data-testid="stFileUploaderDropzone"] button[data-testid="stBaseButton-secondary"] {
    font-size: 0px !important;   /* esconde o texto original */
    padding-left: 14px !important;
    padding-right: 14px !important;
    min-width: 160px !important;
}
/* Insere o texto traduzido */
section[data-testid="stFileUploaderDropzone"] button[data-testid="stBaseButton-secondary"]::after {
    content: "Selecionar arquivo";
    font-size: 14px !important;
    color: inherit;
}
</style>
""", unsafe_allow_html=True)


###########################################################################################################
# CONEXÃO COM O BANCO DE DADOS MONGODB
###########################################################################################################

# Conecta-se ao banco de dados MongoDB (usa cache automático para melhorar performance)
db = conectar_mongo_cepf_gestao()




###########################################################################################################
# CARREGAMENTO DOS DADOS
###########################################################################################################

col_projetos = db["projetos"]

col_editais = db["editais"]

col_beneficios = db["beneficios"]

col_publicos = db["publicos"]

col_pessoas = db["pessoas"]

lista_publicos = list(col_publicos.find({}, {"_id": 0, "publico": 1}))

# SEMPRE insere a opção Outros
opcoes_publicos = sorted({p["publico"] for p in lista_publicos} - {"Outros"})
opcoes_publicos.append("Outros")

codigo_projeto_atual = st.session_state.projeto_atual

df_projeto = pd.DataFrame(
    list(
        col_projetos.find(
            {"codigo": codigo_projeto_atual}
        )
    )
)

if df_projeto.empty:
    st.error("Projeto não encontrado.")
    st.stop()

projeto = df_projeto.iloc[0]

relatorios = projeto.get("relatorios", [])

edital = col_editais.find_one({"codigo_edital": projeto["edital"]})

tipo_usuario = st.session_state.get("tipo_usuario")




###########################################################################################################
# FUNÇÕES
###########################################################################################################



# Função para configurar hyperlink no docx do relatório exportado.
def adicionar_hyperlink(paragraph, url, texto):
    """
    Adiciona um hyperlink clicável em um parágrafo do docx.
    """

    part = paragraph.part
    r_id = part.relate_to(
        url,
        RELATIONSHIP_TYPE.HYPERLINK,
        is_external=True
    )

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Estilo de link (azul + sublinhado)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')

    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')

    rPr.append(color)
    rPr.append(u)

    new_run.append(rPr)

    text = OxmlElement('w:t')
    text.text = texto

    new_run.append(text)
    hyperlink.append(new_run)

    paragraph._element.append(hyperlink)


def gerar_docx_relatorio(relatorio, projeto):
    """
    Gera um documento .docx contendo as informações do relatório,
    utilizando apenas o objeto de projeto como fonte principal.
    """

    with st.spinner("Gerando relatório..."):


        # ------------------------------
        # BUSCA ORGANIZAÇÃO
        # ------------------------------
        organizacao = db["organizacoes"].find_one(
            {"_id": projeto.get("id_organizacao")}
        )

        # ------------------------------
        # BUSCA EDITAL
        # ------------------------------
        edital = db["editais"].find_one(
            {"codigo_edital": projeto.get("edital")}
        )

        # Criação do documento
        doc = Document()




        # ------------------------------
        # CABEÇALHO DO RELATÓRIO
        # ------------------------------

        # Título principal
        titulo = doc.add_heading(
            f"Relatório {relatorio.get('numero')}",
            level=1
        )
        titulo.alignment = 1  # centralizado

        doc.add_paragraph("")

        # Edital
        doc.add_paragraph(
            f"Edital {edital.get('codigo_edital') if edital else ''}"
        )

        # Código + sigla do projeto
        doc.add_heading(
            f"{projeto.get('codigo')} - {projeto.get('sigla')}",
            level=2
        )

        doc.add_paragraph("")

        # Datas e informações gerais
        doc.add_paragraph(
            f"Data de envio: {relatorio.get('data_envio', 'N/A')}"
        )

        # Data de aprovação (não existe campo estruturado, então tenta extrair)
        data_aprovacao = "---"
        for componente in projeto.get("plano_trabalho", {}).get("componentes", []):
            for entrega in componente.get("entregas", []):
                for atividade in entrega.get("atividades", []):
                    for relato in atividade.get("relatos", []):
                        status = relato.get("status_aprovacao", "")
                        if "em" in status:
                            data_aprovacao = status.split("em")[-1].strip()

        doc.add_paragraph(f"Data de aprovação: {data_aprovacao}")

        # Data de exportação (data atual)
        data_exportacao = datetime.datetime.now().strftime("%d/%m/%Y")
        doc.add_paragraph(f"Data de exportação: {data_exportacao}")

        doc.add_paragraph("")

        # Organização
        nome_org = organizacao.get("nome_organizacao") if organizacao else ""
        doc.add_paragraph(f"Organização: {nome_org}")

        # Nome do projeto
        doc.add_paragraph(f"Projeto: {projeto.get('nome_do_projeto')}")

        # Status do projeto (não existe explícito, então inferência simples)
        status_projeto = "Em andamento"

        if projeto.get("data_fim_contrato"):
            try:
                data_fim = datetime.datetime.strptime(
                    projeto.get("data_fim_contrato"),
                    "%d/%m/%Y"
                )
                if data_fim < datetime.datetime.now():
                    status_projeto = "Encerrado"
            except:
                pass

        doc.add_paragraph(f"Status do projeto: {status_projeto}")

        doc.add_paragraph("")
        doc.add_paragraph("")





        # ------------------------------
        # RELATOS DE ATIVIDADES
        # ------------------------------
        doc.add_heading("Relatos de Atividades", level=2)

        tem_relato = False

        # Percorre estrutura do plano de trabalho
        for componente in projeto.get("plano_trabalho", {}).get("componentes", []):
            for entrega in componente.get("entregas", []):
                for atividade in entrega.get("atividades", []):


                    # Filtra apenas relatos do relatório atual
                    relatos = [
                        r for r in atividade.get("relatos", [])
                        if r.get("relatorio_numero") == relatorio.get("numero")
                    ]

                    # Verifica se existem relatos após o filtro
                    if relatos:

                        tem_relato = True

                        # Título da atividade
                        doc.add_heading(
                            f"Atividade: {atividade.get('atividade')}",
                            level=3
                        )

                        doc.add_paragraph("")

                        # Lista relatos
                        for relato in relatos:

                            doc.add_paragraph(
                                f"{relato.get('id_relato')}: {relato.get('relato')}"
                            )

                            doc.add_paragraph(
                                f"Status: {relato.get('status_relato')}"
                            )

                            doc.add_paragraph(
                                f"Data de início: {relato.get('data_inicio')}"
                            )

                            doc.add_paragraph(
                                f"Data de fim: {relato.get('data_fim')}"
                            )

                            doc.add_paragraph(
                                f"Progresso da atividade informado: {relato.get('porc_ativ_relato')}"
                            )




                            # ------------------------------
                            # ANEXOS DO RELATO
                            # ------------------------------
                            anexos = relato.get("anexos", [])

                            if anexos:
                                doc.add_paragraph("Anexos:")
                                    
                                for anexo in anexos:

                                    id_arquivo = anexo.get("id_arquivo")
                                    nome = anexo.get("nome_arquivo", "Arquivo")

                                    url = f"https://drive.google.com/file/d/{id_arquivo}/view"

                                    p = doc.add_paragraph()
                                    adicionar_hyperlink(p, url, nome)


                            # ------------------------------
                            # FOTOS DO RELATO
                            # ------------------------------
                            fotos = relato.get("fotos", [])

                            if fotos:
                                doc.add_paragraph("Fotos:")


                                for foto in fotos:

                                    id_arquivo = foto.get("id_arquivo")
                                    nome = foto.get("nome_arquivo", "Foto")

                                    url = f"https://drive.google.com/file/d/{id_arquivo}/view"

                                    p = doc.add_paragraph()
                                    adicionar_hyperlink(p, url, nome)




                            doc.add_paragraph("")  # espaçamento

        # Caso não existam relatos
        if not tem_relato:
            doc.add_paragraph("Não há relatos de atividades registrados para este projeto.")







        # ------------------------------
        # ESPAÇAMENTO
        # ------------------------------
        doc.add_paragraph("")
        doc.add_paragraph("")




        # ------------------------------
        # REGISTROS FINANCEIROS
        # ------------------------------
        doc.add_heading("Registros Financeiros", level=1)

        tem_despesa = False

        # Percorre orçamento do projeto
        for despesa in projeto.get("financeiro", {}).get("orcamento", []):

            # Filtra lançamentos do relatório atual
            lancamentos = [
                l for l in despesa.get("lancamentos", [])
                if l.get("relatorio_numero") == relatorio.get("numero")
            ]

            # Verifica se há lançamentos válidos
            if lancamentos:

                tem_despesa = True

                # Título da despesa
                doc.add_heading(
                    despesa.get("nome_despesa"),
                    level=2
                )

                # Lista lançamentos
                for lanc in lancamentos:

                    doc.add_paragraph(
                        f"{lanc.get('id_lanc_despesa')}: {lanc.get('descricao_despesa')}"
                    )

                    doc.add_paragraph(
                        f"Status: {lanc.get('status_despesa')}"
                    )

                    doc.add_paragraph(
                        f"Data da despesa: {lanc.get('data_despesa')}"
                    )

                    doc.add_paragraph(
                        f"Fornecedor: {lanc.get('fornecedor')}"
                    )

                    doc.add_paragraph(
                        f"CPF/CNPJ: {lanc.get('cpf_cnpj')}"
                    )

                    # ------------------------------
                    # TABELA DE VALORES
                    # ------------------------------
                    tabela = doc.add_table(rows=2, cols=3)

                    # Cabeçalhos
                    tabela.rows[0].cells[0].text = "Quantidade"
                    tabela.rows[0].cells[1].text = "Valor unitário"
                    tabela.rows[0].cells[2].text = "Valor da despesa"

                    # Valores
                    tabela.rows[1].cells[0].text = str(lanc.get("quantidade", ""))
                    tabela.rows[1].cells[1].text = str(lanc.get("valor_unitario", ""))
                    tabela.rows[1].cells[2].text = str(lanc.get("valor_despesa", ""))

                    # ------------------------------
                    # ANEXOS DO LANÇAMENTO
                    # ------------------------------
                    anexos = lanc.get("anexos", [])

                    if anexos:

                        doc.add_paragraph("")

                        doc.add_paragraph("Anexos:")

                        for anexo in anexos:

                            id_arquivo = anexo.get("id_arquivo")
                            nome = anexo.get("nome_arquivo", "Arquivo")

                            url = f"https://drive.google.com/file/d/{id_arquivo}/view"

                            p = doc.add_paragraph()
                            adicionar_hyperlink(p, url, nome)

                    doc.add_paragraph("")  # espaçamento entre lançamentos


        # Caso não existam registros financeiros
        if not tem_despesa:
            doc.add_paragraph("Não há registros financeiros para este relatório.")











        # ------------------------------
        # ESPAÇAMENTO
        # ------------------------------
        doc.add_paragraph("")
        doc.add_paragraph("")

        # ------------------------------
        # RESULTADOS
        # ------------------------------
        doc.add_heading("Resultados", level=1)

        # ------------------------------
        # INDICADORES DE PROJETO
        # ------------------------------
        doc.add_heading("Indicadores de projeto", level=2)

        doc.add_paragraph("")

        tem_indicador = False

        # Percorre componentes
        for componente in projeto.get("plano_trabalho", {}).get("componentes", []):

            entregas = componente.get("entregas", [])

            if entregas:

                # Título do componente
                paragrafo = doc.add_paragraph()
                run = paragrafo.add_run(
                    f"Componente: {componente.get('componente')}"
                )
                run.bold = True


                for entrega in entregas:

                    indicadores = entrega.get("indicadores_projeto", [])

                    if indicadores:

                        tem_indicador = True

                        # Título da entrega
                        doc.add_paragraph(
                            f"Entrega: {entrega.get('entrega')}"
                        )

                        for indicador in indicadores:

                            doc.add_paragraph(
                                f"Indicador: {indicador.get('indicador_projeto')}"
                            )

                            doc.add_paragraph(
                                f"Unidade de medida: {indicador.get('unidade_medida')}"
                            )

                            # ------------------------------
                            # "3 COLUNAS" SIMPLES
                            # ------------------------------
                            doc.add_paragraph(
                                f"Início do projeto: {indicador.get('linha_base')}    |    "
                                f"Meta: {indicador.get('meta')}    |    "
                                f"Resultado atual: {indicador.get('resultado_atual')}"
                            )

                            # ------------------------------
                            # DATA DE COLETA FORMATADA
                            # ------------------------------
                            data_coleta = indicador.get("data_coleta")

                            data_formatada = ""
                            if data_coleta:
                                try:
                                    data_formatada = data_coleta.strftime("%d/%m/%Y")
                                except:
                                    data_formatada = str(data_coleta)

                            doc.add_paragraph(
                                f"Último registro em {data_formatada}"
                            )

                            doc.add_paragraph("")  # espaçamento


        # Caso não existam indicadores
        if not tem_indicador:
            doc.add_paragraph("Não há indicadores registrados para este projeto.")








        # ------------------------------
        # ESPAÇAMENTO
        # ------------------------------
        doc.add_paragraph("")
        doc.add_paragraph("")

        # ------------------------------
        # BENEFICIÁRIOS
        # ------------------------------
        doc.add_heading("Beneficiários", level=1)

        doc.add_heading(
            "Número de beneficiários por gênero e faixa etária",
            level=2
        )

        doc.add_paragraph("")


        # Dados do relatório
        benef = relatorio.get("beneficiarios_quant", {})

        mulheres = benef.get("mulheres", {})
        homens = benef.get("homens", {})
        nao_binarios = benef.get("nao_binarios", {})

        # Totais por faixa etária
        total_jovens = (
            mulheres.get("jovens", 0)
            + homens.get("jovens", 0)
            + nao_binarios.get("jovens", 0)
        )

        total_adultos = (
            mulheres.get("adultas", 0)
            + homens.get("adultos", 0)
            + nao_binarios.get("adultos", 0)
        )

        total_idosos = (
            mulheres.get("idosas", 0)
            + homens.get("idosos", 0)
            + nao_binarios.get("idosos", 0)
        )

        # Totais por gênero
        total_mulheres = sum(mulheres.values())
        total_homens = sum(homens.values())
        total_nb = sum(nao_binarios.values())

        total_geral = total_mulheres + total_homens + total_nb

        # ------------------------------
        # TABELA (8 colunas x 4 linhas)
        # ------------------------------
        tabela = doc.add_table(rows=4, cols=8)

        # Linha 1 - Jovens
        tabela.rows[0].cells[0].text = "Mulheres jovens"
        tabela.rows[0].cells[1].text = str(mulheres.get("jovens", 0))

        tabela.rows[0].cells[2].text = "Homens jovens"
        tabela.rows[0].cells[3].text = str(homens.get("jovens", 0))

        tabela.rows[0].cells[4].text = "Não-binários jovens"
        tabela.rows[0].cells[5].text = str(nao_binarios.get("jovens", 0))

        tabela.rows[0].cells[6].text = "Total de jovens"
        tabela.rows[0].cells[7].text = str(total_jovens)

        # Linha 2 - Adultos
        tabela.rows[1].cells[0].text = "Mulheres adultas"
        tabela.rows[1].cells[1].text = str(mulheres.get("adultas", 0))

        tabela.rows[1].cells[2].text = "Homens adultos"
        tabela.rows[1].cells[3].text = str(homens.get("adultos", 0))

        tabela.rows[1].cells[4].text = "Não-binários adultos"
        tabela.rows[1].cells[5].text = str(nao_binarios.get("adultos", 0))

        tabela.rows[1].cells[6].text = "Total de adultos"
        tabela.rows[1].cells[7].text = str(total_adultos)

        # Linha 3 - Idosos
        tabela.rows[2].cells[0].text = "Mulheres idosas"
        tabela.rows[2].cells[1].text = str(mulheres.get("idosas", 0))

        tabela.rows[2].cells[2].text = "Homens idosos"
        tabela.rows[2].cells[3].text = str(homens.get("idosos", 0))

        tabela.rows[2].cells[4].text = "Não-binários idosos"
        tabela.rows[2].cells[5].text = str(nao_binarios.get("idosos", 0))

        tabela.rows[2].cells[6].text = "Total de idosos"
        tabela.rows[2].cells[7].text = str(total_idosos)

        # Linha 4 - Totais
        tabela.rows[3].cells[0].text = "Total de mulheres"
        tabela.rows[3].cells[1].text = str(total_mulheres)

        tabela.rows[3].cells[2].text = "Total de homens"
        tabela.rows[3].cells[3].text = str(total_homens)

        tabela.rows[3].cells[4].text = "Total de não-binários"
        tabela.rows[3].cells[5].text = str(total_nb)

        tabela.rows[3].cells[6].text = "Total geral"
        tabela.rows[3].cells[7].text = str(total_geral)













        doc.add_paragraph("")  # espaçamento
        doc.add_paragraph("")  # espaçamento


        # ------------------------------
        # SUBSEÇÃO: TIPOS DE BENEFICIÁRIOS
        # ------------------------------
        doc.add_heading(
            "Tipos de Beneficiários e Benefício",
            level=2
        )

        doc.add_paragraph("")  # espaçamento


        localidades = projeto.get("locais", {}).get("localidades", [])

        tem_localidade = False

        for loc in localidades:

            beneficiarios = loc.get("beneficiarios", [])

            if beneficiarios:

                tem_localidade = True

                # Nome da localidade em negrito
                p = doc.add_paragraph()
                run = p.add_run(loc.get("nome_localidade", ""))
                run.bold = True

                # Município
                doc.add_paragraph(loc.get("municipio", ""))

                # Lista de beneficiários
                for ben in beneficiarios:

                    tipo = ben.get("tipo_beneficiario", "")
                    beneficios = ben.get("beneficios", [])


                    # Tipo de beneficiário 
                    p = doc.add_paragraph()
                    run = p.add_run(f"{tipo}")
                    # run.bold = True

                    # Lista de benefícios
                    for b in beneficios:
                        doc.add_paragraph(b, style="List Bullet")

                    # doc.add_paragraph(
                    #     f"{tipo}    |    Benefícios: {', '.join(beneficios)}"
                    # )

                doc.add_paragraph("")  # espaçamento


        # Caso não existam dados
        if not tem_localidade:
            doc.add_paragraph("Não há registros de beneficiários por localidade.")











        # ------------------------------
        # ESPAÇAMENTO
        # ------------------------------
        doc.add_paragraph("")
        doc.add_paragraph("")

        # ------------------------------
        # SEÇÃO: PESQUISAS
        # ------------------------------
        doc.add_heading(
            "Pesquisas / Ferramentas de Monitoramento",
            level=1
        )

        doc.add_paragraph("")

        # ------------------------------
        # BUSCA EDITAL
        # ------------------------------
        edital = db["editais"].find_one(
            {"codigo_edital": projeto.get("edital")}
        )

        pesquisas_edital = edital.get("pesquisas_relatorio", []) if edital else []

        # Pesquisas respondidas no projeto
        pesquisas_projeto = projeto.get("pesquisas", [])

        # Indexa pesquisas do projeto por id (para busca rápida)
        mapa_pesquisas = {
            p.get("id_pesquisa"): p
            for p in pesquisas_projeto
        }

        tem_pesquisa = False

        # ------------------------------
        # LOOP NAS PESQUISAS DO EDITAL
        # ------------------------------
        for pesquisa in pesquisas_edital:

            tem_pesquisa = True

            id_pesquisa = pesquisa.get("id")
            nome = pesquisa.get("nome_pesquisa")

            # Nome da pesquisa
            doc.add_paragraph(nome)

            # Verifica se existe no projeto
            dados_proj = mapa_pesquisas.get(id_pesquisa)

            if dados_proj:
                respondida = "Sim" if dados_proj.get("respondida") else "Não"
                verificada = "Sim" if dados_proj.get("verificada") else "Não"
            else:
                respondida = "Não"
                verificada = "Não"

            # Status
            doc.add_paragraph(f"Respondida? {respondida}")
            doc.add_paragraph(f"Verificada? {verificada}")

            doc.add_paragraph("")  # espaçamento


        # Caso não existam pesquisas
        if not tem_pesquisa:
            doc.add_paragraph("Não há pesquisas vinculadas a este edital.")










        # ------------------------------
        # ESPAÇAMENTO
        # ------------------------------
        doc.add_paragraph("")
        doc.add_paragraph("")

        # ------------------------------
        # SEÇÃO: FORMULÁRIO
        # ------------------------------
        doc.add_heading("Formulário", level=1)

        doc.add_paragraph("")


        respostas = relatorio.get("respostas_formulario", {})

        tem_resposta = False

        # Ordena perguntas pela ordem
        perguntas_ordenadas = sorted(
            respostas.values(),
            key=lambda x: x.get("ordem", 0)
        )

        for item in perguntas_ordenadas:

            pergunta = item.get("pergunta", "")
            resposta = item.get("resposta", "")

            tem_resposta = True

            # Pergunta em negrito
            p = doc.add_paragraph()
            run = p.add_run(pergunta)
            run.bold = True

            # Resposta
            if isinstance(resposta, list):
                for r in resposta:
                    nome = r.get("nome", "")
                    doc.add_paragraph(nome, style="List Bullet")
            else:
                doc.add_paragraph(str(resposta))

            doc.add_paragraph("")  # espaçamento


        # Caso não existam respostas
        if not tem_resposta:
            doc.add_paragraph("Não há respostas registradas para este formulário.")












        # ------------------------------
        # SALVAMENTO TEMPORÁRIO
        # ------------------------------
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            caminho = tmp.name
            doc.save(caminho)

        # Garantia de integridade do arquivo
        time.sleep(3)

        # ------------------------------
        # DOWNLOAD
        # ------------------------------

        numero = relatorio.get("numero")
        codigo_projeto = projeto.get("codigo")
        sigla_org = organizacao.get("sigla_organizacao") if organizacao else "org"

        nome_arquivo = f"relatorio_{numero}_{codigo_projeto}_{sigla_org}.docx"



        # Remoção do arquivo temporário
        os.remove(caminho)

        return doc






def calcular_saldo_parcela():
    # ==================================================
    # CÁLCULO DO SALDO DA PARCELA
    # ==================================================
    # Regra:
    # - parcela = relatorio_numero
    # - saldo = valor da parcela - total gasto na parcela
    # - exibir em porcentagem (%)

    parcela_atual = next(
        (p for p in projeto.get("financeiro", {}).get("parcelas", [])
        if p.get("numero") == st.session_state.get("relatorio_numero")),
        None
    )

    if parcela_atual:

        valor_parcela = parcela_atual.get("valor", 0)

        # Soma todas as despesas desta parcela
        total_gasto = 0
        for despesa in projeto.get("financeiro", {}).get("orcamento", []):
            for lanc in despesa.get("lancamentos", []):
                if lanc.get("relatorio_numero") == relatorio_numero:
                    total_gasto += lanc.get("valor_despesa", 0)

        saldo = valor_parcela - total_gasto

        if valor_parcela > 0:
            saldo_pct = (saldo / valor_parcela) * 100
        else:
            saldo_pct = 0

        # Exibição 

        return saldo_pct




# Texto do status da avaliação de Relatos de Atividades ou de Despesas de relatório
def texto_verificacao():
    nome = st.session_state.get("nome", "Usuário")
    data = datetime.datetime.now().strftime("%d/%m/%Y")
    return f"Verificado por {nome} em {data}"


# Atualiza o status da avaliação de Relatos de Atividades ou de Despesas
def atualizar_verificacao_relatorio(projeto_codigo, relatorio_numero, campo, checkbox_key):
    marcado = st.session_state.get(checkbox_key, False)

    nome = st.session_state.get("nome", "Usuário")
    data = datetime.datetime.now().strftime("%d/%m/%Y")

    if marcado:
        col_projetos.update_one(
            {
                "codigo": projeto_codigo,
                "relatorios.numero": relatorio_numero
            },
            {
                "$set": {
                    f"relatorios.$.{campo}": f"Verificado por {nome} em {data}"
                }
            }
        )
    else:
        col_projetos.update_one(
            {
                "codigo": projeto_codigo,
                "relatorios.numero": relatorio_numero
            },
            {
                "$unset": {
                    f"relatorios.$.{campo}": ""
                }
            }
        )





def todos_relatos_aceitos(projeto, relatorio_numero):
    """
    Retorna True se TODOS os relatos do relatório informado
    estiverem com status_relato == 'aceito'.

    Se existir ao menos um relato do relatório que não seja aceito,
    retorna False.

    Se não existir nenhum relato nesse relatório, retorna False.
    """

    relatos_encontrados = []

    componentes = projeto.get("plano_trabalho", {}).get("componentes", [])

    for componente in componentes:
        for entrega in componente.get("entregas", []):
            for atividade in entrega.get("atividades", []):
                for relato in atividade.get("relatos", []):
                    if relato.get("relatorio_numero") == relatorio_numero:
                        relatos_encontrados.append(relato)

    # Se não existe nenhum relato nesse relatório, não aprova
    if not relatos_encontrados:
        return False

    # Todos precisam estar aceitos
    return all(r.get("status_relato") == "aceito" for r in relatos_encontrados)



def todas_despesas_aceitas(projeto, relatorio_numero):
    """
    Retorna True se TODOS os lançamentos de despesas do relatório
    estiverem com status_despesa == 'aceito'.

    Se existir ao menos uma despesa não aceita, retorna False.
    Se não existir nenhuma despesa nesse relatório, retorna False.
    """

    lancamentos = []

    orcamento = projeto.get("financeiro", {}).get("orcamento", [])

    for item in orcamento:
        for lanc in item.get("lancamentos", []):
            if lanc.get("relatorio_numero") == relatorio_numero:
                lancamentos.append(lanc)

    if not lancamentos:
        return False

    return all(l.get("status_despesa") == "aceito" for l in lancamentos)




def gerar_email_relatorio_aprovado(
    nome_do_contato: str,
    relatorio_numero: int,
    projeto: dict,
    organizacao: str,
    logo_url: str
):

    return f"""
<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <style>
        body {{
            font-family: Arial, Helvetica, sans-serif;
            background-color: #f5f5f5;
            margin: 0;
            padding: 0;
        }}
        .container {{
            max-width: 600px;
            margin: 0 auto;
            background: white;
            border-top: 6px solid #A0C256;
            padding: 30px;
        }}
        .logo {{
            text-align: center;
            margin-bottom: 30px;
        }}
        .content {{
            color: #333;
            font-size: 15px;
            line-height: 1.6;
        }}
        .footer {{
            margin-top: 40px;
            font-size: 12px;
            color: #777;
            text-align: center;
        }}
        .highlight {{
            color: #A0C256;
            font-weight: bold;
        }}
    </style>
</head>
<body>

    <div class="container">

        <div class="logo">
            <img src="{logo_url}" height="70" alt="IEB">
        </div>

        <div class="content">

            <p>Olá <strong>{nome_do_contato}</strong>,</p>

            <p>
                Informamos que o <span class="highlight">Relatório {relatorio_numero}</span>
                do projeto <span class="highlight">{projeto['nome_do_projeto']}</span>
                da organização <strong>{organizacao}</strong> foi <strong>aprovado</strong>.
            </p>

            <p>
                O relatório já está validado no sistema e segue para os próximos encaminhamentos.
            </p>

            <p>
                Atenciosamente,<br>
                <strong>Sistema Veredas</strong>
            </p>
        </div>

        <div class="footer">
            Este é um e-mail automático. Não responda.
        </div>

    </div>

</body>
</html>
"""




def notificar_padrinhos_relatorio(
    col_pessoas,
    numero_relatorio,
    projeto,
    logo_url
):
    padrinhos = buscar_padrinhos_do_projeto(col_pessoas, projeto["codigo"])

    if not padrinhos:
        return False

    for padrinho in padrinhos:
        html = montar_email_relatorio_envio(
            nome=padrinho["nome_completo"],
            numero_relatorio=numero_relatorio,
            codigo=projeto["codigo"],
            sigla=projeto["sigla"],
            logo_url=logo_url
        )

        enviar_email(
            corpo_html=html,
            destinatarios=[padrinho["e_mail"]],
            assunto=f"CEPF - Relatório {numero_relatorio} recebido - Projeto {projeto['codigo']} - {projeto['sigla']}"
        )

    return True




def montar_email_relatorio_envio(
    nome: str,
    numero_relatorio: int,
    codigo: str,
    sigla: str,
    logo_url: str
):
    return f"""
<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <style>
        body {{
            font-family: Arial, Helvetica, sans-serif;
            background-color: #f5f5f5;
            margin: 0;
            padding: 0;
        }}
        .container {{
            max-width: 600px;
            margin: 0 auto;
            background: white;
            border-top: 6px solid #A0C256;
            padding: 30px;
        }}
        .logo {{
            text-align: center;
            margin-bottom: 30px;
        }}
        .content {{
            color: #333;
            font-size: 15px;
            line-height: 1.6;
        }}
        .footer {{
            margin-top: 40px;
            font-size: 12px;
            color: #777;
            text-align: center;
        }}
        .highlight {{
            color: #A0C256;
            font-weight: bold;
        }}
    </style>
</head>
<body>

    <div class="container">
        <div class="logo">
            <img src="{logo_url}" height="70" alt="CEPF">
        </div>

        <div class="content">
            <br>
            <p>Olá <strong>{nome}</strong>,</p>

            <p>
                O relatório <span class="highlight">{numero_relatorio}</span> do projeto
                <span class="highlight">{codigo} - {sigla}</span> está disponível para análise.
            </p>

            <p>
                Por favor, acesse o sistema para realizar a avaliação.
            </p>

            <p>Atenciosamente,<br>
            <strong>Sistema Veredas</strong></p>
        </div>

        <div class="footer">
            Este é um e-mail automático. Não responda.
        </div>
    </div>

</body>
</html>
"""




def buscar_padrinhos_do_projeto(col_pessoas, codigo_projeto: str):
    """
    Retorna lista de pessoas (dict) que são padrinhos do projeto.
    Regra:
      - tipo_usuario != beneficiario
      - tipo_usuario != visitante
      - projetos contém o código do projeto
    """

    padrinhos = list(
        col_pessoas.find(
            {
                "tipo_usuario": {"$nin": ["beneficiario", "visitante"]},
                "projetos": codigo_projeto,
                "status": "ativo"
            },
            {
                "nome_completo": 1,
                "e_mail": 1
            }
        )
    )

    return padrinhos




def gerar_id_lanc_despesa(projeto):
    """
    Gera id sequencial no formato despesa_001, despesa_002...
    """

    numeros = []

    for despesa in projeto.get("financeiro", {}).get("orcamento", []):
        for lanc in despesa.get("lancamentos", []):
            idd = lanc.get("id_lanc_despesa")
            if idd and idd.startswith("despesa_"):
                try:
                    numeros.append(int(idd.split("_")[1]))
                except:
                    pass

    proximo = max(numeros, default=0) + 1
    return f"despesa_{str(proximo).zfill(3)}"




# Diálogo do lançamento de despesa
@st.dialog("Registrar despesa", width="medium")
def dialog_lanc_financ(relatorio_numero, projeto, col_projetos):

    # ==================================================
    # OPÇÕES DE DESPESA
    # ==================================================
    orcamento = projeto["financeiro"]["orcamento"]

    opcoes = sorted([
        f"{o['categoria']} | {o['nome_despesa']}"
        for o in orcamento
    ], key=lambda x: x.lower())

    escolha = st.selectbox(
        "Categoria / Despesa *",
        options=opcoes
    )

    categoria, nome_despesa = escolha.split(" | ")

    # ==================================================
    # DADOS DO LANÇAMENTO
    # ==================================================

    # Gera id sequencial
    id_despesa = gerar_id_lanc_despesa(projeto)

    data_despesa = date_picker(
        label="Data da despesa",
        format="dd/MM/yyyy",
        locale="pt_BR",
        one_tap=True,
        key="data_despesa"
    )


    # Linha de valores

    col1, col2, col3 = st.columns(3)


    with col1:

        quantidade = st.number_input(
            "Quantidade *",
            min_value=0,
            # value=1
        )

    with col2:

        valor_unitario = st.number_input(
            "Valor unitário (reais) *",
            min_value=0.0,
            format="%.2f"
        )


    with col3:

        valor = st.number_input(
            "Valor total (reais) *",
            min_value=0.0,
            format="%.2f"
        )


    descricao = st.text_area("Descrição da despesa *")

    col1, col2 = st.columns([2, 1])

    fornecedor = col1.text_input("Fornecedor *")
    cpf_cnpj = col2.text_input("CPF / CNPJ *")


    # ==================================================
    # LABEL DINÂMICO DOS ANEXOS - se a despesa for do tipo taxa bancária, então o anexo não é obrigatório, não coloca * no label. 
    # ==================================================

    categoria_lower = categoria.lower()

    # Verifica se é taxa bancária
    is_taxa_bancaria = "taxas bancárias" in categoria_lower

    # Define label dinamicamente
    label_anexos = "Anexos" if is_taxa_bancaria else "Anexos *"

    anexos = st.file_uploader(
        label_anexos,
        accept_multiple_files=True
    )


    # ==================================================
    # AÇÕES
    # ==================================================

    with st.container(horizontal=True):


        if st.button("Salvar", type="primary", icon=":material/save:"):

            # ==================================================
            # VALIDAÇÃO DE CAMPOS OBRIGATÓRIOS
            # ==================================================

            # LISTAS DE ERROS SEPARADAS

            erros_campos = []
            erro_consistencia = None



            # Validação da data
            if not data_despesa:
                erros_campos.append("Data da despesa")


            # ==================================================
            # VALIDAÇÃO DE VALORES
            # ==================================================

            if quantidade <= 0:
                erros_campos.append("Quantidade")

            if valor_unitario <= 0:
                erros_campos.append("Valor unitário (reais)")

            if not valor or valor <= 0:
                erros_campos.append("Valor total (reais)")


            # ==================================================
            # VALIDAÇÃO DE CONSISTÊNCIA (QUANTIDADE x VALOR UNITÁRIO)
            # ==================================================

            # Só valida consistência se os três campos foram preenchidos corretamente
            if quantidade > 0 and valor_unitario > 0 and valor > 0:

                valor_calculado = quantidade * valor_unitario

                # Corrige problema de float
                valor_calculado = round(valor_calculado, 2)
                valor_informado = round(valor, 2)

                if valor_informado != valor_calculado:
                    erro_consistencia = (
                        f"Valor total deve ser igual a Quantidade × Valor unitário."
                    )

            # Validação da descrição
            if not descricao or not descricao.strip():
                erros_campos.append("Descrição da despesa")

            # Validação do fornecedor
            if not fornecedor or not fornecedor.strip():
                erros_campos.append("Fornecedor")


            # ==================================================
            # VALIDAÇÃO DO CPF / CNPJ
            # ==================================================

            if not cpf_cnpj or not cpf_cnpj.strip():
                erros_campos.append("CPF / CNPJ")
            else:
                # Mantém apenas números, ponto, barra e traço
                cpf_cnpj_filtrado = "".join(
                    c for c in cpf_cnpj if c.isdigit() or c in [".", "/", "-"]
                )

                # Extrai apenas números para validação de tamanho
                cpf_cnpj_numeros = "".join(filter(str.isdigit, cpf_cnpj_filtrado))

                # Verifica se tem 11 (CPF) ou 14 (CNPJ) dígitos
                if len(cpf_cnpj_numeros) not in [11, 14]:
                    erros_campos.append("CPF / CNPJ inválido.")





            # ==================================================
            # VALIDAÇÃO DOS ANEXOS (COM EXCEÇÃO)
            # ==================================================

            # Regra:
            # - Se for "Taxas bancárias" → anexo NÃO obrigatório
            # - Caso contrário → obrigatório

            categoria_lower = categoria.lower()

            is_taxa_bancaria = "taxas bancárias" in categoria_lower

            if not is_taxa_bancaria:
                if not anexos or len(anexos) == 0:
                    erros_campos.append("Anexos")



            # ==================================================
            # SE HOUVER ERROS → MOSTRA WARNING E NÃO SALVA
            # ==================================================

            # EXIBE ERROS SEPARADAMENTE

            if erros_campos:
                campos = ", ".join(erros_campos)
                st.warning(f"Preencha os seguintes campos obrigatórios: {campos}")

            if erro_consistencia:
                st.warning(erro_consistencia)

            # Se houver qualquer erro → bloqueia
            if erros_campos or erro_consistencia:
                st.stop()




            # ==================================================
            # CONTINUA FLUXO NORMAL (SALVAR)
            # ==================================================
            with st.spinner("Salvando despesa..."):

                novo_lancamento = {
                    "id_lanc_despesa": id_despesa,
                    "relatorio_numero": relatorio_numero,
                    "data_despesa": data_despesa.strftime("%d/%m/%Y"),
                    "descricao_despesa": descricao,
                    "fornecedor": fornecedor,
                    "cpf_cnpj": cpf_cnpj_filtrado,
                    "quantidade": quantidade,
                    "valor_unitario": valor_unitario,
                    "valor_despesa": valor,
                    "status_despesa": "aberto",
                    "anexos": []
                }



                # ==================================================
                # DRIVE
                # ==================================================
                servico = obter_servico_drive()

                pasta_projeto = obter_pasta_projeto(
                    servico,
                    projeto["codigo"],
                    projeto["sigla"]
                )

                pasta_financeiro = obter_pasta_relatos_financeiros(
                    servico,
                    pasta_projeto
                )

                pasta_lanc = obter_ou_criar_pasta(
                    servico,
                    id_despesa,
                    pasta_financeiro
                )

                for arq in anexos:
                    id_drive = enviar_arquivo_drive(servico, pasta_lanc, arq)
                    novo_lancamento["anexos"].append({
                        "nome_arquivo": arq.name,
                        "id_arquivo": id_drive
                    })

                # ==================================================
                # SALVA NO OBJETO
                # ==================================================
                for d in projeto["financeiro"]["orcamento"]:
                    if d["categoria"] == categoria and d["nome_despesa"] == nome_despesa:
                        d.setdefault("lancamentos", []).append(novo_lancamento)
                        break

                # ==================================================
                # SALVA NO MONGO
                # ==================================================
                col_projetos.update_one(
                    {"codigo": projeto["codigo"]},
                    {
                        "$set": {
                            "financeiro.orcamento": projeto["financeiro"]["orcamento"]
                        }
                    }
                )

            st.success("Despesa registrada com sucesso!", icon=":material/check:")
            time.sleep(3)
            st.rerun()

        if st.button("Cancelar"):
            st.rerun()






# ==========================================================
# LOCALIZA UMA ATIVIDADE NO DOCUMENTO DO PROJETO
# ==========================================================
def obter_atividade_mongo(projeto, id_atividade):
    """
    Percorre plano_trabalho → componentes → entregas → atividades
    e retorna a atividade correspondente ao id informado.
    """

    componentes = projeto.get("plano_trabalho", {}).get("componentes", [])

    for componente in componentes:
        for entrega in componente.get("entregas", []):
            for atividade in entrega.get("atividades", []):
                if atividade.get("id") == id_atividade:
                    return atividade

    return None


# ==========================================================
# LISTA OS RELATOS DE UMA ATIVIDADE (UI)
# ==========================================================
def listar_relatos_atividade(atividade, relatorio_numero):
    """
    Lista os relatos cadastrados para a atividade,
    filtrando pelo relatório atual.
    """

    relatos = atividade.get("relatos", [])

    relatos = [
        r for r in relatos
        if r.get("relatorio_numero") == relatorio_numero
    ]

    if not relatos:
        st.info("Nenhum relato cadastrado para esta atividade neste relatório.")
        return

    for relato in relatos:
        with st.expander(
            f"{relato.get('id_relato')} — {relato.get('quando')}"
        ):
            st.write(f"Relato: {relato.get('relato')}")
            st.write(f"Onde: {relato.get('onde', '—')}")
            st.write(f"Autor: {relato.get('autor', '—')}")

            if relato.get("anexos"):
                st.write("Anexos:")
                for a in relato["anexos"]:
                    st.write(f"- {a['nome_arquivo']}")

            if relato.get("fotos"):
                st.write("Fotografias:")
                for f in relato["fotos"]:
                    st.write(
                        f"- {f.get('nome_arquivo')} | "
                        f"{f.get('descricao', '')} | "
                        f"{f.get('fotografo', '')}"
                    )





# Função para salvar o relato
def salvar_relato():
    """
    Salva um relato de atividade:
    - valida campos obrigatórios
    - cria pastas no Google Drive (Relatos_atividades/relato_xxx)
    - envia anexos e fotos
    - grava no MongoDB
    - limpa o session_state
    - executa rerun ao final
    """

    # --------------------------------------------------
    # 1. CAMPOS DO FORMULÁRIO
    # --------------------------------------------------
    texto_relato = st.session_state.get("campo_relato", "")
    data_inicio = st.session_state.get("campo_data_inicio")
    data_fim = st.session_state.get("campo_data_fim")
    anexos = st.session_state.get("campo_anexos", [])
    fotos = st.session_state.get("fotos_relato", [])
    porcentagem_atividade = st.session_state.get("campo_porcentagem_atividade", 0)

    # --------------------------------------------------
    # 2. VALIDAÇÕES
    # --------------------------------------------------
    erros = []
    if not texto_relato.strip():
        erros.append("O campo Relato é obrigatório.")
    # --------------------------------------------------
    # VALIDAÇÃO DAS DATAS
    # --------------------------------------------------

    if not data_inicio:
        erros.append("O campo Data de início é obrigatório.")

    if not data_fim:
        erros.append("O campo Data de fim é obrigatório.")

    if erros:
        for e in erros:
            st.error(e)
        return

    # --------------------------------------------------
    # 3. CONEXÃO COM GOOGLE DRIVE
    # --------------------------------------------------
    servico = obter_servico_drive()

    projeto = st.session_state.get("projeto_mongo")
    if not projeto:
        st.error("Projeto não encontrado na sessão.")
        return

    codigo = projeto["codigo"]
    sigla = projeto["sigla"]

    # Pasta do projeto (padrão já usado em Locais)
    pasta_projeto_id = obter_pasta_projeto(
        servico,
        codigo,
        sigla
    )

    # Pasta Relatos_atividades
    pasta_relatos_id = obter_ou_criar_pasta(
        servico,
        "Relatos_atividades",
        pasta_projeto_id
    )

    # --------------------------------------------------
    # 4. ATIVIDADE SELECIONADA
    # --------------------------------------------------
    atividade = st.session_state.get("atividade_selecionada_drive")
    if not atividade:
        st.error("Atividade não selecionada.")
        return

    id_atividade = atividade.get("id")

    # --------------------------------------------------
    # 5. LOCALIZA ATIVIDADE NO MONGO
    # --------------------------------------------------
    atividade_mongo = obter_atividade_mongo(projeto, id_atividade)
    if not atividade_mongo:
        st.error("Atividade não encontrada no banco de dados.")
        return

   
    # --------------------------------------------------
    # GERA ID DE RELATO GLOBALMENTE ÚNICO
    # --------------------------------------------------
    maior_numero = 0

    for componente in projeto["plano_trabalho"]["componentes"]:
        for entrega in componente["entregas"]:
            for atividade in entrega["atividades"]:
                for relato in atividade.get("relatos", []):
                    id_existente = relato.get("id_relato", "")
                    if id_existente.startswith("relato_"):
                        try:
                            numero = int(id_existente.replace("relato_", ""))
                            maior_numero = max(maior_numero, numero)
                        except ValueError:
                            pass

    # Próximo número disponível
    novo_numero = maior_numero + 1
    id_relato = f"relato_{novo_numero:03d}"




    # --------------------------------------------------
    # 6. PASTA DO RELATO (DIRETAMENTE EM Relatos_atividades)
    # --------------------------------------------------
    pasta_relato_id = obter_ou_criar_pasta(
        servico,
        id_relato,
        pasta_relatos_id
    )


    # --------------------------------------------------
    # 7. UPLOAD DE ANEXOS
    # --------------------------------------------------
    lista_anexos = []

    if anexos:
        pasta_anexos_id = obter_ou_criar_pasta(
            servico,
            "anexos",
            pasta_relato_id
        )

        for arq in anexos:
            id_drive = enviar_arquivo_drive(
                servico,
                pasta_anexos_id,
                arq
            )

            if id_drive:
                lista_anexos.append({
                    "nome_arquivo": arq.name,
                    "id_arquivo": id_drive
                })



    # --------------------------------------------------
    # 8. UPLOAD DE FOTOGRAFIAS
    # --------------------------------------------------
    lista_fotos = []

    fotos_validas = [
        f for f in fotos
        if f.get("arquivo") is not None
    ]



    if fotos_validas:

        # --------------------------------------------------
        # CRIA PASTA FOTOS (SE NÃO EXISTIR)
        # --------------------------------------------------
        # Verifica se já existe antes
        consulta = (
            f"name='fotos' and "
            f"'{pasta_relato_id}' in parents and "
            f"mimeType='application/vnd.google-apps.folder' and trashed=false"
        )

        resultado = servico.files().list(
            q=consulta,
            fields="files(id)",
            supportsAllDrives=True,
            includeItemsFromAllDrives=True
        ).execute()

        arquivos = resultado.get("files", [])

        if arquivos:
            pasta_fotos_id = arquivos[0]["id"]
        else:
            # Cria pasta
            pasta_fotos_id = obter_ou_criar_pasta(
                servico,
                "fotos",
                pasta_relato_id
            )


            # DEFINE PERMISSÃO PÚBLICA na pasta de fotos de cada relato
            garantir_permissao_publica_leitura(servico, pasta_fotos_id)


        for foto in fotos_validas:
            arq = foto["arquivo"]

            id_drive = enviar_arquivo_drive(
                servico,
                pasta_fotos_id,
                arq
            )

            if id_drive:
                lista_fotos.append({
                    "nome_arquivo": arq.name,
                    "descricao": foto.get("descricao", ""),
                    "fotografo": foto.get("fotografo", ""),
                    "id_arquivo": id_drive
                })




    # --------------------------------------------------
    # 9. OBJETO FINAL DO RELATO
    # --------------------------------------------------
    

    data_inicio_str = data_inicio.strftime("%d/%m/%Y") if data_inicio else None
    data_fim_str = data_fim.strftime("%d/%m/%Y") if data_fim else None    
    

    novo_relato = {
        "id_relato": id_relato,
        "status_relato": "aberto",
        "relatorio_numero": st.session_state.get("relatorio_numero"),
        "relato": texto_relato.strip(),

        "data_inicio": data_inicio_str,
        "data_fim": data_fim_str,

        "porc_ativ_relato": int(porcentagem_atividade),

        "autor": st.session_state.get("nome", "Usuário")
    }



    if lista_anexos:
        novo_relato["anexos"] = lista_anexos

    if lista_fotos:
        novo_relato["fotos"] = lista_fotos

    atividade_mongo.setdefault("relatos", []).append(novo_relato)

    col_projetos.update_one(
        {"codigo": codigo},
        {
            "$set": {
                "plano_trabalho.componentes": projeto["plano_trabalho"]["componentes"]
            }
        }
    )

    # --------------------------------------------------
    # 10. LIMPEZA DO SESSION_STATE (CRÍTICO)
    # --------------------------------------------------

    for chave in [
        "campo_relato",
        "campo_data_inicio",
        "campo_data_fim",
        "campo_porcentagem_atividade",
        "campo_anexos",
        "fotos_relato"
    ]:

        if chave in st.session_state:
            del st.session_state[chave]

    # Remove chaves dinâmicas das fotos
    for k in list(st.session_state.keys()):
        if k.startswith("foto_"):
            del st.session_state[k]

    # --------------------------------------------------
    # 11. FINALIZAÇÃO
    # --------------------------------------------------
    st.success("Relato salvo com sucesso.", icon=":material/check:")
    time.sleep(3)
    st.rerun()


# Função auxiliar para o salvar_relato, que dá permissão de leitura pública para a pasta de fotos no ato da criação da pasta no drivce
def garantir_permissao_publica_leitura(servico, pasta_id):
    """
    Define permissão:
    Qualquer pessoa com o link → Leitor
    (somente se ainda não existir)
    """

    try:
        servico.permissions().create(
            fileId=pasta_id,
            body={
                "type": "anyone",
                "role": "reader"
            },
            supportsAllDrives=True
        ).execute()
    except Exception as e:
        st.error(f"Erro ao definir permissão: {str(e)}")
        raise




# ==========================================================================================
# DIÁLOGO: RELATAR ATIVIDADE
# ==========================================================================================
@st.dialog("Relatar atividade", width="medium")
def dialog_relatos():

    projeto = st.session_state.get("projeto_mongo")
    if not projeto:
        st.error("Projeto não encontrado.")
        return

    # --------------------------------------------------
    # 1. MONTA LISTA DE ATIVIDADES
    # --------------------------------------------------
    atividades = []

    for componente in projeto["plano_trabalho"]["componentes"]:
        for entrega in componente["entregas"]:
            for atividade in entrega["atividades"]:
                atividades.append({
                    "id": atividade["id"],
                    "atividade": atividade["atividade"],
                    "componente": componente["componente"],
                    "entrega": entrega["entrega"],
                    "data_inicio": atividade.get("data_inicio"),
                    "data_fim": atividade.get("data_fim"),
                    "relatos": atividade.get("relatos", [])
                })

    if not atividades:
        st.info("Nenhuma atividade cadastrada.")
        time.sleep(3)
        return

    # --------------------------------------------------
    # 2. SELECTBOX COM OPÇÃO VAZIA
    # --------------------------------------------------
    atividades_com_placeholder = (
        [{"id": None, "atividade": ""}]
        + atividades
    )

    atividade_selecionada = st.selectbox(
        "Selecione a atividade *",
        atividades_com_placeholder,
        format_func=lambda x: x["atividade"],
        key="atividade_select_dialog"
    )




    # recupera datas da atividade selecionada
    data_inicio_atv = None
    data_fim_atv = None

    if atividade_selecionada and atividade_selecionada.get("id"):

        atividade_mongo = obter_atividade_mongo(
            projeto,
            atividade_selecionada["id"]
        )

        if atividade_mongo:
            data_inicio_atv = atividade_mongo.get("data_inicio")
            data_fim_atv = atividade_mongo.get("data_fim")

    # mostra período programado da atividade
    if data_inicio_atv and data_fim_atv:
        st.write(
            f"Programada para começar em **{data_inicio_atv}** e terminar em **{data_fim_atv}**."
        )

        st.write('')







    # SELECTBOX DE PORCENTAGEM DE EXECUÇÃO DA ATIVIDADE

    # opções de porcentagem
    porcentagens = list(range(0, 101, 10))

    porcentagem_atual = 0

    # busca a porcentagem atual da atividade selecionada
    if atividade_selecionada and atividade_selecionada.get("id"):

        atividade_mongo = obter_atividade_mongo(
            projeto,
            atividade_selecionada["id"]
        )

        if atividade_mongo:
            porcentagem_atual = atividade_mongo.get("porcentagem_atv", 0)

    # garante que esteja dentro das opções
    if porcentagem_atual not in porcentagens:
        porcentagem_atual = 0

    # sincroniza o session_state quando a atividade muda
    if (
        "campo_porcentagem_atividade" not in st.session_state
        or st.session_state.get("atividade_porcentagem_ref") != atividade_selecionada.get("id")
    ):
        st.session_state["campo_porcentagem_atividade"] = porcentagem_atual
        st.session_state["atividade_porcentagem_ref"] = atividade_selecionada.get("id")

    # selectbox de porcentagem
    porcentagem_escolhida = st.selectbox(
        "Atualize a porcentagem de execução da atividade *",
        options=porcentagens,
        format_func=lambda x: f"{x}%",
        key="campo_porcentagem_atividade",
        width=300
    )






    # Salva no session_state (mesmo vazia, para validação)
    st.session_state["atividade_selecionada"] = atividade_selecionada
    st.session_state["atividade_selecionada_drive"] = atividade_selecionada



    

    # ==================================================
    # 3. FORMULÁRIO DO RELATO
    # ==================================================
    @st.fragment
    def corpo_formulario():

        st.divider()

        # -----------------------------
        # CAMPOS BÁSICOS
        # -----------------------------
        st.text_area(
            "Relato *",
            placeholder="Descreva o que foi feito",
            key="campo_relato"
        )


        # --------------------------------------------------
        # CAMPOS DE DATA DO RELATO
        # --------------------------------------------------


        col1, col2 = st.columns(2)

        # ---------- DATA DE INÍCIO ----------
        with col1:
            data_inicio = date_picker(
                label="Data de início",
                format="dd/MM/yyyy",
                locale="pt_BR",
                one_tap=True,
                key="campo_data_inicio"
            )

        # ---------- DATA DE FIM ----------
        with col2:
            data_fim = date_picker(
                label="Data de fim",
                format="dd/MM/yyyy",
                locale="pt_BR",
                one_tap=True,
                key="campo_data_fim"
            )

        # col1.date_input(
        #     "Data de início *",
        #     key="campo_data_inicio",
        #     format="DD/MM/YYYY"
        # )

        # col2.date_input(
        #     "Data de fim *",
        #     key="campo_data_fim",
        #     format="DD/MM/YYYY"
        # )


        st.divider()

        # -----------------------------
        # ANEXOS
        # -----------------------------
        st.write("**Anexos**")
        st.write("Adicione aqui todos os anexos relevantes para esse relato: listas de presença, relatórios, publicações, etc.")
        st.write("Você pode adicionar vários arquivos de uma só vez. **Não inclua fotos aqui**.")
        st.write("")
        
        st.file_uploader(
            "Selecione um ou vários arquivos.",
            type=["pdf", "docx", "xlsx", "csv", "jpg", "jpeg", "png"],
            accept_multiple_files=True,
            key="campo_anexos"
        )

        st.divider()


        # -----------------------------
        # FOTOGRAFIAS
        # -----------------------------
        st.write("**Fotografias**")

        st.write("Selecione aqui todas as **fotografias** relevantes para esse relato.")
        st.write("Você pode adicionar várias fotografias, mas uma de cada vez. Clique no botão **'Adicionar fotografia'** sempre que quiser adicionar mais uma.")


        if "fotos_relato" not in st.session_state:
            st.session_state["fotos_relato"] = []

        st.write('')
        # Botão para adicionar
        if st.button("Adicionar fotografia", icon=":material/add_a_photo:"):
            # Usamos um ID único para cada foto em vez de apenas o índice

            st.session_state["fotos_relato"].append({
                "id": str(uuid.uuid4()), 
                "arquivo": None,
                "descricao": "",
                "fotografo": ""
            })
            st.rerun(scope="fragment") # Atualiza APENAS o fragmento

        # Iteramos sobre uma cópia da lista para evitar erros de índice ao deletar
        for i, foto in enumerate(st.session_state["fotos_relato"]):
            # Criamos uma chave única baseada no ID gerado, não apenas no índice i
            # Isso evita que o Streamlit confunda os campos após uma remoção
            foto_id = foto["id"]
            
            with st.container(border=True):
                col_info, col_delete = st.columns([8, 2])
                col_info.write(f"Fotografia {i+1}")
                

                with col_delete.container(horizontal=True, horizontal_alignment="right"):

                    if st.button("", 
                                        key=f"btn_del_{foto_id}", 
                                        help="Remover foto", 
                                        icon=":material/close:",
                                        type="tertiary"):
                        
                        st.session_state["fotos_relato"].pop(i)
                        st.rerun(scope="fragment") # O "pulo do gato": atualiza só o fragmento

                arquivo_foto = st.file_uploader(
                    "Selecione a foto",
                    type=["jpg", "jpeg", "png"],
                    key=f"file_{foto_id}"
                )

                descricao = st.text_input(
                    "Descrição da foto",
                    key=f"desc_{foto_id}"
                )

                fotografo = st.text_input(
                    "Nome do(a) fotógrafo(a)",
                    key=f"autor_{foto_id}"
                )

            # Sincronização
            foto["arquivo"] = arquivo_foto
            foto["descricao"] = descricao
            foto["fotografo"] = fotografo




        # --------------------------------------------------
        # AÇÕES FINAIS: BOTÕES + VALIDAÇÃO + SPINNER
        # --------------------------------------------------
        st.write('')
        with st.container(horizontal=True):

            # Botão salvar
            salvar = st.button(
                "Salvar relato",
                type="primary",
                icon=":material/save:"
            )

            # Botão cancelar
            cancelar = st.button("Cancelar")

        if salvar:

            erros = []

            # Valida atividade
            if not atividade_selecionada.get("id"):
                erros.append("Selecione uma atividade.")

            # Valida campos obrigatórios
            if not st.session_state.get("campo_relato", "").strip():
                erros.append("O campo Relato é obrigatório.")

            # VALIDAÇÃO DAS DATAS
            # Verifica se ambas as datas foram informadas.

            data_inicio = st.session_state.get("campo_data_inicio")
            data_fim = st.session_state.get("campo_data_fim")

            if not data_inicio:
                erros.append("O campo Data de início é obrigatório.")

            if not data_fim:
                erros.append("O campo Data de fim é obrigatório.")


            # Mostra erros (mesma funcionalidade de antes)
            if erros:
                for e in erros:
                    st.error(e)
                return

            # Se passou na validação, salva
            with st.spinner("Salvando, aguarde..."):
                salvar_relato()

            st.success("Relato salvo com sucesso.")
            st.rerun()

        # Cancelar apenas faz rerun
        if cancelar:
            st.rerun()

    corpo_formulario()






# Função para liberar o próximo relatório quando o relatório anterior for aprovado
def liberar_proximo_relatorio(projeto_codigo, relatorios):
    """
    Se um relatório estiver aprovado, libera o próximo
    caso ele esteja como 'aguardando'.
    """
    for i in range(len(relatorios) - 1):
        status_atual = relatorios[i].get("status_relatorio")
        status_proximo = relatorios[i + 1].get("status_relatorio")

        if status_atual == "aprovado" and status_proximo == "aguardando":
            col_projetos.update_one(
                {
                    "codigo": projeto_codigo,
                    "relatorios.numero": relatorios[i + 1]["numero"]
                },
                {
                    "$set": {
                        "relatorios.$.status_relatorio": "modo_edicao"
                    }
                }
            )




# Renderiza as perguntas em modo visualização
def renderizar_visualizacao(pergunta, resposta):
    """
    Renderiza pergunta em negrito e resposta em texto normal
    """
    st.markdown(f"**{pergunta}**")
    if resposta in [None, "", [], {}]:
        st.write("—")
    else:
        st.write(resposta)
    st.write("")



# Atualiza o status do relatório no banco de dados, apoiando o segmented_control

STATUS_UI_TO_DB = {
    "Modo edição": "modo_edicao",
    "Em análise": "em_analise",
    "Aprovado": "aprovado",
}

STATUS_DB_TO_UI = {v: k for k, v in STATUS_UI_TO_DB.items()}




def atualizar_status_relatorio(idx, relatorio_numero, projeto_codigo):
    """
    Atualiza o status do relatório no MongoDB quando o segmented_control muda.

    Regras de sincronização dos relatos:

    A) Se o relatório voltar de 'em_analise' ou 'aprovado' para 'modo_edicao':
       - relatos deste relatório com status 'em_analise' voltam para 'aberto'

    B) Se o relatório sair de 'modo_edicao' para 'em_analise' ou 'aprovado':
       - relatos deste relatório com status 'aberto' passam para 'em_analise'
    """

    # --------------------------------------------------
    # 1. STATUS SELECIONADO NA UI
    # --------------------------------------------------
    status_ui = st.session_state.get(f"status_relatorio_{idx}")
    status_novo = STATUS_UI_TO_DB.get(status_ui)

    if not status_novo:
        return  # segurança extra

    # --------------------------------------------------
    # 2. BUSCA STATUS ATUAL DO RELATÓRIO NO BANCO
    # --------------------------------------------------
    projeto = col_projetos.find_one(
        {
            "codigo": projeto_codigo,
            "relatorios.numero": relatorio_numero
        },
        {
            "relatorios.$": 1
        }
    )

    if not projeto or "relatorios" not in projeto:
        return

    relatorio = projeto["relatorios"][0]
    status_anterior = relatorio.get("status_relatorio")

    # --------------------------------------------------
    # 3. ATUALIZA STATUS DO RELATÓRIO
    # --------------------------------------------------
    col_projetos.update_one(
        {
            "codigo": projeto_codigo,
            "relatorios.numero": relatorio_numero
        },
        {
            "$set": {
                "relatorios.$.status_relatorio": status_novo
            }
        }
    )

    # --------------------------------------------------
    # 4. VERIFICA SE ALGUMA REGRA DE RELATOS SE APLICA
    # --------------------------------------------------
    aplica_regra_a = (
        status_novo == "modo_edicao"
        and status_anterior in ["em_analise", "aprovado"]
    )

    aplica_regra_b = (
        status_anterior == "modo_edicao"
        and status_novo in ["em_analise", "aprovado"]
    )

    if not (aplica_regra_a or aplica_regra_b):
        return  # nada a fazer nos relatos

    # --------------------------------------------------
    # 5. RECARREGA O PROJETO COMPLETO
    # --------------------------------------------------
    projeto_atualizado = col_projetos.find_one(
        {"codigo": projeto_codigo}
    )

    componentes = projeto_atualizado["plano_trabalho"]["componentes"]
    houve_alteracao = False

    # --------------------------------------------------
    # 6. APLICA AS REGRAS NOS RELATOS
    # --------------------------------------------------
    # ------------------------------------------------------------------
    # Percorre todos os componentes
    # ------------------------------------------------------------------
    for componente in componentes:

        # Recupera entregas de forma segura
        entregas = componente.get("entregas", [])

        # Se não houver entregas, continua
        if not entregas:
            continue

        # ------------------------------------------------------------------
        # Percorre entregas
        # ------------------------------------------------------------------
        for entrega in entregas:

            # Recupera atividades de forma segura
            atividades = entrega.get("atividades", [])

            # --------------------------------------------------------------
            # Se a entrega não tiver atividades, apenas ignora
            # --------------------------------------------------------------
            if not atividades:
                continue

            # --------------------------------------------------------------
            # Percorre atividades
            # --------------------------------------------------------------
            for atividade in atividades:

                # Recupera relatos de forma segura
                relatos = atividade.get("relatos", [])

                # Se não houver relatos, continua
                if not relatos:
                    continue

                # ----------------------------------------------------------
                # Percorre relatos
                # ----------------------------------------------------------
                for relato in relatos:

                    # Apenas relatos do relatório atual
                    if relato.get("relatorio_numero") != relatorio_numero:
                        continue

                    # Regra A: em_analise -> aberto
                    if aplica_regra_a and relato.get("status_relato") == "em_analise":
                        relato["status_relato"] = "aberto"
                        houve_alteracao = True

                    # Regra B: aberto -> em_analise
                    if aplica_regra_b and relato.get("status_relato") == "aberto":
                        relato["status_relato"] = "em_analise"
                        houve_alteracao = True


    # --------------------------------------------------
    # 7. SALVA NO BANCO APENAS SE HOUVE ALTERAÇÃO
    # --------------------------------------------------
    if houve_alteracao:
        col_projetos.update_one(
            {"codigo": projeto_codigo},
            {
                "$set": {
                    "plano_trabalho.componentes": componentes
                }
            }
        )







def extrair_atividades(projeto):
    atividades = []

    plano = projeto.get("plano_trabalho", {})
    componentes = plano.get("componentes", [])

    for componente in componentes:
        for entrega in componente.get("entregas", []):
            for atividade in entrega.get("atividades", []):
                atividades.append({
                    "id": atividade.get("id"),
                    "nome": atividade.get("atividade"),
                    "data_inicio": atividade.get("data_inicio"),
                    "data_fim": atividade.get("data_fim"),
                    "componente": componente.get("componente"),
                    "entrega": entrega.get("entrega"),
                })

    return atividades



# Função para formatar números no padrão brasileiro, com poucas casas decimais (dinamicamente)
def formatar_numero_br_dinamico(valor):
    """
    Formata número no padrão brasileiro:
    - Sem decimais → não mostra casas
    - 1 decimal → mostra 1 casa
    - 2+ decimais → mostra até 2 casas (sem zeros desnecessários)
    """
    try:
        valor = float(valor)
    except (TypeError, ValueError):
        return "—"

    # Verifica parte decimal
    inteiro = int(valor)
    decimal = abs(valor - inteiro)

    # Define casas decimais dinamicamente
    if decimal == 0:
        casas = 0
    elif round(decimal * 10) == decimal * 10:
        casas = 1
    else:
        casas = 2

    texto = f"{valor:,.{casas}f}"

    # Converte para padrão pt-BR
    return texto.replace(",", "X").replace(".", ",").replace("X", ".")

# Função para formatar números no padrão brasileiro na aba de Indicadores de Resultados
def parse_numero_br(valor_str):
    """
    Converte string no formato brasileiro para float.
    Ex:
    - '50,15' -> 50.15
    - '1.234,56' -> 1234.56
    """
    if valor_str is None:
        return None

    valor_str = valor_str.strip()

    if not valor_str:
        return None

    try:
        return float(
            valor_str.replace(".", "").replace(",", ".")
        )
    except ValueError:
        return None


def data_hoje_br():
    return datetime.datetime.now().strftime("%d/%m/%Y")




###########################################################################################################
# TRATAMENTO DOS DADOS E CONTROLES DE SESSÃO
###########################################################################################################


# Libera automaticamente o próximo relatório, se aplicável
liberar_proximo_relatorio(projeto["codigo"], relatorios)

# Recarrega o projeto para refletir possíveis mudanças
projeto = col_projetos.find_one({"codigo": projeto["codigo"]})
relatorios = projeto.get("relatorios", [])




# -------------------------------------------
# CONTROLE DE STEP DO RELATÓRIO
# -------------------------------------------

if "step_relatorio" not in st.session_state:
    st.session_state.step_relatorio = "Atividades"




###########################################################################################################
# INTERFACE PRINCIPAL DA PÁGINA
###########################################################################################################

# Logo hospedada no site do IEB para renderizar nos e-mails.
logo_cepf = "https://cepfcerrado.iieb.org.br/wp-content/uploads/2025/02/LogoConjuntaCEPFIEBGREEN-768x140.png"


# Logo do sidebar
st.logo("images/ieb_logo.svg", size='large')

# Título da página e identificação
col_titulo, col_identificacao = st.columns([3, 2])

with col_titulo:
    st.header("Relatórios")

with col_identificacao:
    st.markdown(
        f"<div style='text-align: right; margin-top: 30px;'>{df_projeto['codigo'].values[0]} - {df_projeto['sigla'].values[0]}</div>",
        unsafe_allow_html=True
    )


st.write('')
st.write('')







###########################################################################################################
# CONFIGURAÇÃO DOS STEPS DO RELATÓRIO
###########################################################################################################

if tipo_usuario in ["admin", "equipe"]:
    steps_relatorio = [
        "Atividades",
        "Despesas",
        "Resultados",
        "Beneficiários",
        "Pesquisas",
        "Formulário",
        "Avaliação"
    ]
else:
    steps_relatorio = [
        "Atividades",
        "Despesas",
        "Resultados",
        "Beneficiários",
        "Pesquisas",
        "Formulário",
        "Enviar"
    ]


###########################################################################################################
# VERIFICA SE EXISTEM RELATÓRIOS
###########################################################################################################

if not relatorios:
    st.warning("Este projeto ainda não possui relatórios cadastrados.")
    st.stop()

###########################################################################################################
# ABAS DOS RELATÓRIOS (sac.tabs)
###########################################################################################################

labels_relatorios = [f"Relatório {r.get('numero')}" for r in relatorios]

aba_selecionada = sac.tabs(
    items=[sac.TabsItem(label=l) for l in labels_relatorios],
    align="left",
    variant="outline",
    key="tabs_relatorios"
    # size="xl"
)

idx = labels_relatorios.index(aba_selecionada)
relatorio = relatorios[idx]

###########################################################################################################
# DADOS DO RELATÓRIO
###########################################################################################################

relatorio_numero = relatorio["numero"]
projeto_codigo = projeto["codigo"]


###########################################################################################################
# LINHA COM TÍTULO E BOTÃO DE OPÇÕES
###########################################################################################################

with st.container(horizontal=True):

    st.subheader(f"Relatório {relatorio_numero}")

    # Popover de exportação do relatório
    with st.popover(
        "Exportar", 
        icon=":material/print:", 
        type="tertiary"):

        # Fragment para isolar renderização
        @st.fragment
        def fragment_exportacao_relatorio():

            # Inicializa estado
            if "docx_relatorio" not in st.session_state:
                st.session_state.docx_relatorio = None

            if "relatorio_gerado" not in st.session_state:
                st.session_state.relatorio_gerado = False


            # BOTÃO GERAR RELATÓRIO ----------------------------------------
            if st.button(
                "Gerar relatório",
                icon=":material/settings:",
                type="secondary",
                width="stretch"
            ):

                # Gera o documento em memória
                buffer = BytesIO()

                doc = gerar_docx_relatorio(relatorio, projeto)

                doc.save(buffer)
                buffer.seek(0)

                # Armazena no session_state
                st.session_state.docx_relatorio = buffer
                st.session_state.relatorio_gerado = True


            # BOTÃO DOWNLOAD -----------------------------------------------
            if st.session_state.relatorio_gerado:

                st.caption("Relatório gerado! Clique para baixar.")

                numero = relatorio.get("numero")
                codigo = projeto.get("codigo", "").replace("/", "-")

                # Busca organização para sigla
                organizacao = db["organizacoes"].find_one(
                    {"_id": projeto.get("id_organizacao")}
                )
                sigla_org = organizacao.get("sigla_organizacao") if organizacao else "org"

                nome_arquivo = f"relatorio_{numero}_{codigo}_{sigla_org}.docx"

                st.download_button(
                    label="Baixar relatório",
                    data=st.session_state.docx_relatorio,
                    file_name=nome_arquivo,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    icon=":material/download:",
                    type="primary",
                    width="stretch"
                )


        # Executa fragment
        fragment_exportacao_relatorio()



###########################################################################################################
# STATUS ATUAL DO RELATÓRIO
###########################################################################################################

status_atual_db = relatorio.get("status_relatorio", "modo_edicao")
status_atual_ui = STATUS_DB_TO_UI.get(status_atual_db, "Modo edição")

aguardando = False

###########################################################################################################
# CONTROLE CENTRAL DE PERMISSÃO DE EDIÇÃO
###########################################################################################################

pode_editar_relatorio = (
    status_atual_db == "modo_edicao"
    and tipo_usuario == "beneficiario"
)

###########################################################################################################
# CONTROLE DE ESTADO – NOVA COMUNIDADE
###########################################################################################################

if f"mostrar_nova_comunidade_{idx}" not in st.session_state:
    st.session_state[f"mostrar_nova_comunidade_{idx}"] = False

###########################################################################################################
# REGRA DE BLOQUEIO (a partir do 2º relatório)
###########################################################################################################

if idx > 0:
    status_anterior = relatorios[idx - 1].get("status_relatorio")

    if status_anterior != "aprovado":
        aguardando = True

        col_projetos.update_one(
            {
                "codigo": projeto_codigo,
                "relatorios.numero": relatorio_numero,
                "relatorios.status_relatorio": {"$ne": "aguardando"}
            },
            {
                "$set": {
                    "relatorios.$.status_relatorio": "aguardando"
                }
            }
        )

        status_atual_ui = "Modo edição"

###########################################################################################################
# MENSAGEM DE STATUS DO RELATÓRIO PARA BENEFICIÁRIO E VISITANTE
###########################################################################################################

if tipo_usuario in ["beneficiario", "visitante"]:

    if status_atual_db == "em_analise":
        st.write("")
        st.warning("Relatório em análise. Aguarde o retorno.", icon=":material/manage_search:")

    elif status_atual_db == "aprovado":
        st.write("")
        st.success("Relatório aprovado", icon=":material/check:")

###########################################################################################################
# SINCRONIZA STATUS DO RELATÓRIO COM A UI
###########################################################################################################

status_key = f"status_relatorio_{idx}"
status_atual_ui = STATUS_DB_TO_UI.get(status_atual_db, "Modo edição")

if st.session_state.get(status_key) != status_atual_ui:
    st.session_state[status_key] = status_atual_ui

###########################################################################################################
# SEGMENTED CONTROL (somente equipe interna)
###########################################################################################################

if tipo_usuario in ["equipe", "admin"]:
    with st.container(horizontal=True, horizontal_alignment="center"):
        st.segmented_control(
            label="",
            options=["Modo edição", "Em análise", "Aprovado"],
            key=f"status_relatorio_{idx}",
            disabled=aguardando,
            on_change=atualizar_status_relatorio if not aguardando else None,
            args=(idx, relatorio_numero, projeto_codigo) if not aguardando else None
        )




###########################################################################################################
# MENSAGEM DE AGUARDO + STEPS DO RELATÓRIO
###########################################################################################################

labels_steps = steps_relatorio

if aguardando:
    st.write("")
    st.info(
        "Aguardando a aprovação do relatório anterior.",
        icon=":material/nest_clock_farsight_analog:"
    )
    step_selecionado = None  # importante para evitar uso depois
else:
    st.write("")
    st.write("")

    step_selecionado = sac.tabs(
        items=[sac.TabsItem(label=s) for s in labels_steps],
        align="start",
        use_container_width=True,
        key=f"steps_relatorio_{idx}"
        # size="md"
    )





###########################################################################################################
# CONTEÚDO DOS STEPS
###########################################################################################################










# ---------- ATIVIDADES ----------

if step_selecionado == "Atividades":

    # Guarda para uso no diálogo e no salvar_relato
    st.session_state["projeto_mongo"] = projeto
    st.session_state["relatorio_numero"] = relatorio_numero

    st.write("")
    st.write("")

    st.markdown("#### Relatos de atividades")
    st.write('')

    # --------------------------------------------------
    # BOTÃO PARA CRIAR NOVO RELATO
    # --------------------------------------------------
    with st.container(horizontal=True, horizontal_alignment="right"):

        if pode_editar_relatorio:
            if st.button(
                "Relatar atividade",
                type="primary",
                key=f"btn_relatar_{idx}",
                icon=":material/add:",
                width=260
            ):
                # Limpa qualquer resíduo antigo de formulário
                for chave in [
                    "campo_relato",
                    "campo_quando",
                    "campo_onde",
                    "campo_anexos",
                    "fotos_relato",
                    "atividade_select_dialog"
                ]:
                    if chave in st.session_state:
                        del st.session_state[chave]

                dialog_relatos()


    # --------------------------------------------------
    # LISTAGEM DE TODOS OS RELATOS DO RELATÓRIO
    # AGRUPADOS POR ATIVIDADE
    # --------------------------------------------------

    if "relato_editando_id" not in st.session_state:
        st.session_state["relato_editando_id"] = None

    tem_relato = False

    # ------------------------------------------------------------------
    # Percorre componentes do plano de trabalho
    # ------------------------------------------------------------------
    for componente in projeto["plano_trabalho"]["componentes"]:

        # Recupera entregas do componente de forma segura
        entregas = componente.get("entregas", [])

        # Caso o componente não tenha entregas, apenas continua
        if not entregas:
            continue

        # ------------------------------------------------------------------
        # Percorre entregas
        # ------------------------------------------------------------------
        for entrega in entregas:

            # Recupera atividades de forma segura
            atividades = entrega.get("atividades", [])

            # --------------------------------------------------------------
            # Caso não existam atividades, mostra aviso e continua
            # --------------------------------------------------------------
            if not atividades:
                continue

            # --------------------------------------------------------------
            # Percorre atividades normalmente
            # --------------------------------------------------------------
            for atividade in atividades:

                # ----------------------------------------------------------
                # Filtra relatos do relatório atual
                # ----------------------------------------------------------
                relatos = [
                    r for r in atividade.get("relatos", [])
                    if r.get("relatorio_numero") == relatorio_numero
                ]

                # Se não há relatos para essa atividade, pula
                if not relatos:
                    continue

                tem_relato = True

                st.write("")
                st.markdown(f"#### {atividade['atividade']}")




                for relato in relatos:

                    id_relato = relato["id_relato"]
                    editando = st.session_state["relato_editando_id"] == id_relato

                    # --------------------------------------------------
                    # GARANTE QUE WIDGETS DE VISUALIZAÇÃO NÃO EXISTAM EM EDIÇÃO
                    # --------------------------------------------------
                    if editando:
                        # remove qualquer state de devolutiva para evitar conflito
                        st.session_state.pop(f"devolutiva_{id_relato}", None)
                        st.session_state.pop(f"status_relato_ui_{id_relato}", None)


                    with st.container(border=True):

                        # ==================================================
                        # MODO VISUALIZAÇÃO DO RELATO
                        # ==================================================
                        if not editando:

                            # --------------------------------------------------
                            # Lógica de status visual (depende de devolutiva)
                            # --------------------------------------------------
                            status_relato_db = relato.get("status_relato", "em_analise")
                            tem_devolutiva = bool(relato.get("devolutiva"))

                            # Regras visuais:
                            # - aberto + devolutiva → Pendente (vermelho)
                            # - aberto sem devolutiva → Aberto (amarelo)
                            # - em_analise → Em análise (azul)
                            # - aceito → Aceito (verde)

                            if status_relato_db == "aberto" and tem_devolutiva:
                                badge = {
                                    "label": "Pendente",
                                    "bg": "#F8D7DA",
                                    "color": "#721C24"
                                }
                            elif status_relato_db == "aberto":
                                badge = {
                                    "label": "Aberto",
                                    "bg": "#FFF3CD",
                                    "color": "#856404"
                                }
                            elif status_relato_db == "aceito":
                                badge = {
                                    "label": "Aceito",
                                    "bg": "#D4EDDA",
                                    "color": "#155724"
                                }
                            else:
                                badge = {
                                    "label": "Em análise",
                                    "bg": "#D1ECF1",
                                    "color": "#0C5460"
                                }

                            # --------------------------------------------------
                            # BADGE VISUAL
                            # --------------------------------------------------
                            
                            col1, col2 = st.columns([9, 1])
                            
                            col2.markdown(
                                f"""
                                <div style="margin-bottom:6px;">
                                    <span style="
                                        background:{badge['bg']};
                                        color:{badge['color']};
                                        padding:4px 10px;
                                        border-radius:20px;
                                        font-size:12px;
                                        font-weight:600;
                                    ">
                                        {badge['label']}
                                    </span>
                                </div>
                                """,
                                unsafe_allow_html=True
                                )


                            # --------------------------------------------------
                            # CONTEÚDO DO RELATO
                            # --------------------------------------------------
                            st.write(f"**{id_relato.upper()}:** {relato.get("relato")}")

                            col1, col2 = st.columns([2, 3])

                            col1.write(f"**Data de início:** {relato.get('data_inicio')}")
                            col2.write(f"**Data de fim:** {relato.get('data_fim')}")

                            if relato.get("porc_ativ_relato") is not None:
                                st.write(f"**Progresso da atividade informado:** {relato.get('porc_ativ_relato')}%")

                            # col1.write(f"**Quando:** {relato.get('quando')}")
                            # col2.write(f"**Onde:** {relato.get('onde')}")

                            # --------------------------------------------------
                            # ANEXOS (links do Drive)
                            # --------------------------------------------------
                            if relato.get("anexos"):
                                with col1:
                                    c1, c2 = st.columns([1, 5])
                                    c1.write("**Anexos:**")
                                    for a in relato["anexos"]:
                                        if a.get("id_arquivo"):
                                            link = gerar_link_drive(a["id_arquivo"])
                                            c2.markdown(
                                                f"[{a['nome_arquivo']}]({link})",
                                                unsafe_allow_html=True
                                            )

                            # --------------------------------------------------
                            # FOTOGRAFIAS (links + metadados)
                            # --------------------------------------------------
                            if relato.get("fotos"):
                                with col2:
                                    c1, c2 = st.columns([1, 5])
                                    c1.write("**Fotografias:**")
                                    for f in relato["fotos"]:
                                        if f.get("id_arquivo"):
                                            link = gerar_link_drive(f["id_arquivo"])
                                            linha = f"[{f['nome_arquivo']}]({link})"
                                            if f.get("descricao"):
                                                linha += f" | {f['descricao']}"
                                            if f.get("fotografo"):
                                                linha += f" | {f['fotografo']}"
                                            c2.markdown(linha, unsafe_allow_html=True)









                            # ==========================
                            # STATUS DO RELATO (ADMIN/EQUIPE)
                            # ==========================

                            STATUS_RELATO_LABEL = {
                                "em_analise": "Em análise",
                                "aberto": "Devolver",
                                "aceito": "Aceito"
                            }

                            STATUS_RELATO_LABEL_INV = {v: k for k, v in STATUS_RELATO_LABEL.items()}

                            usuario_admin = tipo_usuario == "admin"
                            usuario_equipe = tipo_usuario == "equipe"

                            if (usuario_admin or usuario_equipe) and status_atual_db == "em_analise":

                                status_relato_db = relato.get("status_relato", "em_analise")
                                status_relato_label = STATUS_RELATO_LABEL.get(status_relato_db, "Em análise")

                                status_key = f"status_relato_ui_{id_relato}"
                                devolutiva_key = f"devolutiva_{id_relato}"

                                if status_key not in st.session_state:
                                    st.session_state[status_key] = status_relato_label

                                # --------------------------------------------------
                                # CONTROLE DE STATUS
                                # --------------------------------------------------
                                with st.container(horizontal=True, horizontal_alignment="right"):
                                    novo_status_label = st.segmented_control(
                                        label="",
                                        options=["Em análise", "Devolver", "Aceito"],
                                        key=status_key
                                    )

                                novo_status_db = STATUS_RELATO_LABEL_INV.get(novo_status_label)

                                # --------------------------------------------------
                                # TEXTO DE AUDITORIA (status_aprovacao)
                                # --------------------------------------------------
                                status_aprovacao = relato.get("status_aprovacao")
                                if status_aprovacao:

                                    st.markdown(
                                        f"""
                                        <div style="
                                            text-align: right;
                                            color: #6c757d;
                                            font-size: 0.85rem;
                                            margin-top: 4px;
                                        ">
                                            {status_aprovacao}
                                        </div>
                                        """,
                                        unsafe_allow_html=True
                                    )
                                    st.write('')


                                    # with st.container(horizontal=True, horizontal_alignment="right"):
                                    #     st.caption(status_aprovacao)

                                # ==================================================
                                # CASO DEVOLVER
                                # ==================================================
                                if novo_status_label == "Devolver":

                                    if devolutiva_key not in st.session_state:
                                        st.session_state[devolutiva_key] = relato.get("devolutiva", "")

                                    st.text_area(
                                        "Devolutiva:",
                                        key=devolutiva_key,
                                        placeholder="Explique o que precisa ser ajustado neste relato..."
                                    )

                                    tem_devolutiva = bool(st.session_state.get(devolutiva_key, "").strip())
                                    label_botao = "Atualizar" if tem_devolutiva else "Salvar devolutiva"

                                    with st.container(horizontal=True):

                                        if st.button(
                                            label_botao,
                                            key=f"btn_salvar_devolutiva_{id_relato}",
                                            type="primary",
                                            icon=":material/save:"
                                        ):

                                            nome = st.session_state.get("nome", "Usuário")
                                            data = data_hoje_br()

                                            relato["status_relato"] = "aberto"
                                            relato["devolutiva"] = st.session_state.get(devolutiva_key, "")
                                            relato["status_aprovacao"] = f"Devolvido por {nome} em {data}"

                                            col_projetos.update_one(
                                                {"codigo": projeto["codigo"]},
                                                {
                                                    "$set": {
                                                        "plano_trabalho.componentes": projeto["plano_trabalho"]["componentes"]
                                                    }
                                                }
                                            )

                                            st.session_state.pop(status_key, None)
                                            st.session_state.pop(devolutiva_key, None)

                                            st.success("Devolutiva salva.", icon=":material/check:")
                                            time.sleep(3)
                                            st.rerun()

                                # ==================================================
                                # CASO EM ANÁLISE OU ACEITO
                                # ==================================================
                                elif novo_status_db != status_relato_db:

                                    nome = st.session_state.get("nome", "Usuário")
                                    data = data_hoje_br()

                                    relato["status_relato"] = novo_status_db

                                    if novo_status_db == "aceito":
                                        relato.pop("devolutiva", None)
                                        relato["status_aprovacao"] = f"Verificado por {nome} em {data}"

                                        # Atualiza progresso da atividade com base no relato aprovado
                                        if "porc_ativ_relato" in relato:

                                            atividade_id = atividade["id"]
                                            porcentagem_relato = int(relato["porc_ativ_relato"])

                                            atividade_mongo = obter_atividade_mongo(projeto, atividade_id)

                                            if atividade_mongo:
                                                atividade_mongo["porcentagem_atv"] = porcentagem_relato



                                    elif novo_status_db == "em_analise":
                                        relato.pop("status_aprovacao", None)

                                    col_projetos.update_one(
                                        {"codigo": projeto["codigo"]},
                                        {
                                            "$set": {
                                                "plano_trabalho.componentes": projeto["plano_trabalho"]["componentes"]
                                            }
                                        }
                                    )

                                    st.session_state.pop(status_key, None)
                                    st.rerun()







                            # ==================================================
                            # MOSTRA DEVOLUTIVA SE EXISTIR (em_analise ou aberto)
                            # ==================================================

                            status_relato_db = relato.get("status_relato")
                            devolutiva = relato.get("devolutiva")

                            mostrar_devolutiva = False

                            # --------------------------------------------------
                            # REGRA 1: relatório em modo edição
                            # --------------------------------------------------
                            if status_atual_db == "modo_edicao":
                                mostrar_devolutiva = bool(devolutiva)

                            # --------------------------------------------------
                            # REGRA 2: relatório em análise
                            # --------------------------------------------------
                            elif status_atual_db == "em_analise":
                                # se for admin/equipe E relato está devolvido → não mostra
                                if (
                                    tipo_usuario in ["admin", "equipe"]
                                    and status_relato_db == "aberto"
                                ):
                                    mostrar_devolutiva = False
                                else:
                                    mostrar_devolutiva = bool(devolutiva)

                            if mostrar_devolutiva:

                                texto = devolutiva.replace("\n", "\n> ")

                                st.markdown(
                                    f"""
                                <blockquote style="
                                    color: #000000;
                                    opacity: 0.9;
                                    border-left: 4px solid #F8D7DA;
                                    padding-left: 12px;
                                    margin-left: 0;
                                ">
                                <strong>Ajuste necessário:</strong><br>
                                {texto.replace('\n', '<br>')}
                                </blockquote>
                                """,
                                    unsafe_allow_html=True
                                )


                            # --------------------------------------------------
                            # BOTÃO EDITAR (somente se o relato estiver aberto)
                            # --------------------------------------------------
                            if (
                                pode_editar_relatorio
                                and relato.get("status_relato") == "aberto"
                            ):
                                with st.container(horizontal=True, horizontal_alignment="right"):
                                    if st.button(
                                        "Editar",
                                        key=f"btn_edit_{id_relato}",
                                        icon=":material/edit:",
                                        type="tertiary"
                                    ):
                                        st.session_state["relato_editando_id"] = id_relato
                                        st.rerun()



                        # ==================================================
                        # MODO EDIÇÃO INLINE DO RELATO DA ATIVIDADE
                        # ==================================================
                        else:
                            st.markdown(f"**Editando {id_relato.upper()}**")



                            # -----------------------------
                            # PORCENTAGEM DA ATIVIDADE (RELATO)
                            # -----------------------------

                            # opções de 0 a 100 de 10 em 10
                            porcentagens = list(range(0, 101, 10))

                            # valor atual salvo (se existir)
                            porc_atual = 0

                            if atividade and atividade.get("id"):
                                atividade_mongo = obter_atividade_mongo(
                                    projeto,
                                    atividade["id"]
                                )

                                if atividade_mongo:
                                    porc_atual = atividade_mongo.get("porcentagem_atv", 0)

                            # garante consistência com opções
                            if porc_atual not in porcentagens:
                                porc_atual = 0

                            # sincroniza session_state ao trocar atividade
                            if (
                                "campo_porcentagem_atividade_relato" not in st.session_state
                                or st.session_state.get("atividade_porcentagem_relato_ref") != atividade.get("id")
                            ):
                                st.session_state["campo_porcentagem_atividade_relato"] = porc_atual
                                st.session_state["atividade_porcentagem_relato_ref"] = atividade.get("id")


                            # selectbox
                            porc_ativ_relato = st.selectbox(
                                "Atualize a porcentagem de execução da atividade *",
                                options=porcentagens,
                                format_func=lambda x: f"{x}%",
                                key="campo_porcentagem_atividade_relato",
                                width=300
                            )





                            # --------------------------------------------------
                            # CAMPOS DE TEXTO
                            # --------------------------------------------------
                            relato_texto = st.text_area(
                                "Relato *",
                                value=relato.get("relato", ""),
                                key=f"edit_relato_{id_relato}"
                            )


                            # --------------------------------------------------
                            # CAMPOS DE DATA NA EDIÇÃO DO RELATO
                            # --------------------------------------------------
                            # Converte as datas armazenadas como string
                            # (dd/mm/yyyy) para objeto datetime.date
                            # necessário 


                            data_inicio_str = relato.get("data_inicio")
                            data_fim_str = relato.get("data_fim")

                            # Conversão segura para datetime.date
                            data_inicio_valor = None
                            data_fim_valor = None

                            if data_inicio_str:
                                try:
                                    data_inicio_valor = datetime.datetime.strptime(data_inicio_str, "%d/%m/%Y").date()
                                except Exception:
                                    pass

                            if data_fim_str:
                                try:
                                    data_fim_valor = datetime.datetime.strptime(data_fim_str, "%d/%m/%Y").date()
                                except Exception:
                                    pass


                            # Interface de edição das datas
                            col1, col2 = st.columns(2)


                            # ---------- DATA DE INÍCIO ----------
                            with col1:
                                data_inicio = date_picker(
                                    label="Data de início",
                                    value=data_inicio_valor,
                                    format="dd/MM/yyyy",
                                    locale="pt_BR",
                                    one_tap=True,
                                    key=f"edit_data_inicio_{id_relato}"
                                )

                            # ---------- DATA DE FIM ----------
                            with col2:
                                data_fim = date_picker(
                                    label="Data de fim",
                                    value=data_fim_valor,
                                    format="dd/MM/yyyy",
                                    locale="pt_BR",
                                    one_tap=True,
                                    key=f"edit_data_fim_{id_relato}"
                                )






                            # data_inicio = col1.date_input(
                            #     "Data de início *",
                            #     value=data_inicio_valor,
                            #     key=f"edit_data_inicio_{id_relato}",
                            #     format="DD/MM/YYYY"
                            # )

                            # data_fim = col2.date_input(
                            #     "Data de fim *",
                            #     value=data_fim_valor,
                            #     key=f"edit_data_fim_{id_relato}",
                            #     format="DD/MM/YYYY"
                            # )


                            st.divider()

                            # --------------------------------------------------
                            # ANEXOS EXISTENTES (REMOVER)
                            # --------------------------------------------------
                            anexos_remover = []
                            anexos_existentes = relato.get("anexos", [])

                            if anexos_existentes:
                                st.markdown("**Anexos:**")

                                for i, a in enumerate(anexos_existentes):
                                    nome = a.get("nome_arquivo", "arquivo")

                                    if st.checkbox(
                                        f"**Remover:** {nome}",
                                        key=f"rm_anexo_{id_relato}_{i}"
                                    ):
                                        anexos_remover.append(a)

                            # --------------------------------------------------
                            # NOVOS ANEXOS
                            # --------------------------------------------------
                            st.write('')
                            novos_anexos = st.file_uploader(
                                "Adicionar novos anexos",
                                type=["pdf", "docx", "xlsx", "csv", "jpg", "jpeg", "png"],
                                accept_multiple_files=True,
                                key=f"novos_anexos_{id_relato}"
                            )

                            st.divider()

                            # --------------------------------------------------
                            # FOTOS EXISTENTES (REMOVER)
                            # --------------------------------------------------
                            fotos_remover = []
                            fotos_existentes = relato.get("fotos", [])

                            if fotos_existentes:
                                st.markdown("**Fotografias:**")

                                for i, f in enumerate(fotos_existentes):
                                    nome = f.get("nome_arquivo", "foto")
                                    descricao = f.get("descricao", "")
                                    fotografo = f.get("fotografo", "")

                                    label = nome
                                    if descricao:
                                        label += f" | {descricao}"
                                    if fotografo:
                                        label += f" | {fotografo}"

                                    if st.checkbox(
                                        f"**Remover:** {label}",
                                        key=f"rm_foto_{id_relato}_{i}"
                                    ):
                                        fotos_remover.append(f)


                            # --------------------------------------------------
                            # NOVAS FOTOS
                            # --------------------------------------------------
                            st.write('')
                            st.write("**Adicionar novas fotografias**")

                            fotos_novas_key = f"fotos_novas_{id_relato}"
                            if fotos_novas_key not in st.session_state:
                                st.session_state[fotos_novas_key] = []

                            if st.button(
                                "Adicionar fotografia",
                                key=f"btn_add_foto_{id_relato}",
                                icon=":material/add_a_photo:"
                            ):
                                st.session_state[fotos_novas_key].append({
                                    "arquivo": None,
                                    "descricao": "",
                                    "fotografo": ""
                                })

                            for i, foto in enumerate(st.session_state[fotos_novas_key]):
                                with st.container(border=True):

                                    foto["arquivo"] = st.file_uploader(
                                        "Arquivo da foto",
                                        type=["jpg", "jpeg", "png"],
                                        key=f"foto_edit_file_{id_relato}_{i}"
                                    )

                                    foto["descricao"] = st.text_input(
                                        "Descrição",
                                        key=f"foto_edit_desc_{id_relato}_{i}"
                                    )

                                    foto["fotografo"] = st.text_input(
                                        "Fotógrafo(a)",
                                        key=f"foto_edit_autor_{id_relato}_{i}"
                                    )

                            st.divider()

                            # --------------------------------------------------
                            # AÇÕES
                            # --------------------------------------------------
                            # col_save, col_cancel = st.columns([1, 1])

                            with st.container(horizontal=True, horizontal_alignment="left"):


                                if st.button(
                                    "Cancelar",
                                    key=f"btn_cancel_{id_relato}"
                                ):
                                    st.session_state["relato_editando_id"] = None
                                    st.session_state.pop(fotos_novas_key, None)
                                    st.rerun()



                                if st.button(
                                    "Salvar alterações",
                                    key=f"btn_save_{id_relato}",
                                    type="primary",
                                    icon=":material/save:"
                                ):

                                    # ==================================================
                                    # VALIDAÇÃO
                                    # ==================================================
                                    erros = []

                                    # Relato obrigatório
                                    if not relato_texto or not relato_texto.strip():
                                        erros.append("Relato")

                                    # Datas obrigatórias
                                    if not data_inicio:
                                        erros.append("Data de início")

                                    if not data_fim:
                                        erros.append("Data de fim")


                                    # Exibe erros
                                    if erros:
                                        campos = ", ".join(erros)
                                        st.warning(f"Preencha os seguintes campos obrigatórios: {campos}")
                                        st.stop()

                                    # ==================================================
                                    # SALVAMENTO
                                    # ==================================================
                                    with st.spinner("Salvando alterações. Aguarde..."):

                                        # -----------------------------
                                        # TEXTO E DATAS
                                        # -----------------------------
                                        relato["relato"] = relato_texto

                                        relato["data_inicio"] = (
                                            data_inicio.strftime("%d/%m/%Y") if data_inicio else None
                                        )

                                        relato["data_fim"] = (
                                            data_fim.strftime("%d/%m/%Y") if data_fim else None
                                        )

                                        # -----------------------------
                                        # PORCENTAGEM DO RELATO
                                        # -----------------------------
                                        relato["porc_ativ_relato"] = int(porc_ativ_relato)

                                        # ==================================================
                                        # REMOVE ITENS MARCADOS
                                        # ==================================================
                                        if anexos_remover:
                                            relato["anexos"] = [
                                                a for a in relato.get("anexos", [])
                                                if a not in anexos_remover
                                            ]

                                        if fotos_remover:
                                            relato["fotos"] = [
                                                f for f in relato.get("fotos", [])
                                                if f not in fotos_remover
                                            ]

                                        # ==================================================
                                        # DRIVE
                                        # ==================================================
                                        servico = obter_servico_drive()

                                        pasta_projeto_id = obter_pasta_projeto(
                                            servico,
                                            projeto["codigo"],
                                            projeto["sigla"]
                                        )

                                        pasta_relatos_id = obter_ou_criar_pasta(
                                            servico,
                                            "Relatos_atividades",
                                            pasta_projeto_id
                                        )

                                        pasta_relato_id = obter_ou_criar_pasta(
                                            servico,
                                            id_relato,
                                            pasta_relatos_id
                                        )

                                        # -----------------------------
                                        # ANEXOS
                                        # -----------------------------
                                        if novos_anexos:
                                            pasta_anexos_id = obter_ou_criar_pasta(
                                                servico,
                                                "anexos",
                                                pasta_relato_id
                                            )

                                            relato.setdefault("anexos", [])

                                            for arq in novos_anexos:
                                                id_drive = enviar_arquivo_drive(servico, pasta_anexos_id, arq)
                                                if id_drive:
                                                    relato["anexos"].append({
                                                        "nome_arquivo": arq.name,
                                                        "id_arquivo": id_drive
                                                    })

                                        # -----------------------------
                                        # FOTOS
                                        # -----------------------------
                                        fotos_validas = [
                                            f for f in st.session_state[fotos_novas_key]
                                            if f.get("arquivo") is not None
                                        ]

                                        if fotos_validas:
                                            pasta_fotos_id = obter_ou_criar_pasta(
                                                servico,
                                                "fotos",
                                                pasta_relato_id
                                            )

                                            relato.setdefault("fotos", [])

                                            for foto in fotos_validas:
                                                arq = foto["arquivo"]
                                                id_drive = enviar_arquivo_drive(servico, pasta_fotos_id, arq)
                                                if id_drive:
                                                    relato["fotos"].append({
                                                        "nome_arquivo": arq.name,
                                                        "descricao": foto.get("descricao", ""),
                                                        "fotografo": foto.get("fotografo", ""),
                                                        "id_arquivo": id_drive
                                                    })

                                        # ==================================================
                                        # SALVA NO MONGO
                                        # ==================================================
                                        col_projetos.update_one(
                                            {"codigo": projeto["codigo"]},
                                            {"$set": {
                                                "plano_trabalho.componentes": projeto["plano_trabalho"]["componentes"]
                                            }}
                                        )

                                        # Limpa estado
                                        st.session_state["relato_editando_id"] = None
                                        st.session_state.pop(fotos_novas_key, None)

                                        st.success("Relato atualizado com sucesso!", icon=":material/check:")
                                        time.sleep(3)
                                        st.rerun()



                    st.write('')


    if not tem_relato:
        st.caption("Nenhum relato cadastrado neste relatório.")












# ==================================================
# ---------- DESPESAS ----------
# ==================================================
if step_selecionado == "Despesas":


    st.write("")
    st.write("")

    st.markdown("#### Registros de despesas")
    st.write("")

    # --------------------------------------------------
    # PERFIS DE USUÁRIO
    # --------------------------------------------------
    usuario_admin = tipo_usuario == "admin"
    usuario_equipe = tipo_usuario == "equipe"
    usuario_beneficiario = tipo_usuario == "beneficiario"
    usuario_visitante = tipo_usuario == "visitante"

    # --------------------------------------------------
    # REGRA: quem pode registrar despesas
    # --------------------------------------------------
    pode_registrar = (
        usuario_beneficiario and status_atual_db == "modo_edicao"
    )

    # ==================================================
    # BOTÃO: REGISTRAR DESPESA
    # ==================================================
    with st.container(horizontal=True, horizontal_alignment="right"):
       

        saldo_parcela = calcular_saldo_parcela()

        saldo_formatado = f"{saldo_parcela:.1f}".replace(".", ",")


        st.markdown(
            f"Saldo disponível da parcela: "
            f"<span style='font-size:22px'><b>{saldo_formatado}%</b></span>",
            unsafe_allow_html=True
        )


        if pode_registrar:
                    if st.button(
                        "+ Registrar despesa",
                        type="primary",
                        icon=":material/add:",
                        width=260
                    ):
                        dialog_lanc_financ(
                            relatorio_numero=relatorio_numero,
                            projeto=projeto,
                            col_projetos=col_projetos
                        )

    st.write("")




    # ==================================================
    # AGRUPAMENTO DE DESPESAS (CATEGORIA > NOME)
    # ==================================================
    grupo = defaultdict(lambda: defaultdict(list))

    for despesa in projeto.get("financeiro", {}).get("orcamento", []):
        for lanc in despesa.get("lancamentos", []):
            if lanc.get("relatorio_numero") == relatorio_numero:
                grupo[despesa["categoria"]][despesa["nome_despesa"]].append(lanc)

    # --------------------------------------------------
    # SE NÃO HÁ DESPESAS
    # --------------------------------------------------
    if not grupo:
        st.caption("Nenhuma despesa registrada neste relatório.")
        st.stop()

    # ==================================================
    # RENDERIZAÇÃO DAS DESPESAS
    # ==================================================
    for categoria, despesas in grupo.items():

        st.markdown(f"##### {categoria}")

        for nome_despesa, lancamentos in despesas.items():

            st.markdown(f"###### {nome_despesa}")

            for lanc in lancamentos:

                id_despesa = lanc["id_lanc_despesa"]

                # --------------------------------------------------
                # CONTROLE DE EDIÇÃO INLINE
                # --------------------------------------------------
                if "despesa_editando_id" not in st.session_state:
                    st.session_state["despesa_editando_id"] = None

                editando = st.session_state["despesa_editando_id"] == id_despesa

                with st.container(border=True):

                    # ==================================================
                    # BADGE DE STATUS
                    # ==================================================
                    status_despesa_db = lanc.get("status_despesa", "em_analise")
                    tem_devolutiva = bool(lanc.get("devolutiva"))

                    if status_despesa_db == "aberto" and tem_devolutiva:
                        badge = {"label": "Pendente", "bg": "#F8D7DA", "color": "#721C24"}
                    elif status_despesa_db == "aberto":
                        badge = {"label": "Aberto", "bg": "#FFF3CD", "color": "#856404"}
                    elif status_despesa_db == "aceito":
                        badge = {"label": "Aceito", "bg": "#D4EDDA", "color": "#155724"}
                    else:
                        badge = {"label": "Em análise", "bg": "#D1ECF1", "color": "#0C5460"}


                    col1, col2 = st.columns([9, 1])

                    col2.markdown(
                        f"""
                        <div style="margin-bottom:6px;">
                            <span style="
                                background:{badge['bg']};
                                color:{badge['color']};
                                padding:4px 10px;
                                border-radius:20px;
                                font-size:12px;
                                font-weight:600;
                            ">
                                {badge['label']}
                            </span>
                        </div>
                        """,
                        unsafe_allow_html=True
                    )





                    # ==================================================
                    # PERMISSÕES
                    # ==================================================
                    pode_editar_despesa = (
                        usuario_beneficiario
                        and status_atual_db == "modo_edicao"
                        and status_despesa_db == "aberto"
                    )

                    pode_avaliar_despesa = (
                        (usuario_admin or usuario_equipe)
                        and status_atual_db == "em_analise"
                    )




                    # ==================================================
                    # VISUALIZAÇÃO DA DESPESA
                    # ==================================================
                    if not editando:

                        st.write(f"**{id_despesa.upper()}:** {lanc.get('descricao_despesa')}")

                        col1, col2 = st.columns(2)

                        with col1:

                            # DADOS DA DESPESA
                            def linha(label, valor):
                                c1, c2 = st.columns([1, 3])
                                c1.write(f"**{label}:**")
                                c2.write(valor if valor else "-")

                            linha("Data", lanc.get("data_despesa"))
                            linha("Fornecedor", lanc.get("fornecedor"))
                            linha("CPF/CNPJ", lanc.get("cpf_cnpj"))

                            valor = lanc.get("valor_despesa", 0)
                            valor_br = f"{valor:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
                            linha("Valor (R$)", valor_br)


                        with col2:

                            anexos = lanc.get("anexos", [])
                            if anexos:
                                st.markdown("**Anexos:**")
                                for a in anexos:
                                    link = gerar_link_drive(a["id_arquivo"])
                                    st.markdown(f"- [{a['nome_arquivo']}]({link})")




                        # ==================================================
                        # MOSTRA DEVOLUTIVA
                        # ==================================================
                        status_despesa_db = lanc.get("status_despesa")
                        devolutiva = lanc.get("devolutiva")

                        mostrar_devolutiva = False

                        # --------------------------------------------------
                        # REGRA 0: se estiver ACEITO, nunca mostra
                        # --------------------------------------------------
                        if status_despesa_db == "aceito":
                            mostrar_devolutiva = False

                        # --------------------------------------------------
                        # REGRA 1: relatório em modo edição
                        # --------------------------------------------------
                        elif status_atual_db == "modo_edicao":
                            mostrar_devolutiva = bool(devolutiva)

                        # --------------------------------------------------
                        # REGRA 2: relatório em análise
                        # --------------------------------------------------
                        elif status_atual_db == "em_analise":

                            # admin/equipe avaliando não veem devolutiva enquanto avaliam
                            if (
                                tipo_usuario in ["admin", "equipe"]
                                and status_despesa_db == "aberto"
                            ):
                                mostrar_devolutiva = False
                            else:
                                mostrar_devolutiva = bool(devolutiva)

                        # --------------------------------------------------
                        # REGRA 3: fallback seguro (ex: visitante)
                        # --------------------------------------------------
                        else:
                            mostrar_devolutiva = bool(devolutiva)

                        # --------------------------------------------------
                        # Renderização visual
                        # --------------------------------------------------
                        if mostrar_devolutiva and devolutiva:
                            texto = devolutiva.replace("\n", "<br>")

                            st.markdown(
                                f"""
                                <blockquote style="
                                    color: #000000;
                                    opacity: 0.9;
                                    border-left: 4px solid #F8D7DA;
                                    padding-left: 12px;
                                    margin-left: 0;
                                ">
                                <strong>Ajuste necessário:</strong><br>
                                {texto}
                                </blockquote>
                                """,
                                unsafe_allow_html=True
                            )


                        # --------------------------------------------------
                        # BOTÃO EDITAR (somente beneficiário, despesa aberta)
                        # --------------------------------------------------
                        if pode_editar_despesa:

                            with st.container(horizontal=True, horizontal_alignment="right"):
                                if st.button(
                                    "Editar",
                                    key=f"btn_edit_despesa_{id_despesa}",
                                    icon=":material/edit:",
                                    type="tertiary"
                                ):
                                    st.session_state["despesa_editando_id"] = id_despesa
                                    st.rerun()







                    # ==================================================
                    # MODO EDIÇÃO INLINE DA DESPESA
                    # ==================================================
                    if editando:

                        st.markdown(f"**Editando {id_despesa.upper()}**")

                        # --------------------------------------------------
                        # CAMPOS PRINCIPAIS
                        # --------------------------------------------------
                        col1, col2, col3, col4 = st.columns(4)



                        # --------------------------------------------------
                        # DATA
                        # --------------------------------------------------

                        with col1:
                            data = date_picker(
                                label="Data da despesa",
                                value=pd.to_datetime(lanc["data_despesa"], dayfirst=True),
                                format="dd/MM/yyyy",
                                locale="pt_BR",
                                one_tap=True,
                                key=f"edit_data_{id_despesa}"
                            )



                        with col2:
                            quantidade = st.number_input(
                                "Quantidade *",
                                min_value=0,
                                value=int(lanc.get("quantidade", 0)),
                                key=f"edit_qtd_{id_despesa}"
                            )

                        with col3:
                            valor_unitario = st.number_input(
                                "Valor unitário (R$) *",
                                min_value=0.0,
                                value=float(lanc.get("valor_unitario", 0)),
                                format="%.2f",
                                key=f"edit_vunit_{id_despesa}"
                            )

                        with col4:
                            valor = st.number_input(
                                "Valor total (R$) *",
                                min_value=0.0,
                                value=float(lanc.get("valor_despesa", 0)),
                                format="%.2f",
                                key=f"edit_valor_{id_despesa}"
                            )



                        descricao = st.text_area(
                            "Descrição da despesa *",
                            value=lanc.get("descricao_despesa", ""),
                            key=f"edit_desc_{id_despesa}"
                        )

                        col1, col2 = st.columns([2, 1])

                        fornecedor = col1.text_input(
                            "Fornecedor *",
                            value=lanc.get("fornecedor", ""),
                            key=f"edit_forn_{id_despesa}"
                        )

                        cpf_cnpj = col2.text_input(
                            "CPF/CNPJ *",
                            value=lanc.get("cpf_cnpj", ""),
                            key=f"edit_doc_{id_despesa}"
                        )

                        st.divider()

                        # --------------------------------------------------
                        # ANEXOS EXISTENTES (REMOVER)
                        # --------------------------------------------------
                        anexos_remover = []
                        anexos_existentes = lanc.get("anexos", [])

                        if anexos_existentes:
                            st.markdown("**Anexos:**")
                            for i, a in enumerate(anexos_existentes):
                                nome = a.get("nome_arquivo", "arquivo")

                                if st.checkbox(
                                    f"Remover: {nome}",
                                    key=f"rm_anexo_desp_{id_despesa}_{i}"
                                ):
                                    anexos_remover.append(a)

                            st.divider()

                        # --------------------------------------------------
                        # NOVOS ANEXOS
                        # --------------------------------------------------
                        novos_anexos = st.file_uploader(
                            "Adicionar novos anexos",
                            accept_multiple_files=True,
                            key=f"novos_anexos_{id_despesa}"
                        )

                        st.divider()




                        # # ==================================================
                        # # ÁREA DE MENSAGENS (ERROS / WARNINGS)
                        # # ==================================================
                        # container_erros = st.container()


                        # --------------------------------------------------
                        # AÇÕES
                        # --------------------------------------------------
                        with st.container(horizontal=True):


                            with st.container(horizontal=True):


                                if st.button(
                                    "Cancelar",
                                    key=f"btn_cancel_desp_{id_despesa}"
                                ):
                                    st.session_state["despesa_editando_id"] = None
                                    st.rerun()



                                if st.button(
                                    "Salvar alterações",
                                    key=f"btn_save_desp_{id_despesa}",
                                    type="primary",
                                    icon=":material/save:"
                                ):

                                    # ==================================================
                                    # VALIDAÇÕES
                                    # ==================================================

                                    erros_campos = []
                                    erro_consistencia = None

                                    # -------------------------------
                                    # CAMPOS OBRIGATÓRIOS
                                    # -------------------------------

                                    if not data:
                                        erros_campos.append("Data da despesa")

                                    if quantidade <= 0:
                                        erros_campos.append("Quantidade")

                                    if valor_unitario <= 0:
                                        erros_campos.append("Valor unitário (R$)")

                                    if not valor or valor <= 0:
                                        erros_campos.append("Valor total (R$)")

                                    if not descricao or not descricao.strip():
                                        erros_campos.append("Descrição da despesa")

                                    if not fornecedor or not fornecedor.strip():
                                        erros_campos.append("Fornecedor")

                                    if not cpf_cnpj or not cpf_cnpj.strip():
                                        erros_campos.append("CPF / CNPJ")

                                    # -------------------------------
                                    # CONSISTÊNCIA (CÁLCULO)
                                    # -------------------------------

                                    if quantidade > 0 and valor_unitario > 0 and valor > 0:

                                        valor_calculado = round(quantidade * valor_unitario, 2)
                                        valor_informado = round(valor, 2)

                                        if valor_calculado != valor_informado:
                                            erro_consistencia = (
                                                f"Valor total deve ser igual a Quantidade × Valor unitário"
                                            )



                                    # ==================================================
                                    # VALIDAÇÃO DE ANEXOS (COM REMOÇÃO + NOVOS)
                                    # ==================================================

                                    # Regra:
                                    # - Se NÃO for taxa bancária → precisa ter pelo menos 1 anexo no final

                                    categoria_lower = categoria.lower()
                                    is_taxa_bancaria = "taxas bancárias" in categoria_lower

                                    if not is_taxa_bancaria:

                                        anexos_existentes = lanc.get("anexos", [])

                                        # Quantos permanecem após remoção
                                        qtd_restantes = len(anexos_existentes) - len(anexos_remover)

                                        # Quantos novos serão adicionados
                                        qtd_novos = len(novos_anexos) if novos_anexos else 0

                                        total_final = qtd_restantes + qtd_novos

                                        if total_final <= 0:
                                            erros_campos.append("Anexos (mínimo de 1 arquivo)")




                                    # ==================================================
                                    # EXIBE ERROS
                                    # ==================================================

                                    if erros_campos:
                                        campos = ", ".join(erros_campos)
                                        st.warning(f"Preencha os seguintes campos obrigatórios: {campos}")

                                    if erro_consistencia:
                                        st.warning(erro_consistencia)

                                    if erros_campos or erro_consistencia:
                                        st.stop()

                                    # ==================================================
                                    # SALVAR
                                    # ==================================================
                                    with st.spinner("Salvando alterações..."):

                                        lanc.update({
                                            "data_despesa": data.strftime("%d/%m/%Y"),
                                            "descricao_despesa": descricao,
                                            "fornecedor": fornecedor,
                                            "cpf_cnpj": cpf_cnpj,
                                            "quantidade": quantidade,
                                            "valor_unitario": valor_unitario,
                                            "valor_despesa": valor
                                        })



                                        # Remove anexos marcados
                                        if anexos_remover:
                                            lanc["anexos"] = [
                                                a for a in lanc.get("anexos", [])
                                                if a not in anexos_remover
                                            ]

                                        # Upload de novos anexos
                                        if novos_anexos:
                                            servico = obter_servico_drive()
                                            pasta_proj = obter_pasta_projeto(
                                                servico,
                                                projeto["codigo"],
                                                projeto["sigla"]
                                            )
                                            pasta_fin = obter_pasta_relatos_financeiros(servico, pasta_proj)
                                            pasta_lanc = obter_ou_criar_pasta(servico, id_despesa, pasta_fin)

                                            lanc.setdefault("anexos", [])

                                            for arq in novos_anexos:
                                                id_drive = enviar_arquivo_drive(servico, pasta_lanc, arq)
                                                lanc["anexos"].append({
                                                    "nome_arquivo": arq.name,
                                                    "id_arquivo": id_drive
                                                })

                                        # Persistência no Mongo
                                        col_projetos.update_one(
                                            {"codigo": projeto["codigo"]},
                                            {"$set": {"financeiro.orcamento": projeto["financeiro"]["orcamento"]}}
                                        )

                                    # Limpa estado
                                    st.session_state["despesa_editando_id"] = None
                                    st.success("Despesa atualizada com sucesso!", icon=":material/check:")
                                    time.sleep(3)
                                    st.rerun()

                        
                            with st.container(horizontal=True):



                                # ==================================================
                                # BOTÃO EXCLUIR DESPESA (SOMENTE SE STATUS = ABERTO)
                                # ==================================================

                                status_despesa_db = lanc.get("status_despesa")

                                # Controle de estado da confirmação
                                confirm_delete_key = f"confirm_delete_despesa_{id_despesa}"

                                if confirm_delete_key not in st.session_state:
                                    st.session_state[confirm_delete_key] = False

                                # Só permite excluir se estiver aberto
                                if status_despesa_db == "aberto":

                                    with st.container(horizontal=True, horizontal_alignment="right"):

                                        # Botão inicial (ícone de lixeira)
                                        if not st.session_state[confirm_delete_key]:
                                            if st.button(
                                                "",
                                                key=f"btn_delete_{id_despesa}",
                                                icon=":material/delete:",
                                                type="secondary"
                                            ):
                                                # Ativa confirmação
                                                st.session_state[confirm_delete_key] = True
                                                st.rerun()

                                        # ==================================================
                                        # CONFIRMAÇÃO DE EXCLUSÃO
                                        # ==================================================
                                        else:

                                            with st.container(horizontal=True):

                                                st.warning("Deseja realmente excluir esta despesa?")

                                                # col1, col2 = st.columns(2)

                                                # Botão CONFIRMAR
                                                if st.button(
                                                    "Sim, excluir",
                                                    key=f"btn_confirm_delete_{id_despesa}",
                                                    type="primary",
                                                    icon=":material/delete:"
                                                ):
                                                    with st.spinner("Excluindo despesa..."):

                                                        # Remove o lançamento da estrutura
                                                        for d in projeto["financeiro"]["orcamento"]:
                                                            if d["categoria"] == categoria and d["nome_despesa"] == nome_despesa:
                                                                d["lancamentos"] = [
                                                                    l for l in d.get("lancamentos", [])
                                                                    if l.get("id_lanc_despesa") != id_despesa
                                                                ]
                                                                break

                                                        # Salva no Mongo
                                                        col_projetos.update_one(
                                                            {"codigo": projeto["codigo"]},
                                                            {"$set": {"financeiro.orcamento": projeto["financeiro"]["orcamento"]}}
                                                        )

                                                    # Limpa estados
                                                    st.session_state["despesa_editando_id"] = None
                                                    st.session_state.pop(confirm_delete_key, None)

                                                    st.success("Despesa excluída com sucesso!", icon=":material/check:")
                                                    time.sleep(3)
                                                    st.rerun()

                                                # Botão CANCELAR
                                                if st.button(
                                                    "Cancelar",
                                                    key=f"btn_cancel_delete_{id_despesa}"
                                                ):
                                                    st.session_state[confirm_delete_key] = False
                                                    st.rerun()









                    # ==================================================
                    # AVALIAÇÃO (ADMIN / EQUIPE) — MESMA REGRA DE ATIVIDADES
                    # ==================================================
                    if pode_avaliar_despesa:

                        STATUS_DESPESA_LABEL = {
                            "em_analise": "Em análise",
                            "aberto": "Devolver",
                            "aceito": "Aceito"
                        }

                        STATUS_DESPESA_LABEL_INV = {v: k for k, v in STATUS_DESPESA_LABEL.items()}

                        status_despesa_db = lanc.get("status_despesa", "em_analise")
                        status_label = STATUS_DESPESA_LABEL.get(status_despesa_db, "Em análise")

                        status_key = f"status_despesa_ui_{id_despesa}"
                        devolutiva_key = f"devolutiva_despesa_{id_despesa}"

                        # --------------------------------------------------
                        # Estado inicial do segmented_control
                        # Regra igual à Atividades:
                        # aberto sem devolutiva → Em análise
                        # --------------------------------------------------
                        if status_despesa_db == "aberto" and not lanc.get("devolutiva"):
                            status_label = "Em análise"

                        if status_key not in st.session_state:
                            st.session_state[status_key] = status_label

                        # --------------------------------------------------
                        # SEGMENTED CONTROL
                        # --------------------------------------------------
                        with st.container(horizontal=True, horizontal_alignment="right"):
                            novo_status_label = st.segmented_control(
                                label="",
                                options=["Em análise", "Devolver", "Aceito"],
                                key=status_key
                            )

                        novo_status_db = STATUS_DESPESA_LABEL_INV.get(novo_status_label)

                        # --------------------------------------------------
                        # TEXTO DE AUDITORIA
                        # --------------------------------------------------
                        status_aprovacao = lanc.get("status_aprovacao")
                        if status_aprovacao:
                            st.markdown(
                                f"""
                                <div style="
                                    text-align: right;
                                    color: rgba(0,0,0,0.55);
                                    font-size: 0.8rem;
                                    margin-top: 4px;
                                ">
                                    {status_aprovacao}
                                </div>
                                """,
                                unsafe_allow_html=True
                            )
                            st.write("")

                        # ==================================================
                        # CASO DEVOLVER (ação, não mudança de status)
                        # ==================================================
                        if novo_status_label == "Devolver":

                            if devolutiva_key not in st.session_state:
                                st.session_state[devolutiva_key] = lanc.get("devolutiva", "")

                            st.text_area(
                                "**Devolutiva:**",
                                key=devolutiva_key,
                                placeholder="Explique o que precisa ser ajustado nesta despesa..."
                            )

                            tem_devolutiva = bool(st.session_state.get(devolutiva_key, "").strip())
                            label_botao = "Atualizar" if tem_devolutiva else "Salvar devolutiva"

                            with st.container(horizontal=True):
                                if st.button(
                                    label_botao,
                                    key=f"btn_save_dev_{id_despesa}",
                                    type="primary",
                                    icon=":material/save:"
                                ):
                                    nome = st.session_state.get("nome", "Usuário")
                                    data = data_hoje_br()

                                    lanc["status_despesa"] = "aberto"
                                    lanc["devolutiva"] = st.session_state.get(devolutiva_key, "")
                                    lanc["status_aprovacao"] = f"Devolvido por {nome} em {data}"

                                    col_projetos.update_one(
                                        {"codigo": projeto["codigo"]},
                                        {"$set": {"financeiro.orcamento": projeto["financeiro"]["orcamento"]}}
                                    )

                                    st.session_state.pop(status_key, None)
                                    st.session_state.pop(devolutiva_key, None)

                                    st.success("Devolutiva salva.", icon=":material/check:")
                                    time.sleep(3)
                                    st.rerun()

                        # ==================================================
                        # CASO EM ANÁLISE OU ACEITO (mudança real de status)
                        # ==================================================
                        elif novo_status_db != status_despesa_db:

                            nome = st.session_state.get("nome", "Usuário")
                            data = data_hoje_br()

                            lanc["status_despesa"] = novo_status_db

                            if novo_status_db == "aceito":
                                lanc.pop("devolutiva", None)
                                lanc["status_aprovacao"] = f"Verificado por {nome} em {data}"

                            elif novo_status_db == "em_analise":
                                lanc.pop("status_aprovacao", None)

                            col_projetos.update_one(
                                {"codigo": projeto["codigo"]},
                                {"$set": {"financeiro.orcamento": projeto["financeiro"]["orcamento"]}}
                            )

                            st.session_state.pop(status_key, None)
                            st.rerun()


                    # ==================================================
                    # MOSTRA DEVOLUTIVA (MESMA REGRA DE ATIVIDADES)
                    # ==================================================
                    status_despesa_db = lanc.get("status_despesa")
                    devolutiva = lanc.get("devolutiva")

                    mostrar_devolutiva = False

                    # Regra 1: relatório em modo edição
                    if status_atual_db == "modo_edicao":
                        mostrar_devolutiva = bool(devolutiva)

                    # Regra 2: relatório em análise
                    elif status_atual_db == "em_analise":
                        if (
                            tipo_usuario in ["admin", "equipe"]
                            and status_despesa_db == "aberto"
                        ):
                            mostrar_devolutiva = False
                        else:
                            mostrar_devolutiva = bool(devolutiva)















# ==================================================
# ---------- RESULTADOS ----------
# ==================================================







if step_selecionado == "Resultados":

    # Espaçamento visual
    st.write("")
    st.write("")

    # Título da seção
    st.markdown("#### Indicadores de projeto")
    st.write("")

    # Recupera os componentes do plano de trabalho
    componentes = projeto.get("plano_trabalho", {}).get("componentes", [])

    # Lista auxiliar para armazenar todas as entregas,
    # independentemente do componente
    entregas = []

    # Percorre os componentes e coleta todas as entregas
    for componente in componentes:
        for entrega in componente.get("entregas", []):
            entregas.append(entrega)

    # Caso não existam entregas
    if not entregas:
        st.info("Este projeto não possui entregas com indicadores.")
    else:
        # Loop por entrega
        for idx_ent, entrega in enumerate(entregas):

            # Título da entrega
            st.markdown(f"##### {entrega.get('entrega')}")

            # Lista de indicadores do projeto dentro da entrega
            indicadores = entrega.get("indicadores_projeto", [])

            # Caso a entrega não tenha indicadores
            if not indicadores:
                st.caption("Esta entrega não possui indicadores de projeto.")
                continue

            # Loop por indicador
            for idx_ind, indicador in enumerate(indicadores):

                # Container visual para cada indicador
                with st.container(border=True):

                    # Nome do indicador
                    st.markdown(
                        f"**Indicador:** {indicador.get('indicador_projeto')}"
                    )

                    # Unidade de medida
                    st.markdown(
                        f"**Unidade de medida:** {indicador.get('unidade_medida')}"
                    )

                    # Layout em colunas
                    col1, col2, col3, col4 = st.columns([1, 1, 1, 3])

                    # Linha base (somente leitura)
                    col1.markdown(
                        f"**Início do projeto:** {indicador.get('linha_base')}"
                    )

                    # Meta (somente leitura)
                    col2.markdown(
                        f"**Meta:** {indicador.get('meta')}"
                    )

                    # ======================================================
                    # KEYS ISOLADAS
                    # ======================================================

                    key_resultado = (
                        f"resultado_"
                        f"{relatorio_numero}_"
                        f"{entrega.get('id')}_"
                        f"{idx_ind}"
                    )

                    key_obs = (
                        f"obs_"
                        f"{relatorio_numero}_"
                        f"{entrega.get('id')}_"
                        f"{idx_ind}"
                    )

                    key_save = (
                        f"save_"
                        f"{relatorio_numero}_"
                        f"{entrega.get('id')}_"
                        f"{idx_ind}"
                    )

                    # ======================================================
                    # PRÉ-CARGA DO ESTADO (somente uma vez)
                    # ======================================================

                    # Resultado atual (float do banco → string pt-BR)
                    if key_resultado not in st.session_state:
                        valor_resultado = indicador.get("resultado_atual")

                        if valor_resultado is None:
                            st.session_state[key_resultado] = ""
                        else:
                            st.session_state[key_resultado] = (
                                formatar_numero_br_dinamico(valor_resultado)
                            )

                    # Observações
                    if key_obs not in st.session_state:
                        valor_observacoes = indicador.get("observacoes_coleta")
                        if valor_observacoes is None or valor_observacoes == "None":
                            valor_observacoes = ""
                        st.session_state[key_obs] = valor_observacoes

                    # ======================================================
                    # RENDERIZAÇÃO CONDICIONAL
                    # ======================================================

                    if pode_editar_relatorio:
                        # Campo editável: resultado atual (texto, formato BR)
                        resultado_atual_str = col3.text_input(
                            "Resultado atual",
                            key=key_resultado,
                            # placeholder="Ex: 1.234,56"
                        )

                        # Campo editável: observações
                        observacoes = col4.text_input(
                            "Observações",
                            key=key_obs
                        )
                    else:
                        # Apenas exibição
                        col3.write(
                            f"**Resultado atual:** "
                            f"{formatar_numero_br_dinamico(indicador.get('resultado_atual'))}"
                        )

                        valor_obs = st.session_state[key_obs]
                        if not valor_obs:
                            valor_obs = "-"

                        col4.write(
                            f"**Observações:** {valor_obs}"
                        )

                        resultado_atual_str = st.session_state[key_resultado]
                        observacoes = st.session_state[key_obs]

                    # ======================================================
                    # BOTÃO SALVAR (sempre aparece quando pode editar)
                    # ======================================================

                    if pode_editar_relatorio:
                        with st.container(horizontal=True, horizontal_alignment="right"):
                            salvar = st.button(
                                "Salvar",
                                key=key_save,
                                icon=":material/save:",
                                width=200,
                                type="primary"
                            )
                    else:
                        salvar = False

                    # ======================================================
                    # Mostra data da última coleta, se existir
                    # ======================================================



                    data_coleta = indicador.get("data_coleta")

                    if data_coleta:

                        with st.container(horizontal=True, horizontal_alignment="right"):

                            if isinstance(data_coleta, datetime.datetime):

                                # 🔥 CORREÇÃO IMPORTANTE
                                if data_coleta.tzinfo is None:
                                    data_coleta = data_coleta.replace(
                                        tzinfo=datetime.timezone.utc
                                    )

                                data_local = data_coleta.astimezone(
                                    ZoneInfo("America/Sao_Paulo")
                                )

                                data_str = data_local.strftime("%d/%m/%Y %H:%M")

                            else:
                                data_str = str(data_coleta)

                            st.caption(f"Último registro em {data_str}")



                    # ======================================================
                    # AÇÃO DE SALVAMENTO
                    # ======================================================

                    if salvar:
                        # Converte string pt-BR para float
                        resultado_float = parse_numero_br(resultado_atual_str)

                        if resultado_float is None:
                            st.error(
                                "Resultado atual inválido. "
                                "Use o formato brasileiro, por exemplo: 1.234,56"
                            )
                            st.stop()



                        data_coleta = datetime.datetime.now(datetime.timezone.utc)
                        # data_coleta = datetime.datetime.now()

                        observacoes_salvar = observacoes
                        if observacoes_salvar is None or observacoes_salvar == "None":
                            observacoes_salvar = ""

                        # Atualiza no MongoDB
                        col_projetos.update_one(
                            {
                                "codigo": projeto_codigo
                            },
                            {
                                "$set": {
                                    "plano_trabalho.componentes.$[c].entregas.$[e].indicadores_projeto.$[i].resultado_atual": resultado_float,
                                    "plano_trabalho.componentes.$[c].entregas.$[e].indicadores_projeto.$[i].observacoes_coleta": observacoes_salvar,
                                    "plano_trabalho.componentes.$[c].entregas.$[e].indicadores_projeto.$[i].data_coleta": data_coleta
                                }
                            },
                            array_filters=[
                                {"c.entregas.id": entrega.get("id")},
                                {"e.id": entrega.get("id")},
                                {"i.indicador_projeto": indicador.get("indicador_projeto")}
                            ]
                        )

                        # Atualiza o objeto em memória
                        indicador["resultado_atual"] = resultado_float
                        indicador["observacoes_coleta"] = observacoes_salvar
                        indicador["data_coleta"] = data_coleta

                        st.success("Indicador salvo com sucesso.", icon=":material/check:")
                        time.sleep(3)
                        st.rerun()

                # Espaçamento entre indicadores
                st.write("")





















# ---------- BENEFÍCIOS ----------

if step_selecionado == "Beneficiários":


    # =====================================================
    # CARREGA TIPOS DE BENEFÍCIO DO BANCO
    # =====================================================

    dados_beneficios = list(
        col_beneficios.find({}, {"beneficio": 1}).sort("beneficio", 1)
    )

    OPCOES_BENEFICIOS = [
        d["beneficio"]
        for d in dados_beneficios
        if d.get("beneficio")
    ]


    # ============================================
    # CONTROLE DE USUÁRIO / STATUS DO RELATÓRIO
    # ============================================

    usuario_admin = tipo_usuario == "admin"
    usuario_equipe = tipo_usuario == "equipe"
    usuario_beneficiario = tipo_usuario == "beneficiario"
    usuario_visitante = tipo_usuario == "visitante"

    # Se o relatório NÃO estiver em modo_edicao,
    # força modo VISUALIZAÇÃO dos beneficiários
    if status_atual_db != "modo_edicao":
        modo_edicao_benef = False
        modo_visualizacao_benef = True
    else:
        modo_edicao_benef = usuario_beneficiario
        modo_visualizacao_benef = not usuario_beneficiario





    # PARTE 1 - QUANTITATIVO DE BENEFICIÁRIOS ---------------------------------------------------------------------------------------------------------------------------
    st.write('')
    st.write('')




    # ======================================================
    # INICIALIZAÇÃO DO ESTADO DA MATRIZ DE BENEFICIÁRIOS
    # ======================================================


    key_benef_quant = f"beneficiarios_quant_rel_{relatorio_numero}"

    if key_benef_quant not in st.session_state:
        st.session_state[key_benef_quant] = (
            relatorio.get("beneficiarios_quant") or {
                "mulheres": {"jovens": 0, "adultas": 0, "idosas": 0},
                "homens": {"jovens": 0, "adultos": 0, "idosos": 0},
                "nao_binarios": {"jovens": 0, "adultos": 0, "idosos": 0}
            }
        )




    # ======================================================
    # TÍTULO DO BLOCO
    # ======================================================

    st.markdown("##### Número de beneficiários por gênero e faixa etária")

    st.write("")


    # ======================================================
    # MODO EDIÇÃO
    # ======================================================

    if pode_editar_relatorio:


        # Coluna à esquerda para diminuir a largura dos inputs de beneficiários
        content, vazio_d = st.columns([8, 4])

        # -------------------------------
        # LINHA: JOVENS
        # -------------------------------
        col_m, col_h, col_nb = content.columns(3)

        with col_m:
            st.session_state[key_benef_quant]["mulheres"]["jovens"] = st.number_input(
                "Mulheres – Jovens (até 24 anos)",
                min_value=0,
                step=1,
                value=st.session_state[key_benef_quant]["mulheres"]["jovens"],
                key="bq_mulheres_jovens"
            )

        with col_h:
            st.session_state[key_benef_quant]["homens"]["jovens"] = st.number_input(
                "Homens – Jovens (até 24 anos)",
                min_value=0,
                step=1,
                value=st.session_state[key_benef_quant]["homens"]["jovens"],
                key="bq_homens_jovens"
            )

        with col_nb:
            st.session_state[key_benef_quant]["nao_binarios"]["jovens"] = st.number_input(
                "Não-binários – Jovens (até 24 anos)",
                min_value=0,
                step=1,
                value=st.session_state[key_benef_quant]["nao_binarios"]["jovens"],
                key="bq_nb_jovens"
            )

        # -------------------------------
        # LINHA: ADULTOS
        # -------------------------------
        col_m, col_h, col_nb = content.columns(3)

        with col_m:
            st.session_state[key_benef_quant]["mulheres"]["adultas"] = st.number_input(
                "Mulheres – Adultas",
                min_value=0,
                step=1,
                value=st.session_state[key_benef_quant]["mulheres"]["adultas"],
                key="bq_mulheres_adultas"
            )

        with col_h:
            st.session_state[key_benef_quant]["homens"]["adultos"] = st.number_input(
                "Homens – Adultos",
                min_value=0,
                step=1,
                value=st.session_state[key_benef_quant]["homens"]["adultos"],
                key="bq_homens_adultos"
            )

        with col_nb:
            st.session_state[key_benef_quant]["nao_binarios"]["adultos"] = st.number_input(
                "Não-binários – Adultos",
                min_value=0,
                step=1,
                value=st.session_state[key_benef_quant]["nao_binarios"]["adultos"],
                key="bq_nb_adultos"
            )

        # -------------------------------
        # LINHA: IDOSOS
        # -------------------------------
        col_m, col_h, col_nb = content.columns(3)

        with col_m:
            st.session_state[key_benef_quant]["mulheres"]["idosas"] = st.number_input(
                "Mulheres – Idosas (60+ anos)",
                min_value=0,
                step=1,
                value=st.session_state[key_benef_quant]["mulheres"]["idosas"],
                key="bq_mulheres_idosas"
            )

        with col_h:
            st.session_state[key_benef_quant]["homens"]["idosos"] = st.number_input(
                "Homens – Idosos (60+ anos)",
                min_value=0,
                step=1,
                value=st.session_state[key_benef_quant]["homens"]["idosos"],
                key="bq_homens_idosos"
            )

        with col_nb:
            st.session_state[key_benef_quant]["nao_binarios"]["idosos"] = st.number_input(
                "Não-binários – Idosos (60+ anos)",
                min_value=0,
                step=1,
                value=st.session_state[key_benef_quant]["nao_binarios"]["idosos"],
                key="bq_nb_idosos"
            )




        # ======================================================
        # BOTÃO DE SALVAR EXCLUSIVO DA MATRIZ
        # ======================================================
        # Este botão salva SOMENTE a matriz de quantitativos

        if pode_editar_relatorio:

            st.write("")

            salvar_matriz = st.button(
                "Atualizar beneficiários",
                type="primary",
                key=f"salvar_beneficiarios_quant_{relatorio_numero}",
                icon=":material/save:"
            )

            if salvar_matriz:

                # Atualiza apenas a chave 'beneficiarios_quant' no relatório correto
                col_projetos.update_one(
                    {
                        "codigo": projeto["codigo"],
                        "relatorios.numero": relatorio_numero
                    },
                    {
                        "$set": {
                            "relatorios.$.beneficiarios_quant":
                                st.session_state[key_benef_quant]
                        }
                    }
                )

                st.success("Beneficiários salvos com sucesso.", icon=":material/check:")
                time.sleep(3)
                st.rerun()











    # ======================================================
    # MODO VISUALIZAÇÃO
    # ======================================================

    else:


        dados = st.session_state[key_benef_quant]


        # -------------------------------
        # Totais por gênero
        # -------------------------------
        total_mulheres = sum(dados["mulheres"].values())
        total_homens = sum(dados["homens"].values())
        total_nb = sum(dados["nao_binarios"].values())

        # -------------------------------
        # Totais por faixa etária
        # -------------------------------
        total_jovens = (
            dados["mulheres"]["jovens"]
            + dados["homens"]["jovens"]
            + dados["nao_binarios"]["jovens"]
        )

        total_adultos = (
            dados["mulheres"]["adultas"]
            + dados["homens"]["adultos"]
            + dados["nao_binarios"]["adultos"]
        )

        total_idosos = (
            dados["mulheres"]["idosas"]
            + dados["homens"]["idosos"]
            + dados["nao_binarios"]["idosos"]
        )

        total_geral = total_mulheres + total_homens + total_nb

        st.write("")

        # -------------------------------
        # LAYOUT EM 4 COLUNAS
        # -------------------------------
        col_m, col_h, col_nb, col_totais = st.columns(4)

        # -------- MULHERES --------
        with col_m:
            l, v = st.columns(2)
            l.write("Mulheres jovens"); v.write(str(dados["mulheres"]["jovens"]))

            l, v = st.columns(2)
            l.write("Mulheres adultas"); v.write(str(dados["mulheres"]["adultas"]))

            l, v = st.columns(2)
            l.write("Mulheres idosas"); v.write(str(dados["mulheres"]["idosas"]))

            l, v = st.columns(2)
            l.markdown("**Total de mulheres**"); v.markdown(f"**{total_mulheres}**")

        # -------- HOMENS --------
        with col_h:
            l, v = st.columns(2)
            l.write("Homens jovens"); v.write(str(dados["homens"]["jovens"]))

            l, v = st.columns(2)
            l.write("Homens adultos"); v.write(str(dados["homens"]["adultos"]))

            l, v = st.columns(2)
            l.write("Homens idosos"); v.write(str(dados["homens"]["idosos"]))

            l, v = st.columns(2)
            l.markdown("**Total de homens**"); v.markdown(f"**{total_homens}**")

        # -------- NÃO-BINÁRIOS --------
        with col_nb:
            l, v = st.columns(2)
            l.write("Não-binários jovens"); v.write(str(dados["nao_binarios"]["jovens"]))

            l, v = st.columns(2)
            l.write("Não-binários adultos"); v.write(str(dados["nao_binarios"]["adultos"]))

            l, v = st.columns(2)
            l.write("Não-binários idosos"); v.write(str(dados["nao_binarios"]["idosos"]))

            l, v = st.columns(2)
            l.markdown("**Total de não-binários**"); v.markdown(f"**{total_nb}**")

        # -------- TOTAIS GERAIS (NEGRITO) --------
        with col_totais:
            l, v = st.columns(2)
            l.markdown("**Total de jovens**"); v.markdown(f"**{total_jovens}**")

            l, v = st.columns(2)
            l.markdown("**Total de adultos**"); v.markdown(f"**{total_adultos}**")

            l, v = st.columns(2)
            l.markdown("**Total de idosos**"); v.markdown(f"**{total_idosos}**")

            l, v = st.columns(2)
            l.markdown("**Total geral**"); v.markdown(f"**{total_geral}**")








    st.divider()

    # ============================================================================================================
    # PARTE 2 - TIPOS DE BENEFICIÁRIOS E BENEFICIOS 
    # ============================================================================================================

    st.write('')
    st.markdown("##### Tipos de Beneficiários e Benefícios")

    if usuario_beneficiario:

        st.write("")
        st.write(
            "Registre aqui os tipos de **Beneficiários** e **Benefícios** do projeto para cada comunidade."
        )

    st.write(
        "Se precisar, cadastre novas comunidades na opção **Locais** no menu lateral."
    )

    st.write("")
    st.write("")


    projeto = col_projetos.find_one({"codigo": projeto["codigo"]})
    localidades = projeto.get("locais", {}).get("localidades", [])

    if not localidades:
        st.info(
            "Nenhuma comunidade cadastrada no projeto. "
            "Adicione comunidades na página **Locais**."
        )
        st.stop()

    # =====================================================
    # LOOP DAS COMUNIDADES
    # =====================================================
    for localidade in localidades:

        nome_localidade = localidade.get("nome_localidade")
        beneficiarios_bd = localidade.get("beneficiarios", []) or []

        # -------------------------------------------------
        # ESTADO ORIGINAL DO BANCO
        # -------------------------------------------------
        estado_original = {
            b["tipo_beneficiario"]: sorted(b.get("beneficios") or [])
            for b in beneficiarios_bd
            if b.get("tipo_beneficiario")
        }

        # -------------------------------------------------
        # PÚBLICOS PARA RENDERIZAÇÃO
        # -------------------------------------------------
        publicos_renderizacao = list(opcoes_publicos[:-1])

        for tipo in estado_original.keys():
            if tipo not in publicos_renderizacao:
                publicos_renderizacao.append(tipo)

        publicos_renderizacao = sorted(publicos_renderizacao)

        estado_atual = {}
        houve_alteracao = False

        col1, col2 = st.columns([1, 3])

        # -------- COLUNA 1 --------

        with col1:
            st.markdown(f"**{nome_localidade}**")

            municipio = localidade.get("municipio")

            if municipio:
                st.write(municipio)




        # -------- COLUNA 2 --------
        with col2:

            st.write("**Tipos de Beneficiários e Benefícios:**")



            # =====================================================
            # MODO VISUALIZAÇÃO COM LISTA EM PILLS
            # =====================================================
            if modo_visualizacao_benef:

                if not beneficiarios_bd:
                    st.write("Nenhum beneficiário cadastrado.")
                else:
                    for b in beneficiarios_bd:

                        tipo = b.get("tipo_beneficiario")
                        beneficios = b.get("beneficios") or []

                        with st.container():
                            st.write(' ')
                            if beneficios:
                                st.pills(
                                    label=tipo,
                                    options=beneficios,
                                    width="content",
                                    key=f"pill_{projeto['codigo']}_{nome_localidade}_{tipo}"
                                )
                            else:
                                st.pills(
                                    label=tipo,
                                    options=["Nenhum benefício informado"],
                                    width="content",
                                    key=f"pill_{projeto['codigo']}_{nome_localidade}_{tipo}"
                                )


            # =====================================================
            # MODO EDIÇÃO
            # =====================================================
            if modo_edicao_benef:

                # =============================================
                # BENEFICIÁRIOS EXISTENTES
                # =============================================
                for publico in publicos_renderizacao:

                    with st.container(horizontal=True):

                        chk_key = f"chk_{projeto['codigo']}_{nome_localidade}_{publico}"

                        marcado_inicial = publico in estado_original

                        marcado = st.checkbox(
                            publico,
                            value=marcado_inicial,
                            key=chk_key,
                            width=300
                        )

                        if marcado:

                            beneficios_iniciais = estado_original.get(publico, [])

                            beneficios = st.multiselect(
                                f"Benefícios para {publico.lower()}",
                                options=OPCOES_BENEFICIOS,
                                default=beneficios_iniciais,
                                key=f"ms_{projeto['codigo']}_{nome_localidade}_{publico}"
                            )

                            estado_atual[publico] = sorted(beneficios)

                            if (
                                publico not in estado_original
                                or sorted(beneficios) != estado_original.get(publico, [])
                            ):
                                houve_alteracao = True

                        else:
                            if publico in estado_original:
                                houve_alteracao = True

                # =============================================
                # CHECKBOX OUTROS
                # =============================================
                with st.container(horizontal=True):

                    chk_outros_key = f"chk_outros_{projeto['codigo']}_{nome_localidade}"

                    outros_marcado = st.checkbox(
                        "Outros",
                        value=False,
                        key=chk_outros_key,
                        width=300
                    )

                # =============================================
                # FORMULÁRIO OUTROS
                # =============================================
                if outros_marcado:

                    with st.container(horizontal=True):

                        st.text_input(
                            "Tipo de beneficiário",
                            key=f"novo_tipo_{projeto['codigo']}_{nome_localidade}"
                        )

                        st.multiselect(
                            "Benefícios",
                            options=OPCOES_BENEFICIOS,
                            key=f"novo_beneficios_{projeto['codigo']}_{nome_localidade}"
                        )

                    novo_tipo = st.session_state.get(
                        f"novo_tipo_{projeto['codigo']}_{nome_localidade}", ""
                    ).strip()

                    novos_beneficios = st.session_state.get(
                        f"novo_beneficios_{projeto['codigo']}_{nome_localidade}", []
                    )

                    if novo_tipo and novos_beneficios:
                        houve_alteracao = True

        # =================================================
        # BOTÃO SALVAR
        # =================================================
        if houve_alteracao:

            st.write("")

            erros = []

            # with st.container(horizontal=True, horizontal_alignment="right"):
            clicou_salvar = st.button(
                f"Atualizar {nome_localidade}",
                type="primary",
                key=f"salvar_{projeto['codigo']}_{nome_localidade}",
                icon=":material/save:"
            )

            if clicou_salvar:

                beneficiarios_para_salvar = []

                # -----------------------------------------
                # BENEFICIÁRIOS EXISTENTES
                # -----------------------------------------
                for tipo, beneficios in estado_atual.items():
                    if not beneficios:
                        erros.append(
                            f"Selecione ao menos um benefício para **{tipo}**."
                        )
                    else:
                        beneficiarios_para_salvar.append({
                            "tipo_beneficiario": tipo,
                            "beneficios": beneficios
                        })

                # -----------------------------------------
                # NOVO BENEFICIÁRIO (OUTROS)
                # -----------------------------------------
                if outros_marcado and novo_tipo:
                    beneficiarios_para_salvar.append({
                        "tipo_beneficiario": novo_tipo,
                        "beneficios": novos_beneficios
                    })

                if erros:
                    for erro in erros:
                        st.error(erro)
                    time.sleep(3)
                    st.rerun()

                # -----------------------------------------
                # SALVA NO BANCO
                # -----------------------------------------
                col_projetos.update_one(
                    {
                        "codigo": projeto["codigo"],
                        "locais.localidades.nome_localidade": nome_localidade
                    },
                    {
                        "$set": {
                            "locais.localidades.$.beneficiarios":
                                beneficiarios_para_salvar
                        }
                    }
                )

                st.success(
                    f"Beneficiários da comunidade "
                    f"**{nome_localidade}** salvos com sucesso.",
                    icon=":material/check:"
                )
                time.sleep(3)
                st.rerun()


        st.divider()





# ---------- PESQUISAS ----------
if step_selecionado == "Pesquisas":

    # ============================
    # CONTROLE DE USUÁRIO
    # ============================

    usuario_admin = tipo_usuario == "admin"
    usuario_equipe = tipo_usuario == "equipe"
    usuario_beneficiario = tipo_usuario == "beneficiario"
    

    pode_editar = usuario_admin or usuario_equipe or usuario_beneficiario
    pode_verificar = usuario_admin or usuario_equipe

    # ============================
    # BUSCA DADOS
    # ============================

    pesquisas = edital.get("pesquisas_relatorio", []) if edital else []

    if not pesquisas:
        st.caption("Nenhuma pesquisa cadastrada.")
        st.stop()

    st.write("")
    st.write("")
    st.markdown("##### Pesquisas / Ferramentas de Monitoramento")
    st.write("")

    pesquisas_projeto = projeto.get("pesquisas", [])
    status_map = {p["id_pesquisa"]: p for p in pesquisas_projeto}

    # ============================
    # RENDERIZAÇÃO DAS LINHAS
    # ============================

    for pesquisa in pesquisas:

        status = status_map.get(pesquisa["id"], {})

        # Valores atuais do banco
        respondida_db = status.get("respondida", False)
        verificada_db = status.get("verificada", False)
        url_anexo_db = status.get("url_anexo")

        # Chaves únicas
        upload_key = f"upload_{relatorio_numero}_{pesquisa['id']}"
        upload_salvo_key = f"upload_salvo_{relatorio_numero}_{pesquisa['id']}"

        col1, col2, col3, col4, col5 = st.columns([4, 3, 2, 2, 2])

        # -------- PESQUISA --------
        with col1:
            st.markdown(f"**{pesquisa['nome_pesquisa']}**")


        # -------- ANEXO --------
        arquivo = None

        with col2:
            # Caso a pesquisa exija upload
            if pesquisa.get("upload_arquivo"):

                # -----------------------------
                # BENEFICIÁRIO → pode anexar
                # -----------------------------
                if (
                    tipo_usuario == "beneficiario"
                    and not verificada_db
                    and status_atual_db == "modo_edicao"
                ):
                    arquivo = st.file_uploader(
                        "Anexo",
                        key=f"upload_{relatorio_numero}_{pesquisa['id']}"
                    )

                # -----------------------------
                # NÃO BENEFICIÁRIO
                # Mostra aviso SOMENTE se não houver anexo salvo
                # -----------------------------
                elif tipo_usuario != "beneficiario" and not url_anexo_db:
                    st.write(":material/attach_file: Demanda anexo")

            # -----------------------------
            # Link do anexo (se existir)
            # -----------------------------
            if url_anexo_db:
                st.markdown(f":material/attach_file: [Ver anexo]({url_anexo_db})")



        # -------- RESPONDIDA --------
        with col3:
            respondida_ui = st.checkbox(
                "Respondida",
                value=respondida_db,
                disabled = (
                    # Visitante nunca pode
                    tipo_usuario == "visitante"

                    # Beneficiário só pode no modo edição
                    or (
                        tipo_usuario == "beneficiario"
                        and status_atual_db != "modo_edicao"
                    )

                    # Beneficiário não pode se já verificada
                    or (
                        tipo_usuario == "beneficiario"
                        and verificada_db
                    )

                    # Admin/equipe não podem no modo edição
                    or (
                        tipo_usuario in ["admin", "equipe"]
                        and status_atual_db == "modo_edicao"
                    )
                ),
                key=f"resp_{relatorio_numero}_{pesquisa['id']}"
            )

        # -------- VERIFICADA --------
        with col4:
            verificada_ui = st.checkbox(
                "Verificada",
                value=verificada_db,
                disabled = (
                    # Visitante nunca pode
                    tipo_usuario == "visitante"

                    # Beneficiário nunca pode verificar
                    or tipo_usuario == "beneficiario"

                    # Relatório em modo edição trava todos
                    or status_atual_db == "modo_edicao"
                ),
                key=f"verif_{relatorio_numero}_{pesquisa['id']}"
            )

        # -------- DETECTA ALTERAÇÃO --------
        linha_modificada = (
            respondida_ui != respondida_db
            or verificada_ui != verificada_db
            or (
                arquivo is not None
                and not st.session_state.get(upload_salvo_key, False)
            )
        )

        # -------- BOTÃO SALVAR --------
        with col5:
            if linha_modificada and pode_editar:

                if st.button(
                    "Salvar",
                    type="primary",
                    key=f"salvar_{relatorio_numero}_{pesquisa['id']}",
                    icon=":material/save:",
                ):


                    with st.spinner("Salvando..."):

                        # Conecta ao Drive SOMENTE aqui
                        servico = obter_servico_drive()

                        # Pasta do projeto
                        pasta_projeto = obter_pasta_projeto(
                            servico,
                            projeto["codigo"],
                            projeto["sigla"]
                        )

                        # Pasta Pesquisas (direto no projeto)
                        pasta_pesquisas = obter_pasta_pesquisas(
                            servico,
                            pasta_projeto,
                            projeto["codigo"]
                        )

                        url_anexo_final = url_anexo_db  # valor já salvo no banco (se existir)

                        # ------------------------------
                        # UPLOAD (somente se houver novo arquivo)
                        # ------------------------------
                        if (
                            arquivo is not None
                            and not st.session_state.get(upload_salvo_key, False)
                        ):
                            id_drive = enviar_arquivo_drive(
                                servico,
                                pasta_pesquisas,
                                arquivo
                            )

                            url_anexo_final = gerar_link_drive(id_drive)

                            # Marca upload como concluído
                            st.session_state[upload_salvo_key] = True

                        # ------------------------------
                        # MONTA O OBJETO DA PESQUISA
                        # ------------------------------
                        pesquisa_obj = {
                            "id_pesquisa": pesquisa["id"],
                            "respondida": respondida_ui,
                            "verificada": verificada_ui
                        }

                        if url_anexo_final:
                            pesquisa_obj["url_anexo"] = url_anexo_final

                        # ------------------------------
                        # VERIFICA SE JÁ EXISTE NO PROJETO
                        # ------------------------------
                        existe = col_projetos.count_documents(
                            {
                                "codigo": codigo_projeto_atual,
                                "pesquisas.id_pesquisa": pesquisa["id"]
                            }
                        ) > 0

                        if existe:
                            col_projetos.update_one(
                                {
                                    "codigo": codigo_projeto_atual,
                                    "pesquisas.id_pesquisa": pesquisa["id"]
                                },
                                {
                                    "$set": {
                                        "pesquisas.$": pesquisa_obj
                                    }
                                }
                            )
                        else:
                            col_projetos.update_one(
                                {"codigo": codigo_projeto_atual},
                                {
                                    "$push": {
                                        "pesquisas": pesquisa_obj
                                    }
                                }
                            )



                    # Limpa estados temporários
                    st.session_state.pop(upload_key, None)
                    st.session_state.pop(upload_salvo_key, None)

                    st.success(":material/check: Salvo!")
                    time.sleep(3)
                    st.rerun()

        st.divider()




# ---------- FORMULÁRIO ----------
if step_selecionado == "Formulário":

    ###########################################################################
    # 1. BUSCA O EDITAL CORRESPONDENTE AO PROJETO
    ###########################################################################

    edital = col_editais.find_one(
        {"codigo_edital": projeto["edital"]}
    )

    if not edital:
        st.error("Edital não encontrado para este projeto.")
        st.stop()

    perguntas = edital.get("perguntas_relatorio", [])

    if not perguntas:
        st.write('')
        st.error("O edital não possui perguntas cadastradas.")
        st.stop()

    # Ordena as perguntas pela ordem definida no edital
    perguntas = sorted(perguntas, key=lambda x: x.get("ordem", 0))


    ###########################################################################
    # 2. CONTROLE DE ESTADO POR RELATÓRIO (EVITA VAZAMENTO ENTRE ABAS)
    ###########################################################################

    # Identificador único do relatório atual
    relatorio_numero = relatorio["numero"]
    chave_relatorio_ativo = f"form_relatorio_{relatorio_numero}"

    # Se mudou de relatório, recarrega respostas do banco
    if st.session_state.get("form_relatorio_ativo") != chave_relatorio_ativo:
        st.session_state.form_relatorio_ativo = chave_relatorio_ativo


        # -------------------------------------------
        # CARREGA RESPOSTAS DO RELATÓRIO (DICT DE OBJETOS)
        # -------------------------------------------

        # Identificador único do relatório
        relatorio_numero = relatorio["numero"]

        # Evita vazamento entre abas
        if st.session_state.get("form_relatorio_ativo") != relatorio_numero:
            st.session_state.form_relatorio_ativo = relatorio_numero

            # Dicionário
            st.session_state.respostas_formulario = (
                relatorio.get("respostas_formulario", {}).copy()
            )




    ###########################################################################
    # 3. RENDERIZAÇÃO DO FORMULÁRIO
    ###########################################################################

    st.write("")
    st.write("")

    # -------------------------------------------------------------------------
    # Armazena uploads temporários em memória (evita múltiplos envios no rerun)
    # Somente no clique do botão "Salvar formulário" os arquivos serão enviados
    # -------------------------------------------------------------------------
    if "temp_uploads" not in st.session_state:
        st.session_state.temp_uploads = {}


    for pergunta in perguntas:
        tipo = pergunta.get("tipo")
        texto = pergunta.get("pergunta")
        opcoes = pergunta.get("opcoes", [])
        ordem = pergunta.get("ordem")

        # Chave única da pergunta dentro do relatório
        chave = f"pergunta_{ordem}"


        # ---------------------------------------------------------------------
        # TÍTULO (não salva resposta)
        # ---------------------------------------------------------------------
        if tipo == "titulo":
            st.subheader(texto)
            st.write("")
            continue


        # ---------------------------------------------------------------------
        # SUBTÍTULO (não salva resposta)
        # ---------------------------------------------------------------------
        elif tipo == "subtitulo":
            st.markdown(f"##### {texto}")
            st.write("")
            continue


        # ---------------------------------------------------------------------
        # PARÁGRAFO → apenas texto informativo
        # ---------------------------------------------------------------------
        elif tipo == "paragrafo":
            st.write(texto)
            st.write("")
            continue


        # ---------------------------------------------------------------------
        # TEXTO CURTO
        # ---------------------------------------------------------------------
        elif tipo == "texto_curto":

            resposta_atual = (
                st.session_state.respostas_formulario
                .get(chave, {})
                .get("resposta", "")
            )

            if pode_editar_relatorio:
                resposta = st.text_input(
                    label=texto,
                    value=resposta_atual,
                    key=f"input_{chave}"
                )

                st.session_state.respostas_formulario[chave] = {
                    "tipo": tipo,
                    "ordem": ordem,
                    "pergunta": texto,
                    "resposta": resposta
                }
            else:
                renderizar_visualizacao(texto, resposta_atual)


        # ---------------------------------------------------------------------
        # TEXTO LONGO
        # ---------------------------------------------------------------------
        elif tipo == "texto_longo":

            resposta_atual = (
                st.session_state.respostas_formulario
                .get(chave, {})
                .get("resposta", "")
            )

            if pode_editar_relatorio:
                resposta = st.text_area(
                    label=texto,
                    value=resposta_atual,
                    height=150,
                    key=f"input_{chave}"
                )

                st.session_state.respostas_formulario[chave] = {
                    "tipo": tipo,
                    "ordem": ordem,
                    "pergunta": texto,
                    "resposta": resposta
                }
            else:
                renderizar_visualizacao(texto, resposta_atual)


        # ---------------------------------------------------------------------
        # NÚMERO
        # ---------------------------------------------------------------------
        elif tipo == "numero":

            resposta_atual = (
                st.session_state.respostas_formulario
                .get(chave, {})
                .get("resposta", 0)
            )

            if pode_editar_relatorio:
                resposta = st.number_input(
                    label=texto,
                    value=float(resposta_atual),
                    step=1.0,
                    format="%g",
                    key=f"input_{chave}"
                )

                st.session_state.respostas_formulario[chave] = {
                    "tipo": tipo,
                    "ordem": ordem,
                    "pergunta": texto,
                    "resposta": resposta
                }
            else:
                renderizar_visualizacao(
                    texto,
                    formatar_numero_br_dinamico(resposta_atual)
                )


        # ---------------------------------------------------------------------
        # ESCOLHA ÚNICA
        # ---------------------------------------------------------------------
        elif tipo == "escolha_unica":

            resposta_atual = (
                st.session_state.respostas_formulario
                .get(chave, {})
                .get("resposta", opcoes[0] if opcoes else "")
            )

            if pode_editar_relatorio:
                resposta = st.radio(
                    label=texto,
                    options=opcoes,
                    index=opcoes.index(resposta_atual) if resposta_atual in opcoes else 0,
                    key=f"input_{chave}"
                )

                st.session_state.respostas_formulario[chave] = {
                    "tipo": tipo,
                    "ordem": ordem,
                    "pergunta": texto,
                    "resposta": resposta
                }
            else:
                renderizar_visualizacao(texto, resposta_atual)


        # ---------------------------------------------------------------------
        # MÚLTIPLA ESCOLHA
        # ---------------------------------------------------------------------
        elif tipo == "multipla_escolha":

            resposta_atual = (
                st.session_state.respostas_formulario
                .get(chave, {})
                .get("resposta", [])
            )

            if pode_editar_relatorio:
                resposta = st.multiselect(
                    label=texto,
                    options=opcoes,
                    default=resposta_atual,
                    key=f"input_{chave}"
                )

                st.session_state.respostas_formulario[chave] = {
                    "tipo": tipo,
                    "ordem": ordem,
                    "pergunta": texto,
                    "resposta": resposta
                }
            else:
                renderizar_visualizacao(texto, ", ".join(resposta_atual))



        # ---------------------------------------------------------------------
        # UPLOAD DE ARQUIVOS
        # ---------------------------------------------------------------------
        elif tipo == "upload_arquivo":

            MAX_MB = 10
            MAX_BYTES = MAX_MB * 1024 * 1024

            resposta_atual = (
                st.session_state.respostas_formulario
                .get(chave, {})
                .get("resposta", [])
            )

            if pode_editar_relatorio:

                arquivos = st.file_uploader(
                    label=f"{texto} (máx. 10 MB por arquivo)",
                    accept_multiple_files=True,
                    key=f"input_{chave}"
                )

                # ---------------------------------------------------------
                # Validação 
                # ---------------------------------------------------------
                if arquivos:
                    validos = [
                        arq for arq in arquivos
                        if arq.size <= MAX_BYTES
                    ]

                    for arq in arquivos:
                        if arq.size > MAX_BYTES:
                            st.warning(
                                f"O arquivo '{arq.name}' excede 10 MB e não será enviado."
                            )

                    # substitui (não acumula)
                    st.session_state.temp_uploads[chave] = validos
                else:
                    # se remover seleção, limpa também
                    st.session_state.temp_uploads.pop(chave, None)

                # ---------------------------------------------------------
                # Lista de arquivos já salvos (após uploader)
                # ---------------------------------------------------------
                if resposta_atual:
                    st.caption("Arquivos já enviados:")
                    for arq in resposta_atual:
                        link = gerar_link_drive(arq["id"])
                        st.markdown(
                            f":material/attach_file: [{arq['nome']}]({link})"
                        )

                st.session_state.respostas_formulario[chave] = {
                    "tipo": tipo,
                    "ordem": ordem,
                    "pergunta": texto,
                    "resposta": resposta_atual
                }

            else:
                st.markdown(f"**{texto}**")

                if resposta_atual:
                    for arq in resposta_atual:
                        link = gerar_link_drive(arq["id"])
                        st.markdown(
                            f":material/attach_file: [{arq['nome']}]({link})"
                        )
                else:
                    st.caption("Nenhum arquivo enviado")


        # ---------------------------------------------------------------------
        # TIPO NÃO SUPORTADO
        # ---------------------------------------------------------------------
        else:
            st.warning(f"Tipo de pergunta não suportado: {tipo}")

        st.write("")




    ###########################################################################
    # 4. BOTÃO PARA SALVAR RESPOSTAS + UPLOAD REAL PARA O DRIVE
    ###########################################################################
    if pode_editar_relatorio:
        if st.button("Salvar formulário", type="primary", icon=":material/save:"):

            with st.spinner("Salvando o formulário..."):

                servico = None

                # ---------------------------------------------------------
                # Upload incremental (somente se houver novos arquivos)
                # ---------------------------------------------------------
                for chave, arquivos in list(st.session_state.temp_uploads.items()):

                    if not arquivos:
                        continue

                    if not servico:
                        servico = obter_servico_drive()

                    pasta_projeto_id = obter_pasta_projeto(
                        servico,
                        projeto["codigo"],
                        projeto["sigla"]
                    )

                    pasta_relatorios_id = obter_pasta_relatorios(
                        servico,
                        pasta_projeto_id
                    )

                    novos_arquivos = []

                    for arquivo in arquivos:
                        arquivo_id = enviar_arquivo_drive(
                            servico,
                            pasta_relatorios_id,
                            arquivo
                        )

                        if arquivo_id:
                            novos_arquivos.append({
                                "id": arquivo_id,
                                "nome": arquivo.name
                            })

                    existentes = (
                        st.session_state.respostas_formulario[chave]
                        .get("resposta", [])
                    )

                    st.session_state.respostas_formulario[chave]["resposta"] = (
                        existentes + novos_arquivos
                    )

                # ---------------------------------------------------------
                # LIMPEZA CRÍTICA (evita duplicação no rerun)
                # ---------------------------------------------------------
                st.session_state.temp_uploads = {}

                # limpa widgets file_uploader
                for k in list(st.session_state.keys()):
                    if k.startswith("input_pergunta_"):
                        del st.session_state[k]

                # ---------------------------------------------------------
                # Salva no Mongo
                # ---------------------------------------------------------
                col_projetos.update_one(
                    {
                        "codigo": projeto["codigo"],
                        "relatorios.numero": relatorio_numero
                    },
                    {
                        "$set": {
                            "relatorios.$.respostas_formulario":
                                st.session_state.respostas_formulario
                        }
                    }
                )

            st.success("Respostas salvas com sucesso!", icon=":material/check:")
            time.sleep(3)
            st.rerun()









# ---------- ENVIAR ----------


if step_selecionado == "Enviar":

    st.write('')
    st.write('')

    # --------------------------------------------------
    # CASO 1: RELATÓRIO JÁ ENVIADO (EM ANÁLISE)
    # --------------------------------------------------
    if status_atual_db == "em_analise":

        # Recupera a data de envio salva no banco
        data_envio = relatorio.get("data_envio")

        # Formata a data para exibição (DD/MM/YYYY)
        if data_envio:

            # Converte string no formato brasileiro (dd/mm/yyyy) para datetime
            data_dt = datetime.datetime.strptime(
                data_envio, "%d/%m/%Y"
            )

            # Mantém o mesmo formato para exibição
            data_formatada = data_dt.strftime("%d/%m/%Y")


            # data_formatada = datetime.datetime.strptime(
            #     data_envio, "%Y-%m-%d"
            # ).strftime("%d/%m/%Y")
        else:
            data_formatada = "—"

        st.markdown(
            f"##### Relatório enviado em {data_formatada}.")

        st.write("Aguardando análise.")
    # --------------------------------------------------
    # CASO 2: RELATÓRIO APROVADO
    # --------------------------------------------------
    elif status_atual_db == "aprovado":
        st.markdown("##### Relatório aprovado.")

    # --------------------------------------------------
    # CASO 3: RELATÓRIO EM MODO EDIÇÃO E USUÁRIO PODE EDITAR
    # --------------------------------------------------
    elif pode_editar_relatorio:

        st.markdown("### Enviar relatório")

        saldo_parcela = calcular_saldo_parcela()

        saldo_formatado = f"{saldo_parcela:.1f}".replace(".", ",")


        # Mensagem do saldo 
        if saldo_parcela > 20:

            st.markdown(
                f"A parcela atual ainda tem "
                f"<span style='font-size:22px'><b>{saldo_formatado}%</b></span> de saldo.",
                unsafe_allow_html=True
            )

            st.markdown(
                "Recomendamos que **envie o relatório** quando o saldo for **menor que 20%**."
            )

        else:
            st.markdown(
                f"A parcela atual tem "
                f"<span style='font-size:22px'><b>{saldo_formatado}%</b></span> de saldo.",
                unsafe_allow_html=True
            )

            st.markdown(
                "**O relatório já pode ser enviado.**"
            )



        st.divider()
        
        st.write(
            "Ao enviar o relatório, ele será encaminhado para análise "
            "e não poderá mais ser editado enquanto estiver em análise."
        )

        enviar = st.button(
            "Enviar relatório",
            type="primary",
            icon=":material/send:"
        )

        if enviar:

            # Gera a data de envio
            data_envio = datetime.datetime.now().strftime("%d/%m/%Y")

            with st.spinner("Enviando relatório ..."):

                # --------------------------------------------------
                # 1. ATUALIZA STATUS E DATA DO RELATÓRIO
                # --------------------------------------------------
                col_projetos.update_one(
                    {
                        "codigo": projeto_codigo,
                        "relatorios.numero": relatorio_numero
                    },
                    {
                        "$set": {
                            "relatorios.$.status_relatorio": "em_analise",
                            "relatorios.$.data_envio": data_envio
                        }
                    }
                )

                # --------------------------------------------------
                # 2. ATUALIZA STATUS DOS RELATOS ABERTOS
                #    (somente os relatos deste relatório)
                # --------------------------------------------------
                projeto_atualizado = col_projetos.find_one(
                    {"codigo": projeto_codigo}
                )

                componentes = projeto_atualizado["plano_trabalho"]["componentes"]

                houve_alteracao = False


                # ------------------------------------------------------
                # Percorre componentes do plano de trabalho
                # ------------------------------------------------------
                for componente in componentes:

                    # Recupera entregas de forma segura
                    entregas = componente.get("entregas", [])

                    # Se não houver entregas, pula o componente
                    if not entregas:
                        continue

                    # --------------------------------------------------
                    # Percorre entregas
                    # --------------------------------------------------
                    for entrega in entregas:

                        # Recupera atividades de forma segura
                        atividades = entrega.get("atividades", [])

                        # --------------------------------------------------
                        # Se não houver atividades, apenas continua
                        # (não quebra o código)
                        # --------------------------------------------------
                        if not atividades:
                            continue

                        # --------------------------------------------------
                        # Percorre atividades
                        # --------------------------------------------------
                        for atividade in atividades:

                            # Recupera relatos de forma segura
                            relatos = atividade.get("relatos", [])

                            # Se não houver relatos, continua
                            if not relatos:
                                continue

                            # --------------------------------------------------
                            # Percorre relatos
                            # --------------------------------------------------
                            for relato in relatos:

                                # ----------------------------------------------
                                # Apenas relatos do relatório atual
                                # e que ainda estejam abertos
                                # ----------------------------------------------
                                if (
                                    relato.get("relatorio_numero") == relatorio_numero
                                    and relato.get("status_relato") == "aberto"
                                ):
                                    relato["status_relato"] = "em_analise"
                                    houve_alteracao = True



                # Salva no Mongo apenas se houve mudança
                if houve_alteracao:
                    col_projetos.update_one(
                        {"codigo": projeto_codigo},
                        {
                            "$set": {
                                "plano_trabalho.componentes": componentes
                            }
                        }
                    )


                # --------------------------------------------------
                # ENVIA E-MAIL PARA PADRINHOS
                # --------------------------------------------------
                
                
                notificar_padrinhos_relatorio(
                    col_pessoas=col_pessoas,
                    numero_relatorio=relatorio_numero,
                    projeto=projeto_atualizado,
                    logo_url=logo_cepf
                )


            st.success("Relatório enviado para análise.", icon=":material/check:")

            # Reseta para o rerun não se perder.
            st.session_state.step_relatorio = "Atividades"

            time.sleep(3)
            st.rerun()

    # --------------------------------------------------
    # CASO 4: USUÁRIO NÃO PODE EDITAR
    # --------------------------------------------------
    else:
    
        if "visitante" in st.session_state.tipo_usuario:
            st.caption("Função indisponível para visitantes.")
    
        else:

            st.caption("Este relatório não pode ser editado no momento.")











# ---------- AVALIAÇÃO ----------
if step_selecionado == "Avaliação":

    st.write("")
    st.write("")

    relatos_ok = todos_relatos_aceitos(projeto, relatorio_numero)
    despesas_ok = todas_despesas_aceitas(projeto, relatorio_numero)

    relatorio_db = next(
        r for r in projeto["relatorios"]
        if r["numero"] == relatorio_numero
    )

    col1, col2, col3 = st.columns(3, gap="large")

    # Checklist
    with col1:
        st.write("**Checklist**")

        st.checkbox(
            "Relatos de atividades (auto)",
            value=relatos_ok,
            disabled=True,
            key=f"chk_relatos_{relatorio_numero}"
        )

        st.checkbox(
            "Registros de despesas (auto)",
            value=despesas_ok,
            disabled=True,
            key=f"chk_despesas_{relatorio_numero}"
        )




        # -----------------------------
        # RESULTADOS
        # -----------------------------
        res_key = f"chk_res_{relatorio_numero}"

        st.checkbox(
            "Resultados",
            value="res_verif_por" in relatorio_db,
            key=res_key,
            on_change=atualizar_verificacao_relatorio,
            args=(
                projeto_codigo,
                relatorio_numero,
                "res_verif_por",
                res_key
            )
        )

        if relatorio_db.get("res_verif_por"):
            st.caption(relatorio_db["res_verif_por"])








        # -----------------------------
        # BENEFICIÁRIOS
        # -----------------------------
        benef_key = f"chk_benef_{relatorio_numero}"
        st.checkbox(
            "Beneficiários e Benefícios",
            value="benef_verif_por" in relatorio_db,
            key=benef_key,
            on_change=atualizar_verificacao_relatorio,
            args=(
                projeto_codigo,
                relatorio_numero,
                "benef_verif_por",
                benef_key
            )
        )

        if relatorio_db.get("benef_verif_por"):
            st.caption(relatorio_db["benef_verif_por"])

        # -----------------------------
        # PESQUISAS
        # -----------------------------
        pesq_key = f"chk_pesq_{relatorio_numero}"
        st.checkbox(
            "Pesquisas",
            value="pesq_verif_por" in relatorio_db,
            key=pesq_key,
            on_change=atualizar_verificacao_relatorio,
            args=(
                projeto_codigo,
                relatorio_numero,
                "pesq_verif_por",
                pesq_key
            )
        )

        if relatorio_db.get("pesq_verif_por"):
            st.caption(relatorio_db["pesq_verif_por"])

        # -----------------------------
        # FORMULÁRIO
        # -----------------------------
        form_key = f"chk_form_{relatorio_numero}"
        st.checkbox(
            "Formulário",
            value="form_verif_por" in relatorio_db,
            key=form_key,
            on_change=atualizar_verificacao_relatorio,
            args=(
                projeto_codigo,
                relatorio_numero,
                "form_verif_por",
                form_key
            )
        )

        if relatorio_db.get("form_verif_por"):
            st.caption(relatorio_db["form_verif_por"])



    # Anotações
    with col2:

        st.write("**Anotações**")

        # --------------------------------------------------
        # DIALOG DE NOVA ANOTAÇÃO
        # --------------------------------------------------
        @st.dialog("Nova anotação")
        def dialog_nova_anotacao():
            texto = st.text_area(
                "Anotação",
                placeholder="Digite sua anotação sobre este relatório..."
            )

            if st.button("Salvar anotação", type="primary", icon=":material/save:"):
                if not texto.strip():
                    st.warning("A anotação não pode estar vazia.")
                    return

                nova = {
                    "texto_anotacao": texto.strip(),
                    "data_anotacao": datetime.datetime.now().strftime("%d/%m/%Y"),
                    "autor_anotacao": st.session_state.get("nome", "Usuário")
                }

                col_projetos.update_one(
                    {
                        "codigo": projeto_codigo,
                        "relatorios.numero": relatorio_numero
                    },
                    {
                        "$push": {
                            "relatorios.$.anotacoes_avaliacao": nova
                        }
                    }
                )

                st.success("Anotação salva com sucesso.", icon=":material/check:")
                time.sleep(3)
                st.rerun()

        # --------------------------------------------------
        # BOTÃO NOVA ANOTAÇÃO
        # --------------------------------------------------
        if st.button(
            "+ Nova anotação",
            type="secondary",
            icon=":material/add:"
        ):
            dialog_nova_anotacao()







        # --------------------------------------------------
        # RENDERIZAÇÃO DAS ANOTAÇÕES (POPOVER COM AÇÕES)
        # --------------------------------------------------

        if "anotacao_editando" not in st.session_state:
            st.session_state["anotacao_editando"] = None

        if "anotacao_apagando" not in st.session_state:
            st.session_state["anotacao_apagando"] = None

        anotacoes = relatorio_db.get("anotacoes_avaliacao", [])

        if not anotacoes:
            st.caption("Nenhuma anotação registrada.")
        else:
            for i, a in enumerate(reversed(anotacoes)):

                idx_real = len(anotacoes) - 1 - i
                autor = a.get("autor_anotacao")
                data = a.get("data_anotacao")
                texto = a.get("texto_anotacao")

                with st.container(border=True):

                    # Cabeçalho
                    col_h1, col_h2 = st.columns([9, 1])
                    col_h1.markdown(f"**{autor}** · {data}")

                    # --------------------------------------------------
                    # POPOVER DE AÇÕES (somente autor)
                    # --------------------------------------------------
                    if st.session_state.get("nome") == autor:

                        with col_h2.popover("⋮", type="tertiary"):

                            if st.button(
                                "Editar anotação",
                                key=f"btn_edit_anot_{relatorio_numero}_{idx_real}",
                                icon=":material/edit:",
                                type="tertiary"
                            ):
                                st.session_state["anotacao_editando"] = idx_real
                                st.session_state["anotacao_apagando"] = None
                                st.rerun()

                            if st.button(
                                "Apagar anotação",
                                key=f"btn_del_anot_{relatorio_numero}_{idx_real}",
                                icon=":material/delete:",
                                type="tertiary"
                            ):
                                st.session_state["anotacao_apagando"] = idx_real
                                st.session_state["anotacao_editando"] = None
                                st.rerun()

                    # --------------------------------------------------
                    # CONFIRMAÇÃO DE EXCLUSÃO
                    # --------------------------------------------------
                    if st.session_state["anotacao_apagando"] == idx_real:

                        st.warning("Tem certeza que deseja apagar esta anotação? Esta ação não pode ser desfeita.", icon=":material/warning:")

                        with st.container(horizontal=True):

                            if st.button(
                                "Sim, apagar anotação",
                                key=f"btn_confirm_del_{relatorio_numero}_{idx_real}",
                                type="primary",
                                icon=":material/delete:"
                            ):
                                del projeto["relatorios"][idx]["anotacoes_avaliacao"][idx_real]

                                col_projetos.update_one(
                                    {"codigo": projeto_codigo},
                                    {"$set": {"relatorios": projeto["relatorios"]}}
                                )

                                st.success("Anotação apagada.", icon=":material/check:")
                                time.sleep(3)

                                st.session_state["anotacao_apagando"] = None
                                st.rerun()

                            if st.button(
                                "Cancelar",
                                key=f"btn_cancel_del_{relatorio_numero}_{idx_real}"
                            ):
                                st.session_state["anotacao_apagando"] = None
                                st.rerun()

                    # --------------------------------------------------
                    # MODO EDIÇÃO
                    # --------------------------------------------------
                    elif st.session_state["anotacao_editando"] == idx_real:

                        text_key = f"text_anot_{relatorio_numero}_{idx_real}"

                        if text_key not in st.session_state:
                            st.session_state[text_key] = texto

                        novo_texto = st.text_area(
                            "Editar anotação",
                            key=text_key
                        )

                        with st.container(horizontal=True):

                            if st.button(
                                "Atualizar",
                                key=f"btn_upd_{relatorio_numero}_{idx_real}",
                                type="primary",
                                icon=":material/save:"
                            ):
                                projeto["relatorios"][idx]["anotacoes_avaliacao"][idx_real]["texto_anotacao"] = novo_texto

                                col_projetos.update_one(
                                    {"codigo": projeto_codigo},
                                    {"$set": {"relatorios": projeto["relatorios"]}}
                                )

                                st.success("Anotação atualizada.")
                                time.sleep(3)

                                st.session_state["anotacao_editando"] = None
                                st.session_state.pop(text_key, None)
                                st.rerun()

                            if st.button(
                                "Cancelar",
                                key=f"btn_cancel_edit_{relatorio_numero}_{idx_real}"
                            ):
                                st.session_state["anotacao_editando"] = None
                                st.session_state.pop(text_key, None)
                                st.rerun()

                    # --------------------------------------------------
                    # MODO VISUALIZAÇÃO
                    # --------------------------------------------------
                    else:
                        st.markdown(
                            texto.replace("\n", "<br>"),
                            unsafe_allow_html=True
                        )
                        




    # ==================================================
    # COLUNA 3 — APROVAÇÃO DO RELATÓRIO
    # ==================================================
    with col3:

        st.write("**Aprovação**")
        st.write("")

        # --------------------------------------------------
        # REGRA: só pode aprovar se TODO checklist estiver OK
        # --------------------------------------------------
        pode_aprovar = all([
            relatos_ok,
            despesas_ok,
            "res_verif_por" in relatorio_db,
            "benef_verif_por" in relatorio_db,
            "pesq_verif_por" in relatorio_db,
            "form_verif_por" in relatorio_db
        ])



        # --------------------------------------------------
        # BOTÃO DE APROVAÇÃO
        # --------------------------------------------------
        if st.button(
            "Aprovar e enviar e-mail",
            type="primary",
            icon=":material/check_circle:",
            disabled=not pode_aprovar
        ):

            with st.spinner("Aprovando relatório..."):

                # Data atual (dd/mm/yyyy)
                data_hoje = datetime.datetime.now().strftime("%d/%m/%Y")

                # Nome do aprovador
                nome_aprovador = st.session_state.get("nome", "Usuário")

                # --------------------------------------------------
                # ATUALIZA RELATÓRIO EM MEMÓRIA
                # --------------------------------------------------
                projeto["relatorios"][idx]["status_relatorio"] = "aprovado"
                projeto["relatorios"][idx]["data_aprovacao"] = data_hoje
                projeto["relatorios"][idx]["aprovado_por"] = nome_aprovador

                # --------------------------------------------------
                # PERSISTE NO BANCO DE DADOS
                # --------------------------------------------------
                col_projetos.update_one(
                    {"codigo": projeto_codigo},
                    {"$set": {"relatorios": projeto["relatorios"]}}
                )

                # --------------------------------------------------
                # ENVIO DE E-MAIL PARA TODOS OS CONTATOS
                # --------------------------------------------------
                contatos_notificados = []

                for contato in projeto.get("contatos", []):

                    email = contato.get("email")
                    nome_contato = contato.get("nome", "Olá")

                    if not email:
                        continue

                    corpo_html = gerar_email_relatorio_aprovado(
                        nome_do_contato=nome_contato,
                        relatorio_numero=relatorio_numero,
                        projeto=projeto,
                        organizacao=projeto.get("organizacao", ""),
                        logo_url=logo_cepf
                    )

                    enviar_email(
                        corpo_html=corpo_html,
                        destinatarios=[email],
                        assunto=f"Relatório {relatorio_numero} aprovado!"
                    )

                    # Guarda nome para feedback final
                    contatos_notificados.append(nome_contato)

            # --------------------------------------------------
            # FEEDBACK VISUAL E RECARREGAMENTO
            # --------------------------------------------------
            if contatos_notificados:
                nomes = ", ".join(contatos_notificados)
                st.success(
                    f"Relatório aprovado e e-mails enviados com sucesso para {nomes}.",
                    icon=":material/check:"
                )
            else:
                st.success(
                    "Relatório aprovado, mas não havia contatos com e-mail para notificação.",
                    icon=":material/check:"
                )

            time.sleep(10)
            st.rerun()

        # --------------------------------------------------
        # INFORMAÇÃO DE APROVAÇÃO (APÓS APROVAR)
        # --------------------------------------------------
        if relatorio_db.get("status_relatorio") == "aprovado":

            data_aprov = relatorio_db.get("data_aprovacao")
            nome_aprov = relatorio_db.get("aprovado_por", "")

            if data_aprov:
                st.caption(f"Aprovado em {data_aprov} por {nome_aprov}")
                st.caption("Os contatos do projeto foram notificados por e-mail.")









# ###################################################################################################
# SIDEBAR DA PÁGINA DO PROJETO
# ###################################################################################################

sidebar_projeto()
