import streamlit as st
import pandas as pd
import datetime
import time
from docx import Document
from io import BytesIO
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Pt

from funcoes_auxiliares import (
    conectar_mongo_cepf_gestao,
    sidebar_projeto,
    obter_pasta_planos_mitigacao,
    obter_pasta_projeto,
    obter_servico_drive,
    enviar_arquivo_drive,
    gerar_link_drive
)


st.set_page_config(page_title="Salvaguardas", page_icon=":material/health_and_safety:")




# ###########################################################################################################
# # CONFIGURAÇÕES DO STREAMLIT
# ###########################################################################################################


# # Traduzindo o texto do st.file_uploader
# # Texto interno
# st.markdown("""
# <style>
# /* Esconde o texto padrão */
# [data-testid="stFileUploaderDropzone"] div div::before {
#     content: "Arraste e solte os arquivos aqui";
#     color: rgba(49, 51, 63, 0.7);
#     font-size: 0.9rem;
#     font-weight: 400;
#     position: absolute;
#     top: 50px;              /* fixa no topo */
#     left: 50%;
#     transform: translate(-50%, 10%);
#     pointer-events: none;
# }
# /* Esconde o texto original */
# [data-testid="stFileUploaderDropzone"] div div span {
#     visibility: hidden !important;
# }
# </style>
# """, unsafe_allow_html=True)

# # Traduzindo Botão do file_uploader
# st.markdown("""
# <style>
# /* Alvo: apenas o botão dentro do componente de upload */
# section[data-testid="stFileUploaderDropzone"] button[data-testid="stBaseButton-secondary"] {
#     font-size: 0px !important;   /* esconde o texto original */
#     padding-left: 14px !important;
#     padding-right: 14px !important;
#     min-width: 160px !important;
# }
# /* Insere o texto traduzido */
# section[data-testid="stFileUploaderDropzone"] button[data-testid="stBaseButton-secondary"]::after {
#     content: "Selecionar arquivo";
#     font-size: 14px !important;
#     color: inherit;
# }
# </style>
# """, unsafe_allow_html=True)






###########################################################################################################
# CONEXÃO COM O BANCO DE DADOS MONGODB
###########################################################################################################

# Conecta-se ao banco de dados MongoDB (usa cache automático para melhorar performance)
db = conectar_mongo_cepf_gestao()

# Coleção de projetos
col_projetos = db["projetos"]

# Coleção de organizações
col_organizacoes = db["organizacoes"]

###########################################################################################################
# CARREGAMENTO DE DADOS
###########################################################################################################


# Capturando o código do projeto e os dados do projeto
codigo_projeto_atual = st.session_state.get("projeto_atual")

if not codigo_projeto_atual:
    st.error("Nenhum projeto selecionado.")
    st.stop()


df_projeto = pd.DataFrame(
    list(
        col_projetos.find(
            {"codigo": codigo_projeto_atual}
        )
    )
)


# ###################################################################################################
# SIDEBAR DA PÁGINA DO PROJETO
# ###################################################################################################

sidebar_projeto()




###########################################################################################################
# FUNÇÕES 
###########################################################################################################


def formatar_tabela_docx(tabela):
    """
    Aplica formatação padrão em toda a tabela.
    """

    for row in tabela.rows:

        for cell in row.cells:

            for paragrafo in cell.paragraphs:

                paragrafo.paragraph_format.space_before = Pt(6)
                paragrafo.paragraph_format.space_after = Pt(6)




# FUNÇÃO - RENDERIZAÇÃO DO PLANO DE MITIGAÇÃO
def renderizar_plano_mitigacao(
    salvaguardas_doc,
    mapa_politicas,
    nome_politica,
    codigo_projeto_atual,
    col_projetos
):
    """
    Renderiza o link do plano de mitigação e o fluxo de verificação.
    """

    # Recupera chave da política no MongoDB
    chave_politica = mapa_politicas[nome_politica]

    # Recupera dados da política
    dados_politica = salvaguardas_doc.get(
        chave_politica,
        {}
    )

    # Recupera plano salvo
    dados_plano = dados_politica.get("plano_mitigacao")

    # Não renderiza nada se não existir plano
    if not dados_plano:
        return

    nome_arquivo = dados_plano.get("nome")
    url_arquivo = dados_plano.get("url")

    # Validação de integridade
    if not nome_arquivo or not url_arquivo:
        return

    st.write("")

    # Colunas do plano de mitigação
    col1, col2 = st.columns([3, 2], gap="medium")

    # Coluna 1 — link do arquivo
    with col1:

        st.markdown(
            f"**Plano de mitigação:** [{nome_arquivo}]({url_arquivo})"
        )

    # Coluna 2 — verificação
    with col2:

        # Checkbox apenas para equipe/admin
        if st.session_state.get("tipo_usuario") in ["equipe", "admin"]:

            st.write("")

            # Recupera informação de verificação
            verificado_por = dados_politica.get("verificado_por")

            # Valor inicial do checkbox
            valor_checkbox = bool(verificado_por)

            # Checkbox (desabilitado para visitante)
            checkbox_verificado = st.checkbox(
                "Verificado",
                value=valor_checkbox,
                key=f"checkbox_verificado_{chave_politica}",
                disabled=st.session_state.get("tipo_usuario") == "visitante"
            )



            # Caminho MongoDB
            caminho_politica = f"salvaguardas.{chave_politica}"

            # Salvamento da verificação
            if checkbox_verificado and not verificado_por:

                data_hoje = datetime.datetime.today().strftime("%d/%m/%Y")

                texto_verificacao = (
                    f"{st.session_state.nome} em {data_hoje}"
                )

                col_projetos.update_one(
                    {"codigo": codigo_projeto_atual},
                    {
                        "$set": {
                            f"{caminho_politica}.verificado_por": texto_verificacao
                        }
                    }
                )

                st.rerun()

            # Remoção da verificação
            elif not checkbox_verificado and verificado_por:

                col_projetos.update_one(
                    {"codigo": codigo_projeto_atual},
                    {
                        "$unset": {
                            f"{caminho_politica}.verificado_por": ""
                        }
                    }
                )

                st.rerun()

            # Informação da verificação
            if verificado_por:

                st.caption(
                    f"Verificado por {verificado_por}"
                )



###########################################################################################################
# INTERFACE PRINCIPAL DA PÁGINA
###########################################################################################################



# Logo do sidebar
st.logo("images/ieb_logo.svg", size='large')

# Título da página e identificação
col_titulo, col_identificacao = st.columns([3, 2])

with col_titulo:
    st.header("Salvaguardas")

with col_identificacao:
    st.markdown(
        f"<div style='text-align: right; margin-top: 30px;'>{df_projeto['codigo'].values[0]} - {df_projeto['sigla'].values[0]}</div>",
        unsafe_allow_html=True
    )


st.write('')
st.write('')




# ===============================================================================================
# PERMISSÃO - SALVAGUARDAS
# -----------------------------------------------------------------------------------------------
# Apenas admin/equipe podem editar. Usuários externos apenas visualizam.
# ===============================================================================================

usuario_interno = st.session_state.tipo_usuario in ["admin", "equipe"]

# Usuários internos podem editar, externos apenas visualizar
modo_edicao = usuario_interno


st.subheader("Avaliação de risco")
st.write('')

# Recupera o documento do projeto como dicionário
projeto = df_projeto.iloc[0].to_dict()

# Recupera o objeto salvaguardas do documento do projeto
# Caso não exista, utiliza um dicionário vazio
salvaguardas_doc = projeto.get("salvaguardas", {})


# Extrai os subdocumentos de cada política
pol2 = salvaguardas_doc.get("pol_2_trabalho", {})
pol3 = salvaguardas_doc.get("pol_3_poluicao", {})
pol4 = salvaguardas_doc.get("pol_4_comunidade", {})
pol5 = salvaguardas_doc.get("pol_5_reassentamento", {})
pol6 = salvaguardas_doc.get("pol_6_biodiversidade", {})
pol7 = salvaguardas_doc.get("pol_7_indigenas", {})
pol8 = salvaguardas_doc.get("pol_8_patrimonio", {})
pol9 = salvaguardas_doc.get("pol_9_genero", {})


# Carrega valores existentes no session_state
# Isso faz com que os widgets apareçam já preenchidos ao abrir a página


if "salv_2_aplicavel" not in st.session_state:
    st.session_state["salv_2_aplicavel"] = pol2.get("aplicavel")

if "salv_2_detalhes" not in st.session_state:
    st.session_state["salv_2_detalhes"] = pol2.get("detalhes")

if "salv_2_categoria_risco" not in st.session_state:
    st.session_state["salv_2_categoria_risco"] = pol2.get("categoria")


if "salv_3_aplicavel" not in st.session_state:
    st.session_state["salv_3_aplicavel"] = pol3.get("aplicavel")

if "salv_3_pesticidas_detalhes" not in st.session_state:
    st.session_state["salv_3_pesticidas_detalhes"] = pol3.get("detalhes_pesticidas")

if "salv_3_poluicao_detalhes" not in st.session_state:
    st.session_state["salv_3_poluicao_detalhes"] = pol3.get("detalhes_poluicao")

if "salv_3_categoria_risco" not in st.session_state:
    st.session_state["salv_3_categoria_risco"] = pol3.get("categoria")


if "salv_4_aplicavel" not in st.session_state:
    st.session_state["salv_4_aplicavel"] = pol4.get("aplicavel")

if "salv_4_detalhes" not in st.session_state:
    st.session_state["salv_4_detalhes"] = pol4.get("detalhes")

if "salv_4_categoria_risco" not in st.session_state:
    st.session_state["salv_4_categoria_risco"] = pol4.get("categoria")


if "salv_5_aplicavel" not in st.session_state:
    st.session_state["salv_5_aplicavel"] = pol5.get("aplicavel")

if "salv_5_detalhes" not in st.session_state:
    st.session_state["salv_5_detalhes"] = pol5.get("detalhes")

if "salv_5_categoria_risco" not in st.session_state:
    st.session_state["salv_5_categoria_risco"] = pol5.get("categoria")


if "salv_6_aplicavel" not in st.session_state:
    st.session_state["salv_6_aplicavel"] = pol6.get("aplicavel")

if "salv_6_detalhes" not in st.session_state:
    st.session_state["salv_6_detalhes"] = pol6.get("detalhes")

if "salv_6_categoria_risco" not in st.session_state:
    st.session_state["salv_6_categoria_risco"] = pol6.get("categoria")


if "salv_7_aplicavel" not in st.session_state:
    st.session_state["salv_7_aplicavel"] = pol7.get("aplicavel")

if "salv_7_detalhes" not in st.session_state:
    st.session_state["salv_7_detalhes"] = pol7.get("detalhes")

if "salv_7_categoria_risco" not in st.session_state:
    st.session_state["salv_7_categoria_risco"] = pol7.get("categoria")


if "salv_8_aplicavel" not in st.session_state:
    st.session_state["salv_8_aplicavel"] = pol8.get("aplicavel")

if "salv_8_detalhes" not in st.session_state:
    st.session_state["salv_8_detalhes"] = pol8.get("detalhes")

if "salv_8_categoria_risco" not in st.session_state:
    st.session_state["salv_8_categoria_risco"] = pol8.get("categoria")


if "salv_9_detalhes" not in st.session_state:
    st.session_state["salv_9_detalhes"] = pol9.get("detalhes")

if "salv_9_categoria_risco" not in st.session_state:
    st.session_state["salv_9_categoria_risco"] = pol9.get("categoria")

if "salv_fortalecimento_capacidades" not in st.session_state:
    st.session_state["salv_fortalecimento_capacidades"] = salvaguardas_doc.get("fortalecimento_capacidades")


# Mapeamento entre nome da política e chave do MongoDB
mapa_politicas = {
    "2. Condições de Trabalho e Trabalhistas": "pol_2_trabalho",
    "3. Eficiência de Recursos e Prevenção de Poluição": "pol_3_poluicao",
    "4. Saúde, Segurança e Proteção da Comunidade": "pol_4_comunidade",
    "5. Restrições de Uso da Terra e Reassentamento Involuntário": "pol_5_reassentamento",
    "6. Conservação da Biodiversidade e Gestão Sustentável de Recursos Naturais Vivos": "pol_6_biodiversidade",
    "7. Povos Indígenas": "pol_7_indigenas",
    "8. Patrimônio Cultural": "pol_8_patrimonio",
    "9. Igualdade de Gênero": "pol_9_genero"
}




# COMEÇO DO FORMULÁRIO COM AS POLÍTICAS DE SALVAGUARDAS ###########################################################


# Colunas de informações da avaliação
col1, col2, col3 = st.columns([3, 3, 1])


# # Recupera o nome do usuário logado no session_state
# nome_usuario_atual = st.session_state.get("nome")


# Recupera o nome da pessoa que fez a última avaliação
nome_avaliador = salvaguardas_doc.get("nome_avaliador_risco")

st.write('')

# COLUNA 1
# Mostra apenas se existir informação no banco
if nome_avaliador:
    col1.write(f"**Responsável pela última avaliação:** {nome_avaliador}")



# COLUNA 2
# Recupera a data da última avaliação salva
data_aval_risco = salvaguardas_doc.get("data_aval_risco")

# Mostra a data apenas se existir no banco
if data_aval_risco:
    col2.write(f"**Data da última avaliação:** {data_aval_risco}")


# COLUNA 3
# ===============================================================================================
# EXPORTAÇÃO DO RELATÓRIO
# ===============================================================================================

with col3:

    with st.popover(
        "Exportar",
        icon=":material/print:",
        width="stretch"
    ):

        # Geração do documento
        if st.button(
            "Gerar relatório",
            key="gerar_relatorio_salvaguardas",
            icon=":material/settings:",
            width="stretch"
            # type="primary"
        ):


            # ----------------------------------------------------------------------------------
            # Cria documento Word
            # ----------------------------------------------------------------------------------


            # Cria documento Word
            doc = Document()


            # Recupera organização do projeto
            organizacao = col_organizacoes.find_one(
                {"_id": projeto.get("id_organizacao")}
            )

            # Nome da organização
            nome_organizacao = ""

            if organizacao:

                nome_organizacao = organizacao.get(
                    "nome_organizacao",
                    ""
                )



            # ===============================================================================================
            # TÍTULO PRINCIPAL
            # ===============================================================================================

            doc.add_heading(
                "Avaliação de risco",
                level=1
            )



            # ===============================================================================================
            # SEÇÃO 1 — DETALHES DO PROJETO
            # ===============================================================================================

            doc.add_heading(
                "Seção 1: Detalhes do Projeto",
                level=2
            )

            # Espaçamento após subtítulo
            doc.add_paragraph("")


            # Código do projeto
            paragrafo = doc.add_paragraph()

            run_label = paragrafo.add_run(
                "Código do projeto: "
            )

            run_label.bold = True

            paragrafo.add_run(
                f"{projeto.get('codigo', '')}"
            )


            # Título do projeto
            paragrafo = doc.add_paragraph()

            run_label = paragrafo.add_run(
                "Título do Projeto: "
            )

            run_label.bold = True

            paragrafo.add_run(
                f"{projeto.get('nome_do_projeto', '')}"
            )


            # Organização solicitante
            paragrafo = doc.add_paragraph()

            run_label = paragrafo.add_run(
                "Organização Solicitante: "
            )

            run_label.bold = True

            paragrafo.add_run(
                nome_organizacao
            )


            # Nome do avaliador
            paragrafo = doc.add_paragraph()

            run_label = paragrafo.add_run(
                "Nome da Pessoa que Completa a Avaliação de Risco: "
            )

            run_label.bold = True

            paragrafo.add_run(
                f"{salvaguardas_doc.get('nome_avaliador_risco', '')}"
            )


            # Data da avaliação
            paragrafo = doc.add_paragraph()

            run_label = paragrafo.add_run(
                "Data da Avaliação de Risco: "
            )

            run_label.bold = True

            paragrafo.add_run(
                f"{salvaguardas_doc.get('data_aval_risco', '')}"
            )


            # Espaçamento final da seção
            doc.add_paragraph("")



            # ===============================================================================================
            # SEÇÃO 2 — AVALIAÇÃO DE RISCO
            # ===============================================================================================

            # Espaçamento antes da seção
            doc.add_paragraph("")

            doc.add_heading(
                "Seção 2: Avaliação de Risco",
                level=2
            )

            # Espaçamento após subtítulo
            doc.add_paragraph("")


            # ===============================================================================================
            # TABELA — AVALIAÇÃO DE RISCO
            # ===============================================================================================

            # Cria tabela com cabeçalho
            tabela = doc.add_table(
                rows=1,
                cols=5
            )

            # Estilo da tabela
            tabela.style = "Table Grid"

            # Cabeçalho
            cabecalho = tabela.rows[0].cells


            # Aplica fundo cinza claro no cabeçalho
            for cell in cabecalho:

                cell._tc.get_or_add_tcPr().append(
                    parse_xml(
                        r'<w:shd {} w:fill="D9D9D9"/>'.format(
                            nsdecls('w')
                        )
                    )
                )

            cabecalho[0].text = "Política de Salvaguardas"
            cabecalho[1].text = "Aplicável?"
            cabecalho[2].text = "Avaliação de Risco"
            cabecalho[3].text = "Categoria de Risco"
            cabecalho[4].text = (
                "Notas (os exemplos abaixo são indicativos "
                "e não preveem todas as eventualidades)"
            )


            # ===============================================================================================
            # LINHA 1 — AVALIAÇÃO AMBIENTAL E SOCIAL
            # ===============================================================================================

            linha = tabela.add_row().cells

            linha[0].text = "1. Avaliação Ambiental e Social"

            linha[1].text = (
                "Sim (Aplica-se a todos os projetos)"
            )

            linha[2].text = "N/A"

            linha[3].text = "N/A"

            linha[4].text = (
                "Não é necessário atribuir uma categoria "
                "de risco individual a esta política."
            )

            # Fundo cinza claro nas colunas 2, 3 e 4
            for indice in [1, 2, 3]:

                linha[indice]._tc.get_or_add_tcPr().append(
                    parse_xml(
                        r'<w:shd {} w:fill="D9D9D9"/>'.format(
                            nsdecls('w')
                        )
                    )
                )



            # ===============================================================================================
            # LINHA 2 — CONDIÇÕES DE TRABALHO E TRABALHISTAS
            # ===============================================================================================

            # Recupera dados da política no banco
            dados_pol_2 = salvaguardas_doc.get(
                mapa_politicas["2. Condições de Trabalho e Trabalhistas"],
                {}
            )

            linha = tabela.add_row().cells

            # Coluna 1
            linha[0].text = (
                "2. Condições de Trabalho e Trabalhistas"
            )

            # Coluna 2
            linha[1].text = str(
                dados_pol_2.get("aplicavel", "")
            )

            # Coluna 3
            paragrafo = linha[2].paragraphs[0]

            paragrafo.add_run(
                "O projeto proposto apresenta riscos significativos "
                "em relação às condições de trabalho e trabalhistas?"
            )

            paragrafo2 = linha[2].add_paragraph()

            detalhes = dados_pol_2.get(
                "detalhes",
                ""
            )

            paragrafo2.add_run(
                f"Detalhes: {detalhes}"
            )



            # Coluna 4
            linha[3].text = str(
                dados_pol_2.get("categoria", "")
            )

            # Coluna 5
            linha[4].text = (
                "Projetos geralmente serão atribuídos à Categoria C "
                "para esta política, a menos que existam riscos elevados "
                "relacionados à saúde e segurança ocupacional, como "
                "mergulho ou trabalho como guardas ecológicos."
            )







            # ===============================================================================================
            # LINHA 3 — EFICIÊNCIA DE RECURSOS E PREVENÇÃO DE POLUIÇÃO
            # ===============================================================================================

            # Recupera dados da política no banco
            dados_pol_3 = salvaguardas_doc.get(
                mapa_politicas["3. Eficiência de Recursos e Prevenção de Poluição"],
                {}
            )

            linha = tabela.add_row().cells

            # Coluna 1
            linha[0].text = (
                "3. Eficiência de Recursos e Prevenção de Poluição"
            )

            # Coluna 2
            linha[1].text = str(
                dados_pol_3.get("aplicavel", "")
            )

            # Coluna 3
            paragrafo = linha[2].paragraphs[0]

            paragrafo.add_run(
                "O projeto proposto apresenta riscos significativos "
                "relacionados a pesticidas?"
            )

            detalhes_pesticidas = dados_pol_3.get(
                "detalhes_pesticidas",
                ""
            )

            paragrafo2 = linha[2].add_paragraph()

            paragrafo2.add_run(
                f"Detalhes: {detalhes_pesticidas}"
            )

            paragrafo3 = linha[2].add_paragraph()

            paragrafo3.add_run(
                "O projeto proposto apresenta riscos significativos "
                "relacionados ao uso insustentável de recursos e/ou "
                "formas de poluição que não sejam pesticidas?"
            )

            detalhes_poluicao = dados_pol_3.get(
                "detalhes_poluicao",
                ""
            )

            paragrafo4 = linha[2].add_paragraph()

            paragrafo4.add_run(
                f"Detalhes: {detalhes_poluicao}"
            )

            # Coluna 4
            linha[3].text = str(
                dados_pol_3.get("categoria", "")
            )

            # Coluna 5
            paragrafo = linha[4].paragraphs[0]

            paragrafo.add_run(
                "Projetos que envolvem a aquisição ou uso de "
                "pesticidas químicos serão atribuídos à Categoria B."
            )

            paragrafo2 = linha[4].add_paragraph()

            paragrafo2.add_run(
                "Os projetos com potencial de exposição da comunidade "
                "a materiais e substâncias perigosos liberados por suas "
                "atividades serão classificados na Categoria A ou B."
            )









            # ===============================================================================================
            # LINHA 4 — SAÚDE, SEGURANÇA E PROTEÇÃO DA COMUNIDADE
            # ===============================================================================================

            # Recupera dados da política no banco
            dados_pol_4 = salvaguardas_doc.get(
                mapa_politicas["4. Saúde, Segurança e Proteção da Comunidade"],
                {}
            )

            linha = tabela.add_row().cells

            # Coluna 1
            linha[0].text = (
                "4. Saúde, Segurança e Proteção da Comunidade"
            )

            # Coluna 2
            linha[1].text = str(
                dados_pol_4.get("aplicavel", "")
            )

            # Coluna 3
            paragrafo = linha[2].paragraphs[0]

            paragrafo.add_run(
                "O projeto proposto apresenta riscos significativos "
                "relacionados à saúde, segurança e proteção da comunidade?"
            )

            detalhes = dados_pol_4.get(
                "detalhes",
                ""
            )

            paragrafo2 = linha[2].add_paragraph()

            paragrafo2.add_run(
                f"Detalhes: {detalhes}"
            )

            # Coluna 4
            linha[3].text = str(
                dados_pol_4.get("categoria", "")
            )

            # Coluna 5
            paragrafo = linha[4].paragraphs[0]

            paragrafo.add_run(
                "Projetos que financiam salários e/ou operações de "
                "guardas florestais, guardas ecológicos ou pessoal "
                "de segurança similar (armado ou desarmado) serão "
                "classificados como Categoria B"
            )

            run_superscript = paragrafo.add_run("1")

            run_superscript.font.superscript = True

            paragrafo.add_run(".")

            paragrafo2 = linha[4].add_paragraph()

            paragrafo2.add_run(
                "Projetos que envolvem atividades de pesquisa"
            )

            run_superscript_1 = paragrafo2.add_run("2")

            run_superscript_1.font.superscript = True

            paragrafo2.add_run(
                " com seres humanos"
            )

            run_superscript_2 = paragrafo2.add_run("2")

            run_superscript_2.font.superscript = True

            paragrafo2.add_run(
                " serão classificados como Categoria B."
            )








            # ===============================================================================================
            # LINHA 5 — RESTRIÇÕES DE USO DA TERRA E REASSENTAMENTO INVOLUNTÁRIO
            # ===============================================================================================

            # Recupera dados da política no banco
            dados_pol_5 = salvaguardas_doc.get(
                mapa_politicas["5. Restrições de Uso da Terra e Reassentamento Involuntário"],
                {}
            )

            linha = tabela.add_row().cells

            # Coluna 1
            linha[0].text = (
                "5. Restrições de Uso da Terra e Reassentamento Involuntário"
            )

            # Coluna 2
            linha[1].text = str(
                dados_pol_5.get("aplicavel", "")
            )

            # Coluna 3
            paragrafo = linha[2].paragraphs[0]

            paragrafo.add_run(
                "O projeto proposto apresenta riscos significativos "
                "relacionados a restrições de acesso associadas "
                "a impactos negativos nos meios de subsistência?"
            )

            detalhes = dados_pol_5.get(
                "detalhes",
                ""
            )

            paragrafo2 = linha[2].add_paragraph()

            paragrafo2.add_run(
                f"Detalhes: {detalhes}"
            )

            # Coluna 4
            linha[3].text = str(
                dados_pol_5.get("categoria", "")
            )

            # Coluna 5
            linha[4].text = (
                "Projetos que envolvem a criação ou expansão "
                "de áreas protegidas serão classificados "
                "como Categoria B."
            )








            # ===============================================================================================
            # LINHA 6 — CONSERVAÇÃO DA BIODIVERSIDADE E GESTÃO SUSTENTÁVEL
            # ===============================================================================================

            # Recupera dados da política no banco
            dados_pol_6 = salvaguardas_doc.get(
                mapa_politicas["6. Conservação da Biodiversidade e Gestão Sustentável de Recursos Naturais Vivos"],
                {}
            )

            linha = tabela.add_row().cells

            # Coluna 1
            linha[0].text = (
                "6. Conservação da Biodiversidade e Gestão Sustentável "
                "de Recursos Naturais Vivos"
            )

            # Coluna 2
            linha[1].text = str(
                dados_pol_6.get("aplicavel", "")
            )

            # Coluna 3
            paragrafo = linha[2].paragraphs[0]

            paragrafo.add_run(
                "O projeto proposto apresenta riscos significativos "
                "relacionados à degradação ou perda de habitat crítico "
                "ou outros habitats naturais?"
            )

            detalhes = dados_pol_6.get(
                "detalhes",
                ""
            )

            paragrafo2 = linha[2].add_paragraph()

            paragrafo2.add_run(
                f"Detalhes: {detalhes}"
            )

            # Coluna 4
            linha[3].text = str(
                dados_pol_6.get("categoria", "")
            )

            # Coluna 5
            paragrafo = linha[4].paragraphs[0]

            paragrafo.add_run(
                "Projetos que envolvem impactos adversos "
                "em habitats críticos"
            )

            run_superscript_3 = paragrafo.add_run("3")

            run_superscript_3.font.superscript = True

            paragrafo.add_run(
                " serão classificados como Categoria A."
            )

            paragrafo2 = linha[4].add_paragraph()

            paragrafo2.add_run(
                "Projetos que envolvem a aquisição de commodities "
                "de recursos naturais (por exemplo, madeira) que "
                "possam contribuir para a conversão ou degradação "
                "significativa de habitats naturais serão "
                "classificados como Categoria B."
            )

            paragrafo3 = linha[4].add_paragraph()

            paragrafo3.add_run(
                "Projetos que envolvem a produção ou colheita "
                "de recursos naturais vivos de populações "
                "selvagens de espécies ameaçadas globalmente "
                "serão classificados como Categoria B."
            )








            # ===============================================================================================
            # LINHA 7 — POVOS INDÍGENAS
            # ===============================================================================================

            # Recupera dados da política no banco
            dados_pol_7 = salvaguardas_doc.get(
                mapa_politicas["7. Povos Indígenas"],
                {}
            )

            linha = tabela.add_row().cells

            # Coluna 1
            linha[0].text = (
                "7. Povos Indígenas"
            )

            # Coluna 2
            linha[1].text = str(
                dados_pol_7.get("aplicavel", "")
            )

            # Coluna 3
            paragrafo = linha[2].paragraphs[0]

            paragrafo.add_run(
                "O projeto proposto apresenta riscos significativos "
                "relacionados aos impactos sobre Povos Indígenas?"
            )

            detalhes = dados_pol_7.get(
                "detalhes",
                ""
            )

            paragrafo2 = linha[2].add_paragraph()

            paragrafo2.add_run(
                f"Detalhes: {detalhes}"
            )

            # Coluna 4
            linha[3].text = str(
                dados_pol_7.get("categoria", "")
            )

            # Coluna 5
            paragrafo = linha[4].paragraphs[0]

            paragrafo.add_run(
                "Projetos que possam afetar Povos Indígenas "
                "em isolamento voluntário ou grupos remotos "
                "com contato externo limitado serão "
                "classificados como Categoria A."
            )

            paragrafo2 = linha[4].add_paragraph()

            paragrafo2.add_run(
                "Projetos que envolvam o uso de ou restrições "
                "de acesso a recursos naturais que sejam "
                "centrais para a identidade, cultura e "
                "subsistência dos Povos Indígenas serão "
                "classificados como Categoria B."
            )

            paragrafo3 = linha[4].add_paragraph()

            paragrafo3.add_run(
                "Projetos que envolvam o desenvolvimento "
                "comercial de terras e recursos naturais "
                "centrais para a identidade e subsistência "
                "dos Povos Indígenas ou o uso comercial "
                "de seu patrimônio cultural serão "
                "classificados como Categoria B."
            )





            # ===============================================================================================
            # LINHA 8 — PATRIMÔNIO CULTURAL
            # ===============================================================================================

            # Recupera dados da política no banco
            dados_pol_8 = salvaguardas_doc.get(
                mapa_politicas["8. Patrimônio Cultural"],
                {}
            )

            linha = tabela.add_row().cells

            # Coluna 1
            linha[0].text = (
                "8. Patrimônio Cultural"
            )

            # Coluna 2
            linha[1].text = str(
                dados_pol_8.get("aplicavel", "")
            )

            # Coluna 3
            paragrafo = linha[2].paragraphs[0]

            paragrafo.add_run(
                "O projeto proposto apresenta riscos significativos "
                "relacionados aos impactos sobre o patrimônio cultural "
                "tangível e/ou intangível?"
            )

            detalhes = dados_pol_8.get(
                "detalhes",
                ""
            )

            paragrafo2 = linha[2].add_paragraph()

            paragrafo2.add_run(
                f"Detalhes: {detalhes}"
            )

            # Coluna 4
            linha[3].text = str(
                dados_pol_8.get("categoria", "")
            )

            # Coluna 5
            paragrafo = linha[4].paragraphs[0]

            paragrafo.add_run(
                "Projetos que introduzam restrições ao acesso "
                "das partes interessadas ao patrimônio cultural "
                "serão classificados como Categoria B."
            )

            paragrafo2 = linha[4].add_paragraph()

            paragrafo2.add_run(
                "Projetos que envolvam o uso comercial de "
                "patrimônio cultural serão classificados "
                "como Categoria B."
            )







            # ===============================================================================================
            # LINHA 9 — IGUALDADE DE GÊNERO
            # ===============================================================================================

            # Recupera dados da política no banco
            dados_pol_9 = salvaguardas_doc.get(
                mapa_politicas["9. Igualdade de Gênero"],
                {}
            )

            linha = tabela.add_row().cells

            # Coluna 1
            linha[0].text = (
                "9. Igualdade de Gênero"
            )

            # Coluna 2
            linha[1].text = (
                "Sim (Aplica-se a todos os projetos)"
            )

            # Aplica fundo cinza claro somente na coluna 2
            linha[1]._tc.get_or_add_tcPr().append(
                parse_xml(
                    r'<w:shd {} w:fill="D9D9D9"/>'.format(
                        nsdecls('w')
                    )
                )
            )

            # Coluna 3
            paragrafo = linha[2].paragraphs[0]

            paragrafo.add_run(
                "O projeto proposto apresenta riscos significativos "
                "relacionados a impactos na promoção, proteção e "
                "respeito à igualdade de gênero?"
            )

            detalhes = dados_pol_9.get(
                "detalhes",
                ""
            )

            paragrafo2 = linha[2].add_paragraph()

            paragrafo2.add_run(
                f"Detalhes: {detalhes}"
            )

            # Coluna 4
            linha[3].text = str(
                dados_pol_9.get("categoria", "")
            )

            # Coluna 5
            linha[4].text = (
                "Projetos serão tipicamente atribuídos à "
                "Categoria C para esta política, a menos "
                "que existam riscos elevados de agravamento "
                "de desigualdades existentes relacionadas ao gênero."
            )






            # ===============================================================================================
            # LINHA 10 — ENGAJAMENTO DE PARTES INTERESSADAS
            # ===============================================================================================

            linha = tabela.add_row().cells

            # Coluna 1
            linha[0].text = (
                "10. Engajamento de Partes Interessadas"
            )

            # Coluna 2
            linha[1].text = (
                "Sim (Aplica-se a todos os projetos)"
            )

            # Fundo cinza claro na coluna 2
            linha[1]._tc.get_or_add_tcPr().append(
                parse_xml(
                    r'<w:shd {} w:fill="D9D9D9"/>'.format(
                        nsdecls('w')
                    )
                )
            )

            # Coluna 3
            linha[2].text = "N/A"

            # Fundo cinza claro na coluna 3
            linha[2]._tc.get_or_add_tcPr().append(
                parse_xml(
                    r'<w:shd {} w:fill="D9D9D9"/>'.format(
                        nsdecls('w')
                    )
                )
            )

            # Coluna 4
            linha[3].text = "N/A"

            # Fundo cinza claro na coluna 4
            linha[3]._tc.get_or_add_tcPr().append(
                parse_xml(
                    r'<w:shd {} w:fill="D9D9D9"/>'.format(
                        nsdecls('w')
                    )
                )
            )

            # Coluna 5
            linha[4].text = (
                "Não é necessário atribuir uma categoria "
                "de risco individual a esta política."
            )





            # ===============================================================================================
            # LINHA 11 — SEPARADOR VISUAL
            # ===============================================================================================

            linha = tabela.add_row().cells

            # Aplica fundo cinza escuro em todas as células
            for cell in linha:

                cell.text = ""

                cell._tc.get_or_add_tcPr().append(
                    parse_xml(
                        r'<w:shd {} w:fill="808080"/>'.format(
                            nsdecls('w')
                        )
                    )
                )





            # ===============================================================================================
            # LINHA 12 — CATEGORIA GERAL DE RISCO
            # ===============================================================================================

            linha = tabela.add_row().cells

            # Coluna 1
            linha[0].text = (
                "CATEGORIA GERAL DE RISCO"
            )

            # Coluna 2
            linha[1].text = "N/A"

            # Fundo cinza claro na coluna 2
            linha[1]._tc.get_or_add_tcPr().append(
                parse_xml(
                    r'<w:shd {} w:fill="D9D9D9"/>'.format(
                        nsdecls('w')
                    )
                )
            )

            # Coluna 3
            paragrafo = linha[2].paragraphs[0]

            paragrafo.add_run("N/A")

            paragrafo2 = linha[2].add_paragraph()

            paragrafo2.add_run(
                "[A categoria de risco segue a categoria "
                "de risco mais alta atribuída às "
                "Políticas de Salvaguarda 2-9.]"
            )

            # Fundo cinza claro na coluna 3
            linha[2]._tc.get_or_add_tcPr().append(
                parse_xml(
                    r'<w:shd {} w:fill="D9D9D9"/>'.format(
                        nsdecls('w')
                    )
                )
            )

            # Coluna 4
            linha[3].text = str(
                salvaguardas_doc.get(
                    "categoria_geral_risco",
                    ""
                )
            )

            # Coluna 5
            linha[4].text = (
                "A categoria geral de risco para o projeto "
                "é equivalente à categoria mais alta atribuída "
                "às políticas individuais de salvaguarda."
            )





            # ===============================================================================================
            # LINHA 13 — SEPARADOR VISUAL
            # ===============================================================================================

            linha = tabela.add_row().cells

            # Aplica fundo cinza escuro em todas as células
            for cell in linha:

                cell.text = ""

                cell._tc.get_or_add_tcPr().append(
                    parse_xml(
                        r'<w:shd {} w:fill="808080"/>'.format(
                            nsdecls('w')
                        )
                    )
                )







            # ===============================================================================================
            # LINHA 14 — FORTALECIMENTO DE CAPACIDADE
            # ===============================================================================================

            linha = tabela.add_row().cells

            # Coluna 1
            linha[0].text = (
                "FORTALECIMENTO DE CAPACIDADE"
            )

            # Mescla colunas 2 a 5
            celula_mesclada = linha[1].merge(linha[4])

            # Primeiro parágrafo
            paragrafo = celula_mesclada.paragraphs[0]

            paragrafo.add_run(
                "O solicitante necessita de fortalecimento "
                "de capacidade para gerenciar os riscos "
                "ambientais e sociais identificados aqui? "
                "Se sim, descreva as atividades de "
                "fortalecimento de capacidade que precisam "
                "ser integradas ao desenho do projeto:"
            )

            # Segundo parágrafo
            fortalecimento = salvaguardas_doc.get(
                "fortalecimento_capacidades",
                ""
            )

            paragrafo2 = celula_mesclada.add_paragraph()

            paragrafo2.add_run(
                str(fortalecimento)
            )








            # Aplica formatação geral da tabela
            formatar_tabela_docx(tabela)




            # ===============================================================================================
            # NOTAS
            # ===============================================================================================

            # Espaçamento após tabela
            doc.add_paragraph("")

            doc.add_heading(
                "Notas",
                level=3
            )

            # Nota 1
            paragrafo = doc.add_paragraph()

            run_numero = paragrafo.add_run("1) ")

            run_numero.bold = True

            paragrafo.add_run(
                "Esta disposição aplica-se independentemente "
                "de o pessoal de segurança ser empregado "
                "ou contratado pela entidade beneficiária, "
                "por uma agência governamental ou por um terceiro."
            )

            # Nota 2
            paragrafo = doc.add_paragraph()

            run_numero = paragrafo.add_run("2) ")

            run_numero.bold = True

            paragrafo.add_run(
                "Pesquisa com Seres Humanos refere-se a qualquer "
                "forma de investigação disciplinada que visa "
                "contribuir para um corpo de conhecimento ou teoria "
                "que envolva a obtenção de (a) dados de indivíduos "
                "vivos por meio de intervenção ou interação com "
                "o indivíduo ou (b) informações pessoais identificáveis. "
                "Projetos de demonstração de campo de conservação "
                "e/ou desenvolvimento geralmente não são considerados "
                "pesquisa, nem métodos participativos padrão usados "
                "para monitorar os impactos desses projetos no "
                "bem-estar humano (por exemplo, discussões em "
                "grupos focais)."
            )

            # Nota 3
            paragrafo = doc.add_paragraph()

            run_numero = paragrafo.add_run("3) ")

            run_numero.bold = True

            paragrafo.add_run(
                "Habitats críticos incluem, entre outros, áreas "
                "protegidas existentes e propostas, áreas "
                "reconhecidas como protegidas por comunidades "
                "locais tradicionais, bem como áreas identificadas "
                "como importantes para a conservação, como "
                "Áreas-Chave de Biodiversidade (KBAs), "
                "Sítios da Aliança para Extinção Zero (AZE), "
                "Áreas Importantes para Aves e Biodiversidade (IBAs), "
                "sítios Ramsar, etc."
            )








            # Salva documento em memória
            buffer_docx = BytesIO()

            doc.save(buffer_docx)

            buffer_docx.seek(0)

            # Armazena no session_state
            st.session_state["relatorio_salvaguardas_docx"] = buffer_docx



        # Exibe botão de download após gerar o documento
        if "relatorio_salvaguardas_docx" in st.session_state:

            st.caption(
                "Relatório gerado. Clique para baixar."
            )

            st.download_button(
                "Baixar",
                data=st.session_state["relatorio_salvaguardas_docx"],
                file_name=f"{codigo_projeto_atual} - Quadro de avaliação de risco.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                icon=":material/download:",
                type="primary",
                width="stretch"
            )



st.write("")
st.write("")
st.write("")




# ===============================================================================================
# VISUALIZAÇÃO DE PLANOS DE MITIGAÇÃO - USUÁRIO BENEFICIÁRIO
# -----------------------------------------------------------------------------------------------
# Exibe apenas as políticas com categoria de risco A ou B para upload do plano de mitigação.
# ===============================================================================================

if st.session_state.tipo_usuario == "beneficiario":



    # Lista das políticas de salvaguardas e respectivas categorias de risco
    politicas_risco = [
        {
            "nome": "2. Condições de Trabalho e Trabalhistas",
            "categoria": pol2.get("categoria")
        },
        {
            "nome": "3. Eficiência de Recursos e Prevenção de Poluição",
            "categoria": pol3.get("categoria")
        },
        {
            "nome": "4. Saúde, Segurança e Proteção da Comunidade",
            "categoria": pol4.get("categoria")
        },
        {
            "nome": "5. Restrições de Uso da Terra e Reassentamento Involuntário",
            "categoria": pol5.get("categoria")
        },
        {
            "nome": "6. Conservação da Biodiversidade e Gestão Sustentável de Recursos Naturais Vivos",
            "categoria": pol6.get("categoria")
        },
        {
            "nome": "7. Povos Indígenas",
            "categoria": pol7.get("categoria")
        },
        {
            "nome": "8. Patrimônio Cultural",
            "categoria": pol8.get("categoria")
        },
        {
            "nome": "9. Igualdade de Gênero",
            "categoria": pol9.get("categoria")
        }
    ]

    # Filtra apenas políticas com categoria A ou B
    politicas_ativas = [
        politica
        for politica in politicas_risco
        if politica["categoria"] in ["Categoria A", "Categoria B"]
    ]

    # Mensagem quando não houver salvaguardas ativadas
    if not politicas_ativas:

        st.write("Não ativou nenhuma salvaguarda")

    # Exibe upload para cada política aplicável
    else:

        st.subheader("Planos de mitigação")
        st.write("")

        for politica in politicas_ativas:

            st.write(f"**{politica['nome']}**")

            col1, col2, col3 = st.columns([3,1,1], gap="large")

            with col1:

                st.file_uploader(
                    "Insira o plano de mitigação",
                    key=f"upload_plano_mitigacao_{politica['nome']}"
                )


            with col2:

                # Recupera chave da política atual
                chave_politica = mapa_politicas.get(
                    politica["nome"]
                )

                # Recupera dados da política
                dados_politica = salvaguardas_doc.get(
                    chave_politica,
                    {}
                )

                # Recupera informação de verificação
                verificado_por = dados_politica.get(
                    "verificado_por"
                )

                # Define estado inicial do checkbox
                valor_checkbox = bool(verificado_por)

                # Checkbox somente leitura para beneficiário
                st.checkbox(
                    "Verificado",
                    value=valor_checkbox,
                    key=f"checkbox_verificado_beneficiario_{chave_politica}",
                    disabled=True
                )

                # Exibe informação de verificação
                if verificado_por:

                    st.caption(
                        f"Verificado por {verificado_por}"
                    )



            with col3:

                with st.container(horizontal=True, horizontal_alignment="right"):

                    salvar = st.button(
                        "Salvar",
                        key=f"salvar_plano_mitigacao_{politica['nome']}",
                        icon=":material/save:",
                        type="primary"
                    )



                # Chave única de controle de processamento
                chave_processando = f"processando_upload_{politica['nome']}"

                # Inicializa a flag no session_state
                if chave_processando not in st.session_state:
                    st.session_state[chave_processando] = False


                if salvar and not st.session_state[chave_processando]:

                    # Ativa trava de processamento
                    st.session_state[chave_processando] = True

                    # Recupera o arquivo enviado no uploader correspondente
                    arquivo_upload = st.session_state.get(
                        f"upload_plano_mitigacao_{politica['nome']}"
                    )

                    # Validação de arquivo obrigatório
                    if not arquivo_upload:

                        st.warning("Selecione um arquivo antes de salvar.")

                        # Libera a trava
                        st.session_state[chave_processando] = False

                    else:

                        with st.spinner("Salvando..."):

                            # Conecta ao Google Drive somente no momento do upload
                            servico = obter_servico_drive()

                            # Obtém a pasta principal do projeto
                            pasta_projeto = obter_pasta_projeto(
                                servico,
                                projeto["codigo"],
                                projeto["sigla"]
                            )

                            # Obtém ou cria a pasta de planos de mitigação
                            pasta_planos = obter_pasta_planos_mitigacao(
                                servico,
                                pasta_projeto
                            )

                            # Faz upload do arquivo
                            id_drive = enviar_arquivo_drive(
                                servico,
                                pasta_planos,
                                arquivo_upload
                            )

                            # Valida sucesso do upload
                            if id_drive:

                                # Gera link público do arquivo
                                url_arquivo = gerar_link_drive(id_drive)


                                # Recupera a chave correspondente da política
                                chave_politica = mapa_politicas.get(politica["nome"])


                                # Atualiza o documento no MongoDB
                                col_projetos.update_one(
                                    {"codigo": codigo_projeto_atual},
                                    {
                                        "$set": {
                                            f"salvaguardas.{chave_politica}.plano_mitigacao": {
                                                "nome": arquivo_upload.name,
                                                "url": url_arquivo
                                            }
                                        },

                                        # Remove verificação anterior ao substituir o arquivo
                                        "$unset": {
                                            f"salvaguardas.{chave_politica}.verificado_por": ""
                                        }
                                    }
                                )


                                # Atualiza o documento no MongoDB
                                col_projetos.update_one(
                                    {"codigo": codigo_projeto_atual},
                                    {
                                        "$set": {
                                            f"salvaguardas.{chave_politica}.plano_mitigacao": {
                                                "nome": arquivo_upload.name,
                                                "url": url_arquivo
                                            }
                                        }
                                    }
                                )

                                st.success(
                                    "Plano de mitigação salvo com sucesso!",
                                    icon=":material/check:"
                                )

                                # Libera a trava antes do rerun
                                st.session_state[chave_processando] = False

                                time.sleep(3)
                                st.rerun()

                            else:

                                # Libera a trava em caso de erro
                                st.session_state[chave_processando] = False



            # Lista o link do arquivo já salvo, se houver

            # Identifica a chave da política atual
            chave_politica = mapa_politicas.get(politica["nome"])

            # Recupera os dados salvos no banco
            dados_plano = salvaguardas_doc.get(chave_politica, {}).get("plano_mitigacao")

            # Exibe link do arquivo salvo
            if dados_plano:

                nome_arquivo = dados_plano.get("nome")
                url_arquivo = dados_plano.get("url")

                if nome_arquivo and url_arquivo:

                    st.markdown(
                        f"[{nome_arquivo}]({url_arquivo})"
                    )


            st.divider()




# Se não for beneficiário

else:





    # Define a largura padrão das colunas que será reutilizada na tabela de salvaguardas
    largura_colunas = [2, 2, 3, 2, 3]

    # Cria a primeira linha de colunas (cabeçalho)
    cab1, cab2, cab3, cab4, cab5 = st.columns(largura_colunas)

    # Cabeçalhos em negrito
    cab1.markdown("**Política de Salvaguardas**")
    cab2.markdown("**Aplicável?**")
    cab3.markdown("**Avaliação de Risco**")
    cab4.markdown("**Categoria de Risco**")
    cab5.markdown("**Notas (os exemplos abaixo são indicativos e não prevêem todas as eventualidades)**")

    st.divider()


    # colunas das linhas de pergutnas e respostas
    col1, col2, col3, col4, col5 = st.columns(largura_colunas, gap="medium")


    # 1. Avaliação Ambiental e Social ------------------------------------------------------------------------------------------------

    # Coluna 1 — nome da política
    col1.write("**1. Avaliação Ambiental e Social**")

    # Coluna 2 — política sempre aplicável
    with col2:

        # Mostra o valor fixo para o usuário
        st.write("Sim")

        # Observação da política
        st.caption("Aplica-se a todos os projetos")

        # Define o valor no session_state para garantir que seja salvo
        st.session_state["salv_1_aplicavel"] = "Sim"

    # Coluna 3 — avaliação de risco
    col3.write("N/A")

    # Coluna 4 — categoria de risco
    col4.write("N/A")

    # Coluna 5 — observação
    col5.write("Não é necessário atribuir uma categoria de risco individual a esta política.")


    st.divider()




    # 2. Condições de Trabalho e Trabalhistas -------------------------------------------------------------

    # colunas da segunda linha de perguntas e respostas
    col1, col2, col3, col4, col5 = st.columns(largura_colunas, gap="medium")


    # Coluna 1 — nome da política
    col1.write("**2. Condições de Trabalho e Trabalhistas**")

    # Coluna 2 — pergunta se a política é aplicável
    with col2:
        st.radio(
            "Aplicável?",
            ["Sim", "Não"],
            key="salv_2_aplicavel",
            horizontal=True,
            disabled=not modo_edicao
        )

    # Coluna 3 — pergunta de avaliação de risco e campo para detalhes
    with col3:
        st.write("O projeto proposto apresenta riscos significativos em relação às condições de trabalho e trabalhistas?")
        st.text_area(
            "Detalhes:",
            key="salv_2_detalhes",
            height=120,
            disabled=not modo_edicao
        )

    # Coluna 4 — categoria de risco
    with col4:
        st.radio(
            "Categoria de risco",
            ["Categoria A", "Categoria B", "Categoria C"],
            key="salv_2_categoria_risco",
            disabled=not modo_edicao
        )

    # Coluna 5 — observação explicativa
    col5.write(
        "Projetos geralmente serão atribuídos à Categoria C para esta política, "
        "a menos que existam riscos elevados relacionados à saúde e segurança "
        "ocupacional, como mergulho ou trabalho como guardas ecológicos."
    )

    # Renderiza o link do plano de mitigação e o checkbox de verificado apenas para admin e equipe
    renderizar_plano_mitigacao(
        salvaguardas_doc,
        mapa_politicas,
        "2. Condições de Trabalho e Trabalhistas",
        codigo_projeto_atual,
        col_projetos
    )

    st.divider()








    # 3. Eficiência de Recursos e Prevenção de Poluição --------------------------------------------------


    # colunas da terceira linha de perguntas e respostas
    col1, col2, col3, col4, col5 = st.columns(largura_colunas, gap="medium")


    # Coluna 1 — nome da política
    col1.write("**3. Eficiência de Recursos e Prevenção de Poluição**")

    # Coluna 2 — pergunta se a política é aplicável
    with col2:
        st.radio(
            "Aplicável?",
            ["Sim", "Não"],
            key="salv_3_aplicavel",
            horizontal=True,
            disabled=not modo_edicao
        )

    # Coluna 3 — perguntas de avaliação de risco com campos de detalhe
    with col3:

        st.write("O projeto proposto apresenta riscos significativos relacionados a pesticidas?")

        st.text_area(
            "Detalhes:",
            key="salv_3_pesticidas_detalhes",
            height=100,
            disabled=not modo_edicao
        )

        st.write("O projeto proposto apresenta riscos significativos relacionados ao uso insustentável de recursos e/ou formas de poluição que não sejam pesticidas?")

        st.text_area(
            "Detalhes:",
            key="salv_3_poluicao_detalhes",
            height=100,
            disabled=not modo_edicao
        )

    # Coluna 4 — categoria de risco
    with col4:
        st.radio(
            "Categoria de risco",
            ["Categoria A", "Categoria B", "Categoria C"],
            key="salv_3_categoria_risco",
            disabled=not modo_edicao
        )

    # Coluna 5 — observações explicativas
    with col5:
        st.write(
            "Projetos que envolvem a aquisição ou uso de pesticidas químicos serão atribuídos à Categoria B."
        )

        st.write(
            "Os projetos com potencial de exposição da comunidade a materiais e substâncias perigosos "
            "liberados por suas atividades serão classificados na Categoria A ou B."
        )



    renderizar_plano_mitigacao(
        salvaguardas_doc,
        mapa_politicas,
        "3. Eficiência de Recursos e Prevenção de Poluição",
        codigo_projeto_atual,
        col_projetos
    )

    st.divider()





    # 4. Saúde, Segurança e Proteção da Comunidade -------------------------------------------------------

    # colunas da quarta linha de perguntas e respostas
    col1, col2, col3, col4, col5 = st.columns(largura_colunas, gap="medium")



    # Coluna 1 — nome da política
    col1.write("**4. Saúde, Segurança e Proteção da Comunidade**")

    # Coluna 2 — pergunta se a política é aplicável
    with col2:
        st.radio(
            "Aplicável?",
            ["Sim", "Não"],
            key="salv_4_aplicavel",
            horizontal=True,
            disabled=not modo_edicao
        )

    # Coluna 3 — pergunta de avaliação de risco com campo para detalhes
    with col3:

        st.write(
            "O projeto proposto apresenta riscos significativos relacionados "
            "à saúde, segurança e proteção da comunidade?"
        )

        st.text_area(
            "Detalhes:",
            key="salv_4_detalhes",
            height=120,
            disabled=not modo_edicao
        )

    # Coluna 4 — categoria de risco
    with col4:
        st.radio(
            "Categoria de risco",
            ["Categoria A", "Categoria B", "Categoria C"],
            key="salv_4_categoria_risco",
            disabled=not modo_edicao
        )

    # Coluna 5 — observações explicativas
    with col5:

        st.markdown(
            "Projetos que financiam salários e/ou operações de guardas florestais, "
            "guardas ecológicos ou pessoal de segurança similar (armado ou desarmado) "
            "serão classificados como Categoria B<sup>1</sup>.",
            unsafe_allow_html=True
        )


        st.markdown(
            "Projetos que envolvem atividades de pesquisa<sup>2</sup> com seres humanos<sup>2</sup> "
            "serão classificados como Categoria B.",
            unsafe_allow_html=True
        )

    st.write('')

    # Nota de rodapé 1
    st.caption(
        "<sup>1</sup> Esta disposição aplica-se independentemente de o pessoal de segurança ser empregado "
        "ou contratado pela entidade beneficiária, por uma agência governamental ou por um terceiro.",
        unsafe_allow_html=True
    )

    # Nota de rodapé 2
    st.caption(
        "<sup>2</sup> Pesquisa com Seres Humanos refere-se a qualquer forma de investigação disciplinada "
        "que visa contribuir para um corpo de conhecimento ou teoria que envolva a obtenção de "
        "(a) dados de indivíduos vivos por meio de intervenção ou interação com o indivíduo ou "
        "(b) informações pessoais identificáveis. Projetos de demonstração de campo de conservação "
        "e/ou desenvolvimento geralmente não são considerados pesquisa, nem métodos participativos "
        "padrão usados para monitorar os impactos desses projetos no bem-estar humano "
        "(por exemplo, discussões em grupos focais).",
        unsafe_allow_html=True
    )




    renderizar_plano_mitigacao(
        salvaguardas_doc,
        mapa_politicas,
        "4. Saúde, Segurança e Proteção da Comunidade",
        codigo_projeto_atual,
        col_projetos
    )

    st.divider()






    # 5. Restrições de Uso da Terra e Reassentamento Involuntário ----------------------------------------

    col1, col2, col3, col4, col5 = st.columns(largura_colunas, gap="medium")

    col1.write("**5. Restrições de Uso da Terra e Reassentamento Involuntário**")

    with col2:
        st.radio(
            "Aplicável?",
            ["Sim", "Não"],
            key="salv_5_aplicavel",
            horizontal=True,
            disabled=not modo_edicao
        )

    with col3:
        st.write(
            "O projeto proposto apresenta riscos significativos relacionados a restrições "
            "de acesso associadas a impactos negativos nos meios de subsistência?"
        )

        st.text_area(
            "Detalhes:",
            key="salv_5_detalhes",
            height=120,
            disabled=not modo_edicao
        )

    with col4:
        st.radio(
            "Categoria de risco",
            ["Categoria A", "Categoria B", "Categoria C"],
            key="salv_5_categoria_risco",
            disabled=not modo_edicao
        )

    with col5:
        st.write(
            "Projetos que envolvem a criação ou expansão de áreas protegidas "
            "serão classificados como Categoria B."
        )



    renderizar_plano_mitigacao(
        salvaguardas_doc,
        mapa_politicas,
        "5. Restrições de Uso da Terra e Reassentamento Involuntário",
        codigo_projeto_atual,
        col_projetos
    )

    st.divider()




    # 6. Conservação da Biodiversidade e Gestão Sustentável de Recursos Naturais Vivos -------------------

    col1, col2, col3, col4, col5 = st.columns(largura_colunas, gap="medium")

    col1.write("**6. Conservação da Biodiversidade e Gestão Sustentável de Recursos Naturais Vivos**")

    with col2:
        st.radio(
            "Aplicável?",
            ["Sim", "Não"],
            key="salv_6_aplicavel",
            horizontal=True,
            disabled=not modo_edicao
        )

    with col3:
        st.write(
            "O projeto proposto apresenta riscos significativos relacionados à "
            "degradação ou perda de habitat crítico ou outros habitats naturais?"
        )

        st.text_area(
            "Detalhes:",
            key="salv_6_detalhes",
            height=120,
            disabled=not modo_edicao
        )

    with col4:
        st.radio(
            "Categoria de risco",
            ["Categoria A", "Categoria B", "Categoria C"],
            key="salv_6_categoria_risco",
            disabled=not modo_edicao
        )

    with col5:
        
        st.markdown(
            "Projetos que envolvem impactos adversos em habitats críticos<sup>3</sup> "
            "serão classificados como Categoria A.",
            unsafe_allow_html=True
        )

        st.write(
            "Projetos que envolvem a aquisição de commodities de recursos naturais "
            "(por exemplo, madeira) que possam contribuir para a conversão ou "
            "degradação significativa de habitats naturais serão classificados como Categoria B."
        )

        st.write(
            "Projetos que envolvem a produção ou colheita de recursos naturais vivos "
            "de populações selvagens de espécies ameaçadas globalmente serão classificados como Categoria B."
        )

    # Nota de rodapé 3
    st.caption(
        "<sup>3</sup> Habitats críticos incluem, entre outros, áreas protegidas existentes "
        "e propostas, áreas reconhecidas como protegidas por comunidades locais tradicionais, "
        "bem como áreas identificadas como importantes para a conservação, como Áreas-Chave "
        "de Biodiversidade (KBAs), Sítios da Aliança para Extinção Zero (AZE), Áreas "
        "Importantes para Aves e Biodiversidade (IBAs), sítios Ramsar, etc.",
        unsafe_allow_html=True
    )




    renderizar_plano_mitigacao(
        salvaguardas_doc,
        mapa_politicas,
        "6. Conservação da Biodiversidade e Gestão Sustentável de Recursos Naturais Vivos",
        codigo_projeto_atual,
        col_projetos
    )

    st.divider()






    # 7. Povos Indígenas ----------------------------------------------------------------------------------

    col1, col2, col3, col4, col5 = st.columns(largura_colunas, gap="medium")

    col1.write("**7. Povos Indígenas**")

    with col2:
        st.radio(
            "Aplicável?",
            ["Sim", "Não"],
            key="salv_7_aplicavel",
            horizontal=True,
            disabled=not modo_edicao
        )

    with col3:
        st.write(
            "O projeto proposto apresenta riscos significativos relacionados "
            "aos impactos sobre Povos Indígenas?"
        )

        st.text_area(
            "Detalhes:",
            key="salv_7_detalhes",
            height=120,
            disabled=not modo_edicao
        )

    with col4:
        st.radio(
            "Categoria de risco",
            ["Categoria A", "Categoria B", "Categoria C"],
            key="salv_7_categoria_risco",
            disabled=not modo_edicao
        )

    with col5:
        st.write(
            "Projetos que possam afetar Povos Indígenas em isolamento voluntário "
            "ou grupos remotos com contato externo limitado serão classificados como Categoria A."
        )

        st.write(
            "Projetos que envolvam o uso de ou restrições de acesso a recursos naturais "
            "que sejam centrais para a identidade, cultura e subsistência dos Povos Indígenas "
            "serão classificados como Categoria B."
        )

        st.write(
            "Projetos que envolvam o desenvolvimento comercial de terras e recursos naturais "
            "centrais para a identidade e subsistência dos Povos Indígenas ou o uso comercial "
            "de seu patrimônio cultural serão classificados como Categoria B."
        )



    renderizar_plano_mitigacao(
        salvaguardas_doc,
        mapa_politicas,
        "7. Povos Indígenas",
        codigo_projeto_atual,
        col_projetos
    )

    st.divider()






    # 8. Patrimônio Cultural -------------------------------------------------------------------------------

    col1, col2, col3, col4, col5 = st.columns(largura_colunas, gap="medium")

    col1.write("**8. Patrimônio Cultural**")

    with col2:
        st.radio(
            "Aplicável?",
            ["Sim", "Não"],
            key="salv_8_aplicavel",
            horizontal=True,
            disabled=not modo_edicao
        )

    with col3:
        st.write(
            "O projeto proposto apresenta riscos significativos relacionados "
            "aos impactos sobre o patrimônio cultural tangível e/ou intangível?"
        )

        st.text_area(
            "Detalhes:",
            key="salv_8_detalhes",
            height=120,
            disabled=not modo_edicao
        )

    with col4:
        st.radio(
            "Categoria de risco",
            ["Categoria A", "Categoria B", "Categoria C"],
            key="salv_8_categoria_risco",
            disabled=not modo_edicao
        )

    with col5:
        st.write(
            "Projetos que introduzam restrições ao acesso das partes interessadas "
            "ao patrimônio cultural serão classificados como Categoria B."
        )

        st.write(
            "Projetos que envolvam o uso comercial de patrimônio cultural "
            "serão classificados como Categoria B."
        )




    renderizar_plano_mitigacao(
        salvaguardas_doc,
        mapa_politicas,
        "8. Patrimônio Cultural",
        codigo_projeto_atual,
        col_projetos
    )

    st.divider()






    # 9. Igualdade de Gênero -------------------------------------------------------------------------------

    col1, col2, col3, col4, col5 = st.columns(largura_colunas, gap="medium")

    col1.write("**9. Igualdade de Gênero**")

    # Coluna 2 — política sempre aplicável
    with col2:

        st.write("Sim")
        st.caption("Aplica-se a todos os projetos")

        # Mantém o valor consistente para o salvamento
        st.session_state["salv_9_aplicavel"] = "Sim"



    with col3:
        st.write(
            "O projeto proposto apresenta riscos significativos relacionados "
            "a impactos na promoção, proteção e respeito à igualdade de gênero?"
        )

        st.text_area(
            "Detalhes:",
            key="salv_9_detalhes",
            height=120,
            disabled=not modo_edicao
        )

    with col4:
        st.radio(
            "Categoria de risco",
            ["Categoria A", "Categoria B", "Categoria C"],
            key="salv_9_categoria_risco",
            disabled=not modo_edicao
        )

    with col5:
        st.write(
            "Projetos serão tipicamente atribuídos à Categoria C para esta política, "
            "a menos que existam riscos elevados de agravamento de desigualdades "
            "existentes relacionadas ao gênero."
        )



    renderizar_plano_mitigacao(
        salvaguardas_doc,
        mapa_politicas,
        "9. Igualdade de Gênero",
        codigo_projeto_atual,
        col_projetos
    )

    st.divider()





    # 10. Engajamento de Partes Interessadas ---------------------------------------------------------------

    col1, col2, col3, col4, col5 = st.columns(largura_colunas, gap="medium")

    col1.write("**10. Engajamento de Partes Interessadas**")

    with col2:
        st.write("Sim")
        st.caption("Aplica-se a todos os projetos")

        # Mantém o valor consistente para o salvamento
        st.session_state["salv_10_aplicavel"] = "Sim"

    col3.write("N/A")
    col4.write("N/A")

    col5.write(
        "Não é necessário atribuir uma categoria de risco individual a esta política."
    )

    st.divider()









    # CATEGORIA GERAL DE RISCO ---------------------------------------------------------------------------

    col1, col2, col3, col4, col5 = st.columns(largura_colunas, gap="medium")

    # Coluna 1
    col1.write("**CATEGORIA GERAL DE RISCO**")

    # Coluna 2
    col2.write("N/A")

    # Coluna 3
    col3.write("N/A")


    # Lê as categorias escolhidas nas perguntas 2 a 9

    categorias_filtradas = []

    for i in range(2, 10):  # perguntas 2 a 9

        aplicavel = st.session_state.get(f"salv_{i}_aplicavel")
        categoria = st.session_state.get(f"salv_{i}_categoria_risco")

        if aplicavel == "Sim" and categoria:
            categorias_filtradas.append(categoria)


    # Inicializa a categoria geral
    categoria_geral = ""

    # Determina automaticamente a categoria geral
    # prioridade: A > B > C

    if "Categoria A" in categorias_filtradas:
        categoria_geral = "Categoria A"
    elif "Categoria B" in categorias_filtradas:
        categoria_geral = "Categoria B"
    elif "Categoria C" in categorias_filtradas:
        categoria_geral = "Categoria C"




    col4.write("Resultado final:")

    if categoria_geral:
        col4.write(f"**{categoria_geral}**")
    else:
        col4.write("")


    # Coluna 5
    col5.write(
        "A categoria geral de risco para o projeto é equivalente à categoria mais alta "
        "atribuída às políticas individuais de salvaguarda."
    )

    st.divider()



    # FORTALECIMENTO DE CAPACIDADE ---------------------------------------------------------------------------

    largura_colunas = [2, 2, 3, 2, 3]

    col1, col2 = st.columns([2, 11], gap="medium")


    col1.write("**FORTALECIMENTO DE CAPACIDADE**")

    col2.write("O solicitante necessita de fortalecimento de capacidade para gerenciar os riscos ambientais e sociais identificados aqui?")

    col2.write("Se sim, descreva as atividades de fortalecimento de capacidade que precisam ser integradas ao desenho do projeto:")

    col2.text_area(
        "Detalhes:",
        key="salv_fortalecimento_capacidades",
        height=120,
        disabled=not modo_edicao
    )

    st.divider()








    st.write("")


    # Botão somente para equipe e adimn

    if st.session_state.get("tipo_usuario") in ["equipe", "admin"]:
        

        # Botão para salvar as respostas no banco
        if st.button("Salvar", icon=":material/save:", width=200, type="primary"):

            # Data da avaliação
            data_avaliacao = datetime.datetime.today().strftime("%d/%m/%Y")

            # Nome do avaliador
            avaliador_atual = st.session_state.get("nome", "Usuário")

            # Estrutura organizada das respostas de salvaguardas
            dados_salvaguardas = {

                "nome_avaliador_risco": avaliador_atual,
                "data_aval_risco": data_avaliacao,
                "categoria_geral_risco": categoria_geral,

                "pol_2_trabalho": {
                    "aplicavel": st.session_state.get("salv_2_aplicavel"),
                    "detalhes": st.session_state.get("salv_2_detalhes"),
                    "categoria": st.session_state.get("salv_2_categoria_risco")
                },

                "pol_3_poluicao": {
                    "aplicavel": st.session_state.get("salv_3_aplicavel"),
                    "detalhes_pesticidas": st.session_state.get("salv_3_pesticidas_detalhes"),
                    "detalhes_poluicao": st.session_state.get("salv_3_poluicao_detalhes"),
                    "categoria": st.session_state.get("salv_3_categoria_risco")
                },

                "pol_4_comunidade": {
                    "aplicavel": st.session_state.get("salv_4_aplicavel"),
                    "detalhes": st.session_state.get("salv_4_detalhes"),
                    "categoria": st.session_state.get("salv_4_categoria_risco")
                },

                "pol_5_reassentamento": {
                    "aplicavel": st.session_state.get("salv_5_aplicavel"),
                    "detalhes": st.session_state.get("salv_5_detalhes"),
                    "categoria": st.session_state.get("salv_5_categoria_risco")
                },

                "pol_6_biodiversidade": {
                    "aplicavel": st.session_state.get("salv_6_aplicavel"),
                    "detalhes": st.session_state.get("salv_6_detalhes"),
                    "categoria": st.session_state.get("salv_6_categoria_risco")
                },

                "pol_7_indigenas": {
                    "aplicavel": st.session_state.get("salv_7_aplicavel"),
                    "detalhes": st.session_state.get("salv_7_detalhes"),
                    "categoria": st.session_state.get("salv_7_categoria_risco")
                },

                "pol_8_patrimonio": {
                    "aplicavel": st.session_state.get("salv_8_aplicavel"),
                    "detalhes": st.session_state.get("salv_8_detalhes"),
                    "categoria": st.session_state.get("salv_8_categoria_risco")
                },

                "pol_9_genero": {
                    "aplicavel": st.session_state.get("salv_9_aplicavel"),
                    "detalhes": st.session_state.get("salv_9_detalhes"),
                    "categoria": st.session_state.get("salv_9_categoria_risco")
                },

                "fortalecimento_capacidades": st.session_state.get("salv_fortalecimento_capacidades"),
            }

            # Atualiza o documento do projeto no MongoDB
            


            ###########################################################################################################
            # FUNÇÃO - CONVERTE DICIONÁRIO EM DOT NOTATION PARA MONGODB
            ###########################################################################################################

            def flatten_dict(d, parent_key=""):
                """
                Converte dicionários aninhados em estrutura dot notation.

                Exemplo:
                {"a": {"b": 1}}

                vira:

                {"a.b": 1}
                """

                items = []

                for k, v in d.items():

                    new_key = f"{parent_key}.{k}" if parent_key else k

                    if isinstance(v, dict):

                        items.extend(
                            flatten_dict(v, new_key).items()
                        )

                    else:

                        items.append((new_key, v))

                return dict(items)


            # Converte estrutura aninhada para dot notation
            dados_flatten = flatten_dict(dados_salvaguardas)

            # Adiciona prefixo "salvaguardas."
            # Ignora campos None para evitar sobrescrever dados existentes
            dados_update = {
                f"salvaguardas.{k}": v
                for k, v in dados_flatten.items()
                if v is not None
            }

            # Atualiza apenas os campos editáveis
            resultado = col_projetos.update_one(
                {"codigo": codigo_projeto_atual},
                {
                    "$set": dados_update
                }
            )
            

            # Mostra mensagem de sucesso
            if resultado.modified_count >= 0:
                st.success("Respostas salvas com sucesso!", icon=":material/check:")
                time.sleep(3)
                st.rerun()





